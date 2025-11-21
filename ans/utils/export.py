"""
Export functionality for the ANS application.

This module provides functions to export story content to various formats
(DOCX and PDF) with proper formatting.
"""
import os
import datetime
from typing import Optional

from ans.utils.constants import (
    PROJECT_FILES,
    FILE_ENCODING,
    TIMESTAMP_FORMAT
)

# Optional imports for export functionality
try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

try:
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.lib.colors import HexColor
    from reportlab.platypus import SimpleDocTemplate, Paragraph, PageBreak, Spacer
    from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY
    HAS_REPORTLAB = True
except ImportError:
    HAS_REPORTLAB = False


def export_story_to_docx(project_path: str, output_filename: Optional[str] = None) -> bool:
    """Export story.txt to Word document with formatting.
    
    Args:
        project_path: Path to project folder
        output_filename: Optional custom output filename (without extension)
    
    Returns:
        bool: True if successful, False otherwise
    """
    if not HAS_DOCX:
        return False
    
    try:
        story_path = os.path.join(project_path, PROJECT_FILES['story'])
        config_path = os.path.join(project_path, PROJECT_FILES['config'])
        
        if not os.path.exists(story_path):
            return False
        
        # Read story content
        with open(story_path, 'r', encoding=FILE_ENCODING) as f:
            story_content = f.read()
        
        # Read project config for title
        project_name = os.path.basename(project_path)
        title = f"Novel: {project_name}"
        
        if os.path.exists(config_path):
            with open(config_path, 'r', encoding=FILE_ENCODING) as f:
                for line in f:
                    if line.startswith('Idea:'):
                        title = line.split(':', 1)[1].strip()
                        break
        
        # Create Word document
        doc = Document()
        
        # Add title
        title_para = doc.add_paragraph()
        title_run = title_para.add_run(title)
        title_run.font.size = Pt(24)
        title_run.font.bold = True
        title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # Add metadata
        doc.add_paragraph(f"Project: {project_name}")
        doc.add_paragraph(f"Generated: {datetime.datetime.now().strftime(TIMESTAMP_FORMAT)}")
        
        doc.add_paragraph()  # Blank line
        
        # Parse and add content with chapter headings
        lines = story_content.split('\n')
        for line in lines:
            line = line.rstrip()
            
            # Check if line is a chapter heading
            if line.startswith('=== CHAPTER'):
                # Add chapter heading
                heading = doc.add_heading(line.strip('=').strip(), level=1)
                heading.paragraph_format.space_before = Pt(12)
                heading.paragraph_format.space_after = Pt(12)
            elif line.strip():
                # Add regular paragraph with justified alignment
                para = doc.add_paragraph(line)
                para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            else:
                # Preserve blank lines as spacing
                doc.add_paragraph()
        
        # Save document
        if not output_filename:
            output_filename = f"{project_name}_novel"
        
        output_path = os.path.join(project_path, f"{output_filename}.docx")
        doc.save(output_path)
        
        return True
    
    except Exception as e:
        return False


def export_story_to_pdf(project_path: str, output_filename: Optional[str] = None) -> bool:
    """Export story.txt to PDF using reportlab.
    
    Args:
        project_path: Path to project folder
        output_filename: Optional custom output filename (without extension)
    
    Returns:
        bool: True if successful, False otherwise
    """
    if not HAS_REPORTLAB:
        return False
    
    try:
        story_path = os.path.join(project_path, PROJECT_FILES['story'])
        config_path = os.path.join(project_path, PROJECT_FILES['config'])
        
        if not os.path.exists(story_path):
            return False
        
        # Read story content
        with open(story_path, 'r', encoding=FILE_ENCODING) as f:
            story_content = f.read()
        
        # Read project config for title
        project_name = os.path.basename(project_path)
        title = f"Novel: {project_name}"
        
        if os.path.exists(config_path):
            with open(config_path, 'r', encoding=FILE_ENCODING) as f:
                for line in f:
                    if line.startswith('Idea:'):
                        title = line.split(':', 1)[1].strip()
                        break
        
        # Prepare output path
        if not output_filename:
            output_filename = f"{project_name}_novel"
        
        output_path = os.path.join(project_path, f"{output_filename}.pdf")
        
        # Create PDF
        doc = SimpleDocTemplate(output_path, pagesize=letter)
        
        # Define styles
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=24,
            textColor=HexColor('000000'),
            spaceAfter=6,
            alignment=TA_CENTER,
            fontName='Helvetica-Bold'
        )
        heading_style = ParagraphStyle(
            'ChapterHeading',
            parent=styles['Heading2'],
            fontSize=14,
            textColor=HexColor('000000'),
            spaceAfter=6,
            spaceBefore=12,
            fontName='Helvetica-Bold'
        )
        body_style = ParagraphStyle(
            'CustomBody',
            parent=styles['BodyText'],
            fontSize=11,
            alignment=TA_JUSTIFY,
            spaceAfter=6,
            leading=14
        )
        
        # Build story elements
        story_elements = []
        
        # Add title
        story_elements.append(Paragraph(title, title_style))
        story_elements.append(Paragraph(f"Project: {project_name}", styles['Normal']))
        story_elements.append(Paragraph(f"Generated: {datetime.datetime.now().strftime(TIMESTAMP_FORMAT)}", styles['Normal']))
        story_elements.append(Spacer(1, 0.3*inch))
        
        # Parse and add content
        lines = story_content.split('\n')
        for line in lines:
            line = line.rstrip()
            
            # Check if line is a chapter heading
            if line.startswith('=== CHAPTER'):
                story_elements.append(Paragraph(line.strip('=').strip(), heading_style))
            elif line.strip():
                story_elements.append(Paragraph(line, body_style))
            else:
                story_elements.append(Spacer(1, 0.1*inch))
        
        # Build PDF
        doc.build(story_elements)
        
        return True
    
    except Exception as e:
        return False


def has_docx_support() -> bool:
    """Check if DOCX export is available."""
    return HAS_DOCX


def has_pdf_support() -> bool:
    """Check if PDF export is available."""
    return HAS_REPORTLAB
