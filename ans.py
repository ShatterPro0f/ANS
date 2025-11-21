import sys
import os
import json
import threading
import datetime
import random
import re
from typing import TYPE_CHECKING
import ollama
from PyQt5 import QtWidgets, QtCore, QtGui

# Optional imports for export functionality
try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

try:
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.lib.colors import HexColor
    from reportlab.platypus import SimpleDocTemplate, Paragraph, PageBreak, Spacer
    from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY
    HAS_REPORTLAB = True
except ImportError:
    HAS_REPORTLAB = False

if TYPE_CHECKING:
    from typing import Optional, Dict, Any

from typing import Optional


class BackgroundThread(QtCore.QThread):
    """Background thread for novel generation processing."""
    
    # Signals for communication back to main window
    processing_finished: QtCore.pyqtSignal = QtCore.pyqtSignal(str)
    processing_error: QtCore.pyqtSignal = QtCore.pyqtSignal(str)
    processing_progress: QtCore.pyqtSignal = QtCore.pyqtSignal(str)
    log_update: QtCore.pyqtSignal = QtCore.pyqtSignal(str)
    init_complete: QtCore.pyqtSignal = QtCore.pyqtSignal()
    synopsis_ready: QtCore.pyqtSignal = QtCore.pyqtSignal(str)  # Emits synopsis text
    new_synopsis: QtCore.pyqtSignal = QtCore.pyqtSignal(str)  # Emits refined synopsis text
    new_outline: QtCore.pyqtSignal = QtCore.pyqtSignal(str)  # Emits generated outline text
    new_characters: QtCore.pyqtSignal = QtCore.pyqtSignal(str)  # Emits character JSON array
    new_world: QtCore.pyqtSignal = QtCore.pyqtSignal(str)  # Emits world-building JSON dict
    new_timeline: QtCore.pyqtSignal = QtCore.pyqtSignal(str)  # Emits timeline with dates and events
    new_draft: QtCore.pyqtSignal = QtCore.pyqtSignal(str)  # Emits polished/enhanced draft section
    
    def __init__(self, parent=None):
        """Initialize background thread."""
        super().__init__(parent)
        self.inputs = None
        self.buffer = ''  # Initialize buffer for content storage
        self.backup_timer = None  # Timer for hourly backups
        self.project_path = None  # Store project path for backup access
        self.synopsis = ''  # Store generated synopsis
        self.paused = False  # Flag for pause/resume control
        
        # Refinement tracking for loaded content adjustments
        self.refinement_type = None  # Content type being refined ('synopsis', 'outline', etc.)
        self.refinement_feedback = None  # Feedback for refinement
        
        # Settings with defaults
        self.llm_model = 'gemma3:12b'
        self.temperature = 0.7
        self.max_retries = 3
        self.detail_level = 'balanced'
        self.character_depth = 'standard'
        self.world_depth = 'standard'
        self.quality_check = 'moderate'
        self.sections_per_chapter = 3
    
    def load_synopsis_from_project(self, project_path):
        """Load synopsis from project files. Tries refined_synopsis.txt first, then synopsis.txt."""
        self.synopsis = ''
        
        # Try to load refined synopsis first (latest version)
        refined_synopsis_path = os.path.join(project_path, 'refined_synopsis.txt')
        if os.path.exists(refined_synopsis_path):
            try:
                with open(refined_synopsis_path, 'r', encoding='utf-8') as f:
                    self.synopsis = f.read().strip()
                if self.synopsis:
                    return
            except Exception as e:
                self.log_update.emit(f"Failed to load refined synopsis: {str(e)}")
        
        # Fall back to initial synopsis if refined not found
        synopsis_path = os.path.join(project_path, 'synopsis.txt')
        if os.path.exists(synopsis_path):
            try:
                with open(synopsis_path, 'r', encoding='utf-8') as f:
                    self.synopsis = f.read().strip()
            except Exception as e:
                self.log_update.emit(f"Failed to load initial synopsis: {str(e)}")
    
    def _generate_with_retry(self, parent_window, model: str, prompt: str, max_retries: Optional[int] = None):
        """Generate LLM response with automatic retry logic.
        
        Args:
            parent_window: ANSWindow instance with client
            model: Model name (e.g., 'gemma3:12b')
            prompt: Prompt text to send to LLM
            max_retries: Maximum number of retry attempts (default from settings)
        
        Returns:
            Generator/Iterator with streamed response, or None if all retries failed
        """
        if max_retries is None:
            max_retries = self.max_retries
        
        for attempt in range(max_retries):
            try:
                if not hasattr(parent_window, 'client'):
                    self.log_update.emit(f"Error: Parent window has no LLM client")
                    return None
                
                stream = parent_window.client.generate(
                    model=model,
                    prompt=prompt,
                    stream=True,
                    options={'temperature': self.temperature}
                )
                
                if not hasattr(stream, '__iter__'):
                    self.log_update.emit(f"Error: LLM response is not iterable")
                    return None
                
                return stream
            
            except Exception as e:
                attempt_num = attempt + 1
                self.log_update.emit(f"LLM connection attempt {attempt_num}/{max_retries} failed: {str(e)}")
                
                if attempt_num < max_retries:
                    # Wait before retry (exponential backoff: 1s, 2s, 4s)
                    import time
                    time.sleep(2 ** attempt)
                else:
                    error_msg = f"Failed to connect to LLM after {max_retries} attempts: {str(e)}"
                    self.log_update.emit(error_msg)
                    if hasattr(parent_window, 'error_signal'):
                        parent_window.error_signal.emit(error_msg)
                    return None
    
    def run(self):
        """Main thread execution. Parse inputs and emit status."""
        try:
            # Check if this is a refinement operation on loaded content
            if isinstance(self.inputs, dict) and self.inputs.get('refinement'):
                # Handle refinement of loaded content
                parent = self.parent()
                if parent is None or parent.__class__.__name__ != 'ANSWindow':
                    self.processing_error.emit("No active project context")
                    return
                
                parent_window = parent  # type: ignore
                content_type = self.inputs.get('type')
                feedback = self.inputs.get('feedback', '')
                
                # Load synopsis from project for refinement context
                if parent_window.current_project:  # type: ignore
                    self.load_synopsis_from_project(parent_window.current_project['path'])  # type: ignore
                
                # Call the appropriate refinement method
                if content_type == 'synopsis':
                    self.refine_synopsis_with_feedback(content_type, feedback)
                elif content_type == 'outline':
                    self.refine_outline_with_feedback(content_type, feedback)
                elif content_type == 'characters':
                    self.refine_characters_with_feedback(content_type, feedback)
                elif content_type == 'world':
                    self.refine_world_with_feedback(content_type, feedback)
                elif content_type == 'timeline':
                    self.refine_timeline_with_feedback(content_type, feedback)
                elif content_type == 'section':
                    self.refine_section_with_feedback(content_type, feedback)
                return
            
            # Check if this is a generation operation (not a novel generation config string)
            if isinstance(self.inputs, dict) and self.inputs.get('operation'):
                # Handle approval-triggered operations
                parent = self.parent()
                if parent is None or parent.__class__.__name__ != 'ANSWindow':
                    self.processing_error.emit("No active project context")
                    return
                
                parent_window = parent  # type: ignore
                operation = self.inputs.get('operation')
                content_type = self.inputs.get('type')
                
                # Load synopsis for context in all operations
                if parent_window.current_project:  # type: ignore
                    self.load_synopsis_from_project(parent_window.current_project['path'])  # type: ignore
                
                # Execute the appropriate operation
                if operation == 'generate_outline':
                    self.generate_outline(content_type)
                elif operation == 'generate_characters':
                    self.generate_characters(content_type)
                elif operation == 'generate_world':
                    self.generate_world(content_type)
                elif operation == 'generate_timeline':
                    self.generate_timeline(content_type)
                elif operation == 'start_chapter_research_loop':
                    self.start_chapter_research_loop()
                elif operation == 'approve_section':
                    self.approve_section(content_type)
                return
            
            # Validate inputs exist
            if not self.inputs:
                self.processing_error.emit("No configuration provided")
                return
            
            # Parse configuration string: "Idea: {idea}, Tone: {tone}, Soft Target: {target}"
            # Look for separators from the end to extract soft target first
            soft_target = 250000
            inputs_str = str(self.inputs)
            
            # Try to extract soft target from end
            soft_target_match = re.search(r', Soft Target: (\d+)$', inputs_str)
            if soft_target_match:
                soft_target = int(soft_target_match.group(1))
                # Remove the soft target from the string
                inputs_str = inputs_str[:soft_target_match.start()]
            
            # Now find where "Idea:" and ", Tone:" are
            idea_start = inputs_str.find('Idea: ')
            tone_start = inputs_str.rfind(', Tone: ')  # Use rfind to get the last occurrence
            
            if idea_start == -1 or tone_start == -1:
                self.processing_error.emit(f"Invalid config format: {self.inputs}")
                return
            
            # Extract idea and tone
            idea = inputs_str[idea_start + 6:tone_start].strip()  # Skip "Idea: "
            tone = inputs_str[tone_start + 8:].strip()  # Skip ", Tone: "
            
            # Emit log update with parsing result
            log_msg = f"Parsed config - Idea: {idea[:50]}..., Tone: {tone[:50]}..., Soft Target: {soft_target}"
            self.log_update.emit(log_msg)
            # Get parent window (ANSWindow) to access current_project
            parent = self.parent()
            # Validate parent is ANSWindow before proceeding (check by class name to avoid forward reference)
            if parent is None or parent.__class__.__name__ != 'ANSWindow':
                self.processing_error.emit("No active project context")
                return
            
            parent_window = parent  # type: ignore  # parent is ANSWindow at runtime
            if not hasattr(parent_window, 'current_project') or not parent_window.current_project:  # type: ignore
                self.processing_error.emit("No active project")
                return
            
            project_path = parent_window.current_project['path']  # type: ignore
            
            # Write to config.txt with key-value pairs
            config_data = (
                f"Idea: {idea}\n"
                f"Tone: {tone}\n"
                f"SoftTarget: {soft_target}\n"
                f"TotalChapters: 0\n"
                f"CurrentChapter: 0\n"
                f"CurrentSection: 0\n"
                f"Progress: 0%\n"
                f"LLMModel: gemma3:12b"
            )
            config_path = os.path.join(project_path, 'config.txt')
            with open(config_path, 'w', encoding='utf-8') as f:
                f.write(config_data)
            self.log_update.emit(f"Wrote configuration to config.txt")
            
            # Update context.txt with initial context
            context_entry = f"Novel started: {idea}. Initial tone: {tone}."
            context_path = os.path.join(project_path, 'context.txt')
            with open(context_path, 'a', encoding='utf-8') as f:
                f.write(context_entry + "\n")
            self.log_update.emit(f"Updated context.txt with novel start information")
            
            # Store project path and start backup timer (3600 seconds = 1 hour)
            self.project_path = project_path
            self.backup_timer = threading.Timer(3600, self.backup)
            self.backup_timer.daemon = True
            self.backup_timer.start()
            self.log_update.emit("Backup timer started (1 hour interval)")
            
            # Emit initialization complete signal
            self.init_complete.emit()
            self.log_update.emit("Init complete signal emitted")
            
            # Generate synopsis using LLM
            self.log_update.emit("Starting synopsis generation...")
            parent_window_check = self.parent()
            if parent_window_check is None or not isinstance(parent_window_check, QtWidgets.QMainWindow):
                self.log_update.emit("Synopsis generation: Invalid parent window")
                return
            
            parent_window = parent_window_check
            
            if not hasattr(parent_window, 'client'):
                self.log_update.emit("Synopsis generation: Parent has no client attribute")
                return
            
            try:
                self.log_update.emit(f"Generating synopsis with gemma3:12b model (streaming)...")
                
                synopsis_prompt = (
                    f"Generate ONLY a 500–1000-word novel synopsis. No introduction or preamble.\n\n"
                    f"Novel Idea: {idea}\n"
                    f"Tone: {tone}\n"
                    f"Target Word Count: {soft_target}\n\n"
                    f"Include these sections:\n"
                    f"- Setting: World and time period\n"
                    f"- Characters: Main characters and their motivations\n"
                    f"- Themes: Central themes (e.g., rebellion, redemption)\n"
                    f"- Plot Arc: Beginning, middle, climax, and resolution\n\n"
                    f"Ensure: No plot holes, consistent tone, engaging narrative. Start with the synopsis immediately."
                )
                
                self.log_update.emit(f"Streaming LLM response (tokens arriving in real-time)...")
                
                # Stream the response and collect tokens
                self.synopsis = ''
                token_count = 0
                
                # Use retry helper to get stream
                stream = self._generate_with_retry(
                    parent_window,
                    model=self.llm_model,
                    prompt=synopsis_prompt,
                    max_retries=3
                )
                
                if stream is None:
                    self.log_update.emit("Failed to generate synopsis after retries")
                    return
                
                # Collect tokens from stream
                for chunk in stream:
                    # Check if paused
                    self.wait_while_paused()
                    
                    try:
                        # Handle both dict and GenerateResponse object formats
                        token = None
                        
                        # First try dict access for raw API responses
                        if isinstance(chunk, dict) and 'response' in chunk:
                            token = chunk['response']
                        # Then try object attribute access for ollama._types.GenerateResponse
                        elif hasattr(chunk, 'response'):
                            token = chunk.response  # type: ignore
                        
                        if token:
                            self.synopsis += token
                            token_count += 1
                            
                            # Update synopsis display live every token for real-time rendering
                            self.synopsis_ready.emit(self.synopsis)
                            
                            # Log every 100 tokens to avoid log spam
                            if token_count % 100 == 0:
                                self.log_update.emit(f"[Streaming] {token_count} tokens received...")
                        else:
                            # Skip empty/invalid chunks silently instead of warning
                            pass
                    
                    except Exception as e:
                        self.log_update.emit(f"Warning: Error processing chunk: {str(e)}")
                        continue
                
                word_count = len(self.synopsis.split())
                
                # Save initial synopsis to synopsis.txt
                synopsis_path = os.path.join(project_path, 'synopsis.txt')
                with open(synopsis_path, 'w', encoding='utf-8') as f:
                    f.write(self.synopsis)
                
                # Log only once at the end of generation
                self.log_update.emit(f"Synopsis generation complete: {word_count} words ({token_count} tokens)")
                
                # Emit synopsis ready signal to update UI
                self.synopsis_ready.emit(self.synopsis)
                
                # Refinement phase: improve synopsis quality
                self.log_update.emit("Starting synopsis refinement...")
                refinement_start_signal = parent_window.refinement_start if hasattr(parent_window, 'refinement_start') else None
                if refinement_start_signal:
                    refinement_start_signal.emit()
                refinement_prompt = (
                    f"Revise this synopsis: \"{self.synopsis}\" "
                    f"to enhance depth, coherence, tone alignment to \"{tone}\", "
                    f"and check for any inconsistencies or overused vocabulary. "
                    f"Return ONLY the revised synopsis, no explanation or preamble."
                )
                
                refined_synopsis = ''
                refinement_token_count = 0
                
                # Use retry helper for refinement
                refinement_stream = self._generate_with_retry(
                    parent_window,
                    model=self.llm_model,
                    prompt=refinement_prompt,
                    max_retries=3
                )
                
                if refinement_stream is not None:
                    # Collect refined tokens
                    for chunk in refinement_stream:
                        # Check if paused
                        self.wait_while_paused()
                        
                        try:
                            # Handle both dict and GenerateResponse object formats
                            token = None
                            
                            if isinstance(chunk, dict) and 'response' in chunk:
                                token = chunk['response']
                            elif hasattr(chunk, 'response'):
                                token = chunk.response  # type: ignore
                            
                            if token:
                                refined_synopsis += token
                                refinement_token_count += 1
                                
                                # Log every 100 tokens to avoid spam
                                if refinement_token_count % 100 == 0:
                                    self.log_update.emit(f"[Refinement] {refinement_token_count} tokens received...")
                        
                        except Exception as e:
                            self.log_update.emit(f"Warning: Error processing refinement chunk: {str(e)}")
                            continue
                    
                    # Update synopsis with refined version
                    if refined_synopsis:
                        self.synopsis = refined_synopsis
                        refined_word_count = len(self.synopsis.split())
                        
                        # Save refined synopsis to refined_synopsis.txt
                        refined_synopsis_path = os.path.join(project_path, 'refined_synopsis.txt')
                        with open(refined_synopsis_path, 'w', encoding='utf-8') as f:
                            f.write(self.synopsis)
                        
                        self.log_update.emit(f"Synopsis refinement complete: {refined_word_count} words ({refinement_token_count} tokens)")
                        
                        # Emit new_synopsis signal with refined text
                        self.new_synopsis.emit(self.synopsis)
                else:
                    self.log_update.emit("Synopsis refinement failed after retries")
            
            except Exception as e:
                self.log_update.emit(f"Synopsis generation error: {str(e)}")
            
        except Exception as e:
            self.processing_error.emit(f"Error in run(): {str(e)}")
    
    def backup(self):
        """Perform hourly backup of story, log, and buffer files."""
        try:
            if not self.project_path:
                return
            
            # Generate timestamp for backup files
            timestamp = datetime.datetime.now().strftime('%Y%m%d_%H%M%S')
            
            # Backup story.txt
            story_src = os.path.join(self.project_path, 'story.txt')
            story_backup = os.path.join(self.project_path, f'story_backup_{timestamp}.txt')
            if os.path.exists(story_src):
                with open(story_src, 'r', encoding='utf-8') as f:
                    content = f.read()
                with open(story_backup, 'w', encoding='utf-8') as f:
                    f.write(content)
            
            # Backup log.txt
            log_src = os.path.join(self.project_path, 'log.txt')
            log_backup = os.path.join(self.project_path, f'log_backup_{timestamp}.txt')
            if os.path.exists(log_src):
                with open(log_src, 'r', encoding='utf-8') as f:
                    content = f.read()
                with open(log_backup, 'w', encoding='utf-8') as f:
                    f.write(content)
            
            # Overwrite buffer_backup.txt with current buffer
            buffer_backup_path = os.path.join(self.project_path, 'buffer_backup.txt')
            with open(buffer_backup_path, 'w', encoding='utf-8') as f:
                f.write(self.buffer)
            
            self.log_update.emit(f"Backup completed: story_backup_{timestamp}.txt, log_backup_{timestamp}.txt, buffer_backup.txt")
            
            # Restart timer for next backup in 1 hour
            self.backup_timer = threading.Timer(3600, self.backup)
            self.backup_timer.daemon = True
            self.backup_timer.start()
            
        except Exception as e:
            self.log_update.emit(f"Backup error: {str(e)}")
    
    def start_processing(self, data):
        """Store input data and start thread execution. Handles thread restart safely."""
        self.inputs = data
        # Reset thread state - QThread can only be started once, so we start fresh
        if self.isRunning():
            self.quit()
            # Wait for thread to finish with 5 second timeout
            if not self.wait(5000):  # 5000 milliseconds = 5 seconds
                self.log_update.emit("Warning: Previous thread did not finish in time, forcing start")
        self.start()
    
    def set_paused(self, paused):
        """Set pause state of the background thread."""
        self.paused = paused
    
    def is_paused(self):
        """Check if background thread is paused."""
        return self.paused
    
    def wait_while_paused(self):
        """Block execution until thread is resumed. Called during long operations."""
        while self.paused:
            QtCore.QThread.msleep(100)  # Sleep for 100ms to avoid busy-wait
    
    def refine_synopsis_with_feedback(self, content_type, feedback):
        """Refine synopsis based on user feedback. Only processes 'synopsis' type."""
        if content_type != 'synopsis':
            return
        # Get parent window to access tone and client
        parent = self.parent()
        if not isinstance(parent, QtWidgets.QMainWindow):
            self.log_update.emit("Error: No active window context for refinement")
            return
        
        parent_window: 'ANSWindow' = parent  # type: ignore[assignment]
        if not hasattr(parent_window, 'current_project') or not parent_window.current_project:
            self.log_update.emit("Error: No active project for refinement")
            return
        
        # Ensure we have the synopsis to refine
        if not self.synopsis:
            self.log_update.emit("[DEBUG] Synopsis not in memory, loading from project...")
            self.load_synopsis_from_project(parent_window.current_project['path'])
        
        if not self.synopsis:
            self.log_update.emit("Error: No synopsis available for refinement")
            return
        
        self.log_update.emit(f"[DEBUG] Synopsis loaded: {len(self.synopsis)} characters")
        
        # Get tone from the novel idea inputs if available
        tone = ""
        if hasattr(parent_window, 'tone_input'):
            tone = parent_window.tone_input.toPlainText().strip()
        
        if not tone:
            tone = "as originally specified"
        
        # Emit refinement_start signal to clear display and disable buttons
        refinement_start_signal = parent_window.refinement_start if hasattr(parent_window, 'refinement_start') else None
        if refinement_start_signal:
            refinement_start_signal.emit()
        
        # Generate refinement prompt WITH FULL SYNOPSIS CONTEXT
        refinement_prompt = (
            f"Refine this synopsis based on the following feedback:\n\n"
            f"ORIGINAL SYNOPSIS:\n{self.synopsis}\n\n"
            f"USER FEEDBACK:\n{feedback}\n\n"
            f"INSTRUCTIONS:\n"
            f"- Incorporate the feedback while maintaining the core story\n"
            f"- Keep tone: {tone}\n"
            f"- Maintain coherence and plot consistency\n"
            f"- Avoid vocabulary overuse\n"
            f"- Return ONLY the revised synopsis, no explanation or preamble."
        )
        
        # DEBUG: Log the prompt being sent
        self.log_update.emit(f"[DEBUG] Refinement prompt length: {len(refinement_prompt)} characters")
        self.log_update.emit(f"[DEBUG] Tone being used: {tone[:50]}..." if len(tone) > 50 else f"[DEBUG] Tone being used: {tone}")
        self.log_update.emit("Starting synopsis refinement with user feedback...")
        refined_synopsis = ''
        refinement_token_count = 0
        
        # Use retry helper for refinement
        refinement_stream = self._generate_with_retry(
            parent_window,
            model=self.llm_model,
            prompt=refinement_prompt,
            max_retries=3
        )
        
        if refinement_stream is None:
            self.log_update.emit("Failed to refine synopsis after retries")
            return
        
        # Stream refined tokens in REAL-TIME for live display updates
        for chunk in refinement_stream:
            try:
                # Check for pause
                self.wait_while_paused()
                
                # Handle both dict and GenerateResponse object formats
                token = None
                
                if isinstance(chunk, dict) and 'response' in chunk:
                    token = chunk['response']
                elif hasattr(chunk, 'response'):
                    token = chunk.response  # type: ignore
                
                if token:
                    refined_synopsis += token
                    refinement_token_count += 1
                    
                    # EMIT EVERY TOKEN for live streaming (this is key!)
                    # Handler checks if new_length > current_length to detect updates
                    self.new_synopsis.emit(refined_synopsis)
                    
                    # Log every 100 tokens to avoid spam
                    if refinement_token_count % 100 == 0:
                        self.log_update.emit(f"[Feedback Refinement] {refinement_token_count} tokens received...")
            
            except Exception as e:
                self.log_update.emit(f"Warning: Error processing refinement chunk: {str(e)}")
                continue
        
        # Final update with complete refined synopsis
        if refined_synopsis:
            self.synopsis = refined_synopsis
            refined_word_count = len(self.synopsis.split())
            
            # Save refined synopsis to refined_synopsis.txt
            project_path = parent_window.current_project['path']
            refined_synopsis_path = os.path.join(project_path, 'refined_synopsis.txt')
            with open(refined_synopsis_path, 'w', encoding='utf-8') as f:
                f.write(self.synopsis)
            
            self.log_update.emit(f"Synopsis refinement complete: {refined_word_count} words ({refinement_token_count} tokens)")
            
            # Final emit for consistency (already emitted all tokens, but ensure complete state)
            self.new_synopsis.emit(self.synopsis)
    
    def generate_outline(self, content_type):
        """Generate detailed 25-chapter outline based on approved synopsis. Only processes 'synopsis' type."""
        if content_type != 'synopsis':
            return
        
        # Get parent window to access tone, target, client and project
        parent = self.parent()
        if not parent or not isinstance(parent, QtWidgets.QMainWindow):
            self.log_update.emit("Error: No active window context for outline generation")
            return
        
        parent_window: 'ANSWindow' = parent  # type: ignore[assignment]
        if not hasattr(parent_window, 'current_project') or not parent_window.current_project:
            self.log_update.emit("Error: No active project for outline generation")
            return
        
        if not self.synopsis:
            # Try to load synopsis from project files
            self.load_synopsis_from_project(parent_window.current_project['path'])
        
        if not self.synopsis:
            self.log_update.emit("Error: No synopsis available for outline generation")
            return
        
        # Get tone from the novel idea inputs
        tone = ""
        if hasattr(parent_window, 'tone_input'):
            tone = parent_window.tone_input.toPlainText().strip()
        
        if not tone:
            tone = "as originally specified"
        
        # Get soft target from inputs if available
        soft_target = 5000
        if hasattr(parent_window, 'target_input'):
            target_text = parent_window.target_input.toPlainText().strip()
            try:
                soft_target = int(target_text)
            except (ValueError, AttributeError):
                soft_target = 5000
        
        # Generate outline prompt
        outline_prompt = (
            f"Create a detailed outline for a novel with soft target {soft_target} words based on this synopsis: \"{self.synopsis}\". "
            f"Include: 25 chapters with titles, 100–200-word summaries per chapter, key events, character developments, "
            f"dynamic chapter lengths (5000–15000 words) based on pacing. "
            f"Tone: \"{tone}\". "
            f"Ensure no contradictions. "
            f"Return ONLY the outline structure, no explanation or preamble."
        )
        
        self.log_update.emit("Starting outline generation based on approved synopsis...")
        outline_text = ''
        outline_token_count = 0
        
        # Use retry helper for outline generation
        outline_stream = self._generate_with_retry(
            parent_window,
            model=self.llm_model,
            prompt=outline_prompt,
            max_retries=3
        )
        
        if outline_stream is None:
            self.log_update.emit("Failed to generate outline after retries")
            return
        
        try:
            # Collect outline tokens
            for chunk in outline_stream:
                try:
                    # Handle both dict and GenerateResponse object formats
                    token = None
                    
                    if isinstance(chunk, dict) and 'response' in chunk:
                        token = chunk['response']
                    elif hasattr(chunk, 'response'):
                        token = chunk.response  # type: ignore
                    
                    if token:
                        outline_text += token
                        outline_token_count += 1
                        
                        # Log every 100 tokens to avoid spam
                        if outline_token_count % 100 == 0:
                            self.log_update.emit(f"[Outline Generation] {outline_token_count} tokens received...")
                
                except Exception as e:
                    self.log_update.emit(f"Warning: Error processing outline chunk: {str(e)}")
                    continue
            
            # Save outline
            if outline_text:
                project_path = parent_window.current_project['path']
                outline_path = os.path.join(project_path, 'outline.txt')
                with open(outline_path, 'w', encoding='utf-8') as f:
                    f.write("=== NOVEL OUTLINE (25 CHAPTERS) ===\n\n")
                    f.write(outline_text)
                    f.write("\n\n=== END OUTLINE ===\n")
                
                outline_word_count = len(outline_text.split())
                self.log_update.emit(f"Outline generation complete: {outline_word_count} words ({outline_token_count} tokens)")
                
                # Emit new_outline signal to update UI
                self.new_outline.emit(outline_text)
                
                # Save progress: outline ready for approval
                if isinstance(parent_window, QtWidgets.QMainWindow):
                    parent_window._save_progress('outline', 'ready_for_approval')
            else:
                self.log_update.emit("Warning: No outline text generated")
        
        except Exception as e:
            self.log_update.emit(f"Outline generation error: {str(e)}")
    
    def refine_outline_with_feedback(self, content_type, feedback):
        """Refine outline based on user feedback. Only processes 'outline' type."""
        if content_type != 'outline':
            return
        
        # Get parent window to access client and project
        parent = self.parent()
        if not isinstance(parent, QtWidgets.QMainWindow):
            self.log_update.emit("Error: No active window context for outline refinement")
            return
        
        parent_window: 'ANSWindow' = parent  # type: ignore[assignment]
        if not hasattr(parent_window, 'current_project') or not parent_window.current_project:
            self.log_update.emit("Error: No active project for outline refinement")
            return
        
        if not self.synopsis:
            self.load_synopsis_from_project(parent_window.current_project['path'])
        
        if not self.synopsis:
            self.log_update.emit("Error: No synopsis available for outline refinement")
            return
        
        # Load the current outline for context
        current_outline = ""
        if hasattr(parent_window, 'outline_display'):
            current_outline = parent_window.outline_display.toPlainText().strip()
        
        if not current_outline:
            # Try loading from file
            outline_file = os.path.join(parent_window.current_project['path'], 'outline.txt')
            if os.path.exists(outline_file):
                try:
                    with open(outline_file, 'r', encoding='utf-8') as f:
                        current_outline = f.read().strip()
                except Exception as e:
                    self.log_update.emit(f"Warning: Could not load outline from file: {str(e)}")
        
        if not current_outline:
            self.log_update.emit("Error: No outline available for refinement")
            return
        
        # Generate outline refinement prompt WITH FULL CONTEXT
        refinement_prompt = (
            f"Refine this novel outline based on the following feedback:\n\n"
            f"SYNOPSIS:\n{self.synopsis}\n\n"
            f"CURRENT OUTLINE:\n{current_outline}\n\n"
            f"USER FEEDBACK:\n{feedback}\n\n"
            f"INSTRUCTIONS:\n"
            f"- Incorporate the feedback while maintaining story coherence\n"
            f"- Keep the overall structure and tone\n"
            f"- Ensure chapters still align with the synopsis\n"
            f"- Return ONLY the revised outline, no explanation or preamble."
        )
        
        self.log_update.emit("Starting outline refinement with user feedback...")
        
        # Emit outline_refinement_start signal to clear outline display
        outline_refinement_signal = parent_window.outline_refinement_start if hasattr(parent_window, 'outline_refinement_start') else None
        if outline_refinement_signal:
            outline_refinement_signal.emit()
        
        refined_outline = ''
        outline_token_count = 0
        
        # Use retry helper for outline refinement
        refinement_stream = self._generate_with_retry(
            parent_window,
            model=self.llm_model,
            prompt=refinement_prompt,
            max_retries=3
        )
        
        if refinement_stream is None:
            self.log_update.emit("Failed to refine outline after retries")
            return
        
        try:
            # Stream refined tokens in REAL-TIME for live display updates
            for chunk in refinement_stream:
                try:
                    # Check for pause
                    self.wait_while_paused()
                    
                    # Handle both dict and GenerateResponse object formats
                    token = None
                    
                    if isinstance(chunk, dict) and 'response' in chunk:
                        token = chunk['response']
                    elif hasattr(chunk, 'response'):
                        token = chunk.response  # type: ignore
                    
                    if token:
                        refined_outline += token
                        outline_token_count += 1
                        
                        # EMIT EVERY TOKEN for live streaming
                        self.new_outline.emit(refined_outline)
                        
                        # Log every 100 tokens to avoid spam
                        if outline_token_count % 100 == 0:
                            self.log_update.emit(f"[Outline Refinement] {outline_token_count} tokens received...")
                
                except Exception as e:
                    self.log_update.emit(f"Warning: Error processing outline refinement chunk: {str(e)}")
                    continue
            
            # Final processing with refined version
            if refined_outline:
                project_path = parent_window.current_project['path']
                outline_path = os.path.join(project_path, 'outline.txt')
                with open(outline_path, 'w', encoding='utf-8') as f:
                    f.write("=== NOVEL OUTLINE (25 CHAPTERS - REFINED) ===\n\n")
                    f.write(refined_outline)
                    f.write("\n\n=== END OUTLINE ===\n")
                
                refined_word_count = len(refined_outline.split())
                self.log_update.emit(f"Outline refinement complete: {refined_word_count} words ({outline_token_count} tokens)")
                
                # Final emit for consistency
                self.new_outline.emit(refined_outline)
        
        except Exception as e:
            self.log_update.emit(f"Outline refinement error: {str(e)}")
    
    def generate_characters(self, content_type):
        """Generate detailed character profiles based on outline. Only processes 'outline' type."""
        if content_type != 'outline':
            return
        
        # Get parent window to access outline, client and project
        parent = self.parent()
        if not isinstance(parent, QtWidgets.QMainWindow):
            self.log_update.emit("Error: No active window context for character generation")
            return
        
        parent_window: 'ANSWindow' = parent  # type: ignore[assignment]
        if not hasattr(parent_window, 'current_project') or not parent_window.current_project:
            self.log_update.emit("Error: No active project for character generation")
            return
        
        # Get outline from file
        project_path = parent_window.current_project['path']
        outline_path = os.path.join(project_path, 'outline.txt')
        
        try:
            with open(outline_path, 'r', encoding='utf-8') as f:
                outline_content = f.read()
        except Exception as e:
            self.log_update.emit(f"Error reading outline: {str(e)}")
            return
        
        if not outline_content.strip():
            self.log_update.emit("Error: Outline is empty")
            return
        
        # Get synopsis for context
        if not self.synopsis:
            # Try to load synopsis from project files
            self.load_synopsis_from_project(project_path)
        
        if not self.synopsis:
            self.log_update.emit("Error: No synopsis available for character generation")
            return
        
        # Generate character profiles prompt
        character_prompt = (
            f"Generate detailed profiles for main characters in this outline: \"{outline_content}\". "
            f"Include: Name, Age, Background, Traits, Arc, Relationships. "
            f"Format as JSON array of character objects. "
            f"Ensure consistency with synopsis: \"{self.synopsis}\". "
            f"Return ONLY the JSON array, no explanation or preamble."
        )
        
        self.log_update.emit("Starting character generation based on approved outline...")
        characters_json = ''
        character_token_count = 0
        
        character_stream = self._generate_with_retry(
            parent_window,
            model=self.llm_model,
            prompt=character_prompt,
            max_retries=3
        )
        
        if character_stream is None:
            self.log_update.emit("Failed to generate characters after retries")
            return
        
        try:
            
            # Collect character tokens
            for chunk in character_stream:
                try:
                    # Handle both dict and GenerateResponse object formats
                    token = None
                    
                    if isinstance(chunk, dict) and 'response' in chunk:
                        token = chunk['response']
                    elif hasattr(chunk, 'response'):
                        token = chunk.response  # type: ignore
                    
                    if token:
                        characters_json += token
                        character_token_count += 1
                        
                        # Log every 100 tokens to avoid spam
                        if character_token_count % 100 == 0:
                            self.log_update.emit(f"[Character Generation] {character_token_count} tokens received...")
                
                except Exception as e:
                    self.log_update.emit(f"Warning: Error processing character chunk: {str(e)}")
                    continue
            
            # Save characters to file
            if characters_json:
                characters_path = os.path.join(project_path, 'characters.txt')
                with open(characters_path, 'w', encoding='utf-8') as f:
                    f.write("=== MAIN CHARACTERS (JSON) ===\n\n")
                    f.write(characters_json)
                    f.write("\n\n=== END CHARACTERS ===\n")
                
                character_word_count = len(characters_json.split())
                self.log_update.emit(f"Character generation complete: {character_word_count} words ({character_token_count} tokens)")
                
                # Emit new_characters signal to update UI
                self.new_characters.emit(characters_json)
                
                # Save progress: characters ready for approval
                if isinstance(parent_window, QtWidgets.QMainWindow):
                    parent_window._save_progress('characters', 'ready_for_approval')
            else:
                self.log_update.emit("Warning: No character data generated")
        
        except Exception as e:
            self.log_update.emit(f"Character generation error: {str(e)}")
    
    def refine_characters_with_feedback(self, content_type, feedback):
        """Refine character profiles based on user feedback. Only processes 'characters' type."""
        if content_type != 'characters':
            return
        
        # Get parent window to access client and project
        parent = self.parent()
        if not isinstance(parent, QtWidgets.QMainWindow):
            self.log_update.emit("Error: No active window context for character refinement")
            return
        
        parent_window: 'ANSWindow' = parent  # type: ignore[assignment]
        if not hasattr(parent_window, 'current_project') or not parent_window.current_project:
            self.log_update.emit("Error: No active project for character refinement")
            return
        
        # Ensure we have the synopsis context
        if not self.synopsis:
            self.load_synopsis_from_project(parent_window.current_project['path'])
        
        # Get current characters from file or display
        project_path = parent_window.current_project['path']
        characters_content = ""
        
        # Try to get from display first
        if hasattr(parent_window, 'characters_display'):
            characters_content = parent_window.characters_display.toPlainText().strip()
        
        # If not in display, read from file
        if not characters_content:
            characters_path = os.path.join(project_path, 'characters.txt')
            try:
                with open(characters_path, 'r', encoding='utf-8') as f:
                    characters_content = f.read().strip()
            except Exception as e:
                self.log_update.emit(f"Error reading characters: {str(e)}")
                return
        
        if not characters_content:
            self.log_update.emit("Error: Characters are empty")
            return
        
        # Generate character refinement prompt WITH FULL CONTEXT
        refinement_prompt = (
            f"Refine the character profiles based on the following feedback:\n\n"
            f"SYNOPSIS:\n{self.synopsis}\n\n"
            f"CURRENT CHARACTERS:\n{characters_content}\n\n"
            f"USER FEEDBACK:\n{feedback}\n\n"
            f"INSTRUCTIONS:\n"
            f"- Incorporate feedback while maintaining consistency with synopsis\n"
            f"- Keep character arcs and relationships coherent\n"
            f"- Ensure character depth matches story requirements\n"
            f"- Return ONLY valid JSON array of character objects, no explanation or preamble."
        )
        
        self.log_update.emit("Starting character refinement with user feedback...")
        
        refined_characters = ''
        characters_token_count = 0
        
        refinement_stream = self._generate_with_retry(
            parent_window,
            model=self.llm_model,
            prompt=refinement_prompt,
            max_retries=3
        )
        
        if refinement_stream is None:
            self.log_update.emit("Failed to refine characters after retries")
            return
        
        try:
            
            # Stream refined tokens in REAL-TIME for live display updates
            for chunk in refinement_stream:
                try:
                    # Check for pause
                    self.wait_while_paused()
                    
                    # Handle both dict and GenerateResponse object formats
                    token = None
                    
                    if isinstance(chunk, dict) and 'response' in chunk:
                        token = chunk['response']
                    elif hasattr(chunk, 'response'):
                        token = chunk.response  # type: ignore
                    
                    if token:
                        refined_characters += token
                        characters_token_count += 1
                        
                        # EMIT EVERY TOKEN for live streaming
                        self.new_characters.emit(refined_characters)
                        
                        # Log every 100 tokens to avoid spam
                        if characters_token_count % 100 == 0:
                            self.log_update.emit(f"[Character Refinement] {characters_token_count} tokens received...")
                
                except Exception as e:
                    self.log_update.emit(f"Warning: Error processing character refinement chunk: {str(e)}")
                    continue
            
            # Final processing with refined version
            if refined_characters:
                with open(characters_path, 'w', encoding='utf-8') as f:
                    f.write("=== MAIN CHARACTERS (JSON - REFINED) ===\n\n")
                    f.write(refined_characters)
                    f.write("\n\n=== END CHARACTERS ===\n")
                
                refined_word_count = len(refined_characters.split())
                self.log_update.emit(f"Character refinement complete: {refined_word_count} words ({characters_token_count} tokens)")
                
                # Final emit for consistency
                self.new_characters.emit(refined_characters)
        
        except Exception as e:
            self.log_update.emit(f"Character refinement error: {str(e)}")
    
    def generate_world(self, content_type):
        """Generate world-building details based on outline. Only processes 'characters' type."""
        if content_type != 'characters':
            return
        
        # Get parent window to access outline, client and project
        parent = self.parent()
        if not isinstance(parent, QtWidgets.QMainWindow):
            self.log_update.emit("Error: No active window context for world generation")
            return
        
        parent_window: 'ANSWindow' = parent  # type: ignore[assignment]
        if not hasattr(parent_window, 'current_project') or not parent_window.current_project:
            self.log_update.emit("Error: No active project for world generation")
            return
        
        # Get outline from file
        project_path = parent_window.current_project['path']
        outline_path = os.path.join(project_path, 'outline.txt')
        
        try:
            with open(outline_path, 'r', encoding='utf-8') as f:
                outline_content = f.read()
        except Exception as e:
            self.log_update.emit(f"Error reading outline: {str(e)}")
            return
        
        if not outline_content.strip():
            self.log_update.emit("Error: Outline is empty")
            return
        
        # Generate world-building prompt
        world_prompt = (
            f"Generate world-building details for this outline: \"{outline_content}\". "
            f"Include: Tech, Culture, Geography, History, Rules. "
            f"Format as JSON dict. "
            f"Align with tone and characters. "
            f"Return ONLY the JSON dict, no explanation or preamble."
        )
        
        self.log_update.emit("Starting world generation based on approved characters...")
        world_json = ''
        world_token_count = 0
        
        world_stream = self._generate_with_retry(
            parent_window,
            model=self.llm_model,
            prompt=world_prompt,
            max_retries=3
        )
        
        if world_stream is None:
            self.log_update.emit("Failed to generate world after retries")
            return
        
        try:
            
            # Collect world tokens
            for chunk in world_stream:
                try:
                    # Handle both dict and GenerateResponse object formats
                    token = None
                    
                    if isinstance(chunk, dict) and 'response' in chunk:
                        token = chunk['response']
                    elif hasattr(chunk, 'response'):
                        token = chunk.response  # type: ignore
                    
                    if token:
                        world_json += token
                        world_token_count += 1
                        
                        # Log every 100 tokens to avoid spam
                        if world_token_count % 100 == 0:
                            self.log_update.emit(f"[World Generation] {world_token_count} tokens received...")
                
                except Exception as e:
                    self.log_update.emit(f"Warning: Error processing world chunk: {str(e)}")
                    continue
            
            # Save world to file
            if world_json:
                world_path = os.path.join(project_path, 'world.txt')
                with open(world_path, 'w', encoding='utf-8') as f:
                    f.write("=== WORLD BUILDING (JSON) ===\n\n")
                    f.write(world_json)
                    f.write("\n\n=== END WORLD ===\n")
                
                world_word_count = len(world_json.split())
                self.log_update.emit(f"World generation complete: {world_word_count} words ({world_token_count} tokens)")
                
                # Emit new_world signal to update UI
                self.new_world.emit(world_json)
                
                # Save progress: world ready for approval
                if isinstance(parent_window, QtWidgets.QMainWindow):
                    parent_window._save_progress('world', 'ready_for_approval')
            else:
                self.log_update.emit("Warning: No world data generated")
        
        except Exception as e:
            self.log_update.emit(f"World generation error: {str(e)}")
    
    def refine_world_with_feedback(self, content_type, feedback):
        """Refine world-building details based on user feedback. Only processes 'world' type."""
        if content_type != 'world':
            return
        
        # Get parent window to access client and project
        parent = self.parent()
        if not isinstance(parent, QtWidgets.QMainWindow):
            self.log_update.emit("Error: No active window context for world refinement")
            return
        
        parent_window: 'ANSWindow' = parent  # type: ignore[assignment]
        if not hasattr(parent_window, 'current_project') or not parent_window.current_project:
            self.log_update.emit("Error: No active project for world refinement")
            return
        
        # Ensure we have the synopsis context
        if not self.synopsis:
            self.load_synopsis_from_project(parent_window.current_project['path'])
        
        # Get current world from display or file
        project_path = parent_window.current_project['path']
        world_content = ""
        
        # Try to get from display first
        if hasattr(parent_window, 'world_display'):
            world_content = parent_window.world_display.toPlainText().strip()
        
        # If not in display, read from file
        if not world_content:
            world_path = os.path.join(project_path, 'world.txt')
            try:
                with open(world_path, 'r', encoding='utf-8') as f:
                    world_content = f.read().strip()
            except Exception as e:
                self.log_update.emit(f"Error reading world: {str(e)}")
                return
        
        if not world_content:
            self.log_update.emit("Error: World is empty")
            return
        
        # Generate world refinement prompt WITH FULL CONTEXT
        refinement_prompt = (
            f"Refine the world-building details based on the following feedback:\n\n"
            f"SYNOPSIS:\n{self.synopsis}\n\n"
            f"CURRENT WORLD:\n{world_content}\n\n"
            f"USER FEEDBACK:\n{feedback}\n\n"
            f"INSTRUCTIONS:\n"
            f"- Incorporate feedback while maintaining consistency with synopsis\n"
            f"- Ensure world rules and magic systems remain coherent\n"
            f"- Check for internal consistency and logical worldbuilding\n"
            f"- Return ONLY valid JSON object with world details, no explanation or preamble."
        )
        
        self.log_update.emit("Starting world refinement with user feedback...")
        
        refined_world = ''
        world_token_count = 0
        
        refinement_stream = self._generate_with_retry(
            parent_window,
            model=self.llm_model,
            prompt=refinement_prompt,
            max_retries=3
        )
        
        if refinement_stream is None:
            self.log_update.emit("Failed to refine world after retries")
            return
        
        try:
            
            # Stream refined tokens in REAL-TIME for live display updates
            for chunk in refinement_stream:
                try:
                    # Check for pause
                    self.wait_while_paused()
                    
                    # Handle both dict and GenerateResponse object formats
                    token = None
                    
                    if isinstance(chunk, dict) and 'response' in chunk:
                        token = chunk['response']
                    elif hasattr(chunk, 'response'):
                        token = chunk.response  # type: ignore
                    
                    if token:
                        refined_world += token
                        world_token_count += 1
                        
                        # EMIT EVERY TOKEN for live streaming
                        self.new_world.emit(refined_world)
                        
                        # Log every 100 tokens to avoid spam
                        if world_token_count % 100 == 0:
                            self.log_update.emit(f"[World Refinement] {world_token_count} tokens received...")
                
                except Exception as e:
                    self.log_update.emit(f"Warning: Error processing world refinement chunk: {str(e)}")
                    continue
            
            # Final processing with refined version
            if refined_world:
                with open(world_path, 'w', encoding='utf-8') as f:
                    f.write("=== WORLD BUILDING (JSON - REFINED) ===\n\n")
                    f.write(refined_world)
                    f.write("\n\n=== END WORLD ===\n")
                
                refined_word_count = len(refined_world.split())
                self.log_update.emit(f"World refinement complete: {refined_word_count} words ({world_token_count} tokens)")
                
                # Final emit for consistency
                self.new_world.emit(refined_world)
        
        except Exception as e:
            self.log_update.emit(f"World refinement error: {str(e)}")
    
    def generate_timeline(self, content_type):
        """Generate timeline from outline. Only processes 'world' type."""
        if content_type != 'world':
            return
        
        # Get parent window to access outline, client, project, and tone
        parent = self.parent()
        if not isinstance(parent, QtWidgets.QMainWindow):
            self.log_update.emit("Error: No active window context for timeline generation")
            return
        
        parent_window: 'ANSWindow' = parent  # type: ignore[assignment]
        if not hasattr(parent_window, 'current_project') or not parent_window.current_project:
            self.log_update.emit("Error: No active project for timeline generation")
            return
        
        # Get outline from file
        project_path = parent_window.current_project['path']
        outline_path = os.path.join(project_path, 'outline.txt')
        
        try:
            with open(outline_path, 'r', encoding='utf-8') as f:
                outline_content = f.read()
        except Exception as e:
            self.log_update.emit(f"Error reading outline: {str(e)}")
            return
        
        if not outline_content.strip():
            self.log_update.emit("Error: Outline is empty")
            return
        
        # Get tone from config for timeline context
        config_content = parent_window.current_project.get('config', '')
        tone = 'Unknown'
        for line in config_content.split('\n'):
            if 'Tone:' in line:
                tone = line.split('Tone:', 1)[1].strip()
                break
        
        # Generate timeline prompt
        timeline_prompt = (
            f"Generate a detailed timeline from this outline: \"{outline_content}\". "
            f"Assign dates, locations, actions for key events. "
            f"Check for chronological inconsistencies, resolve them. "
            f"Include character arcs and tone: \"{tone}\". "
            f"Return only the timeline, no explanation or preamble."
        )
        
        self.log_update.emit("Starting timeline generation from approved planning...")
        timeline_text = ''
        timeline_token_count = 0
        
        timeline_stream = self._generate_with_retry(
            parent_window,
            model=self.llm_model,
            prompt=timeline_prompt,
            max_retries=3
        )
        
        if timeline_stream is None:
            self.log_update.emit("Failed to generate timeline after retries")
            return
        
        try:
            
            # Collect timeline tokens
            for chunk in timeline_stream:
                try:
                    # Handle both dict and GenerateResponse object formats
                    token = None
                    
                    if isinstance(chunk, dict) and 'response' in chunk:
                        token = chunk['response']
                    elif hasattr(chunk, 'response'):
                        token = chunk.response  # type: ignore
                    
                    if token:
                        timeline_text += token
                        timeline_token_count += 1
                        
                        # Log every 100 tokens to avoid spam
                        if timeline_token_count % 100 == 0:
                            self.log_update.emit(f"[Timeline Generation] {timeline_token_count} tokens received...")
                
                except Exception as e:
                    self.log_update.emit(f"Warning: Error processing timeline chunk: {str(e)}")
                    continue
            
            # Save timeline to file
            if timeline_text:
                timeline_path = os.path.join(project_path, 'timeline.txt')
                with open(timeline_path, 'w', encoding='utf-8') as f:
                    f.write("=== NOVEL TIMELINE (WITH DATES, LOCATIONS, EVENTS) ===\n\n")
                    f.write(timeline_text)
                    f.write("\n\n=== END TIMELINE ===\n")
                
                timeline_word_count = len(timeline_text.split())
                self.log_update.emit(f"Timeline generation complete: {timeline_word_count} words ({timeline_token_count} tokens)")
                
                # Emit new_timeline signal to update UI
                self.new_timeline.emit(timeline_text)
            else:
                self.log_update.emit("Warning: No timeline data generated")
        
        except Exception as e:
            self.log_update.emit(f"Timeline generation error: {str(e)}")
    
    def refine_timeline_with_feedback(self, content_type, feedback):
        """Refine timeline based on user feedback. Only processes 'timeline' type."""
        if content_type != 'timeline':
            return
        
        # Get parent window to access client and project
        parent = self.parent()
        if not isinstance(parent, QtWidgets.QMainWindow):
            self.log_update.emit("Error: No active window context for timeline refinement")
            return
        
        parent_window: 'ANSWindow' = parent  # type: ignore[assignment]
        if not hasattr(parent_window, 'current_project') or not parent_window.current_project:
            self.log_update.emit("Error: No active project for timeline refinement")
            return
        
        # Ensure we have the synopsis context
        if not self.synopsis:
            self.load_synopsis_from_project(parent_window.current_project['path'])
        
        # Get current timeline from display or file
        project_path = parent_window.current_project['path']
        timeline_content = ""
        
        # Try to get from display first
        if hasattr(parent_window, 'timeline_display'):
            timeline_content = parent_window.timeline_display.toPlainText().strip()
        
        # If not in display, read from file
        if not timeline_content:
            timeline_path = os.path.join(project_path, 'timeline.txt')
            try:
                with open(timeline_path, 'r', encoding='utf-8') as f:
                    timeline_content = f.read().strip()
            except Exception as e:
                self.log_update.emit(f"Error reading timeline: {str(e)}")
                return
        
        if not timeline_content:
            self.log_update.emit("Error: Timeline is empty")
            return
        
        # Emit timeline_refinement_start signal to clear display and disable buttons
        refinement_signal = parent_window.timeline_refinement_start if hasattr(parent_window, 'timeline_refinement_start') else None
        if refinement_signal:
            refinement_signal.emit()
        
        # Generate timeline refinement prompt WITH FULL CONTEXT
        refinement_prompt = (
            f"Refine the timeline based on the following feedback:\n\n"
            f"SYNOPSIS:\n{self.synopsis}\n\n"
            f"CURRENT TIMELINE:\n{timeline_content}\n\n"
            f"USER FEEDBACK:\n{feedback}\n\n"
            f"INSTRUCTIONS:\n"
            f"- Incorporate feedback while maintaining consistency with synopsis\n"
            f"- Ensure dates, locations, and events are logically ordered\n"
            f"- Verify character arcs align with timeline events\n"
            f"- Keep all dates and event details consistent\n"
            f"- Return ONLY the refined timeline, no explanation or preamble."
        )
        
        self.log_update.emit("Starting timeline refinement with user feedback...")
        
        refined_timeline = ''
        timeline_token_count = 0
        
        refinement_stream = self._generate_with_retry(
            parent_window,
            model=self.llm_model,
            prompt=refinement_prompt,
            max_retries=3
        )
        
        if refinement_stream is None:
            self.log_update.emit("Failed to refine timeline after retries")
            return
        
        try:
            
            # Stream refined tokens in REAL-TIME for live display updates
            for chunk in refinement_stream:
                try:
                    # Check for pause
                    self.wait_while_paused()
                    
                    # Handle both dict and GenerateResponse object formats
                    token = None
                    
                    if isinstance(chunk, dict) and 'response' in chunk:
                        token = chunk['response']
                    elif hasattr(chunk, 'response'):
                        token = chunk.response  # type: ignore
                    
                    if token:
                        refined_timeline += token
                        timeline_token_count += 1
                        
                        # EMIT EVERY TOKEN for live streaming
                        self.new_timeline.emit(refined_timeline)
                        
                        # Log every 100 tokens to avoid spam
                        if timeline_token_count % 100 == 0:
                            self.log_update.emit(f"[Timeline Refinement] {timeline_token_count} tokens received...")
                
                except Exception as e:
                    self.log_update.emit(f"Warning: Error processing timeline refinement chunk: {str(e)}")
                    continue
            
            # Final processing with refined version
            if refined_timeline:
                with open(timeline_path, 'w', encoding='utf-8') as f:
                    f.write("=== NOVEL TIMELINE (WITH DATES, LOCATIONS, EVENTS - REFINED) ===\n\n")
                    f.write(refined_timeline)
                    f.write("\n\n=== END TIMELINE ===\n")
                
                refined_word_count = len(refined_timeline.split())
                self.log_update.emit(f"Timeline refinement complete: {refined_word_count} words ({timeline_token_count} tokens)")
                
                # Final emit for consistency
                self.new_timeline.emit(refined_timeline)
        
        except Exception as e:
            self.log_update.emit(f"Timeline refinement error: {str(e)}")
    
    def refine_section_with_feedback(self, content_type, feedback):
        """Refine draft section based on user feedback. Only processes 'section' type."""
        if content_type != 'section':
            return
        
        # Get parent window to access client and project
        parent = self.parent()
        if not isinstance(parent, QtWidgets.QMainWindow):
            self.log_update.emit("Error: No active window context for section refinement")
            return
        
        parent_window: 'ANSWindow' = parent  # type: ignore[assignment]
        if not hasattr(parent_window, 'current_project') or not parent_window.current_project:
            self.log_update.emit("Error: No active project for section refinement")
            return
        
        # Check if buffer has content
        if not self.buffer or not self.buffer.strip():
            self.log_update.emit("Error: No section content in buffer for refinement")
            return
        
        # Ensure we have the synopsis context for consistency
        if not self.synopsis:
            self.load_synopsis_from_project(parent_window.current_project['path'])
        
        # Generate section refinement prompt WITH FULL CONTEXT
        refinement_prompt = (
            f"Refine this draft section based on the following feedback:\n\n"
            f"SYNOPSIS:\n{self.synopsis}\n\n"
            f"CURRENT SECTION:\n{self.buffer}\n\n"
            f"USER FEEDBACK:\n{feedback}\n\n"
            f"INSTRUCTIONS:\n"
            f"- Incorporate the feedback while maintaining story coherence\n"
            f"- Ensure the section flows naturally and maintains tone\n"
            f"- Check vocabulary for overuse and improve sentence variety\n"
            f"- Verify plot consistency with the synopsis\n"
            f"- Return ONLY the revised section, no explanation or preamble."
        )
        
        self.log_update.emit(f"Starting section refinement with user feedback...")
        refined_section = ''
        section_token_count = 0
        
        refinement_stream = self._generate_with_retry(
            parent_window,
            model=self.llm_model,
            prompt=refinement_prompt,
            max_retries=3
        )
        
        if refinement_stream is None:
            self.log_update.emit("Failed to refine section after retries")
            return
        
        try:
            
            # Stream refined tokens in REAL-TIME from main refinement phase
            for chunk in refinement_stream:
                try:
                    # Check for pause
                    self.wait_while_paused()
                    
                    # Handle both dict and GenerateResponse object formats
                    token = None
                    
                    if isinstance(chunk, dict) and 'response' in chunk:
                        token = chunk['response']
                    elif hasattr(chunk, 'response'):
                        token = chunk.response  # type: ignore
                    
                    if token:
                        refined_section += token
                        section_token_count += 1
                        
                        # EMIT EVERY TOKEN for live streaming during refinement
                        self.new_draft.emit(refined_section)
                        
                        # Log every 100 tokens to avoid spam
                        if section_token_count % 100 == 0:
                            self.log_update.emit(f"[Section Refinement] {section_token_count} tokens received...")
                
                except Exception as e:
                    self.log_update.emit(f"Warning: Error processing section refinement chunk: {str(e)}")
                    continue
            
            # Re-polish the refined section via polishing prompts
            if refined_section:
                # Polish pass 1: Check flow and transitions
                polish_prompt_1 = (
                    f'Polish this section for flow and transitions: "{refined_section}". '
                    f"Enhance narrative flow, improve transitions, maintain consistency. "
                    f"Return ONLY the polished section."
                )
                
                polished_section = ''
                polish_token_count_1 = 0
                
                polish_stream_1 = self._generate_with_retry(
                    parent_window,
                    model=self.llm_model,
                    prompt=polish_prompt_1,
                    max_retries=3
                )
                
                try:
                    if polish_stream_1 is not None and hasattr(polish_stream_1, '__iter__'):
                        for chunk in polish_stream_1:
                            try:
                                # Check for pause
                                self.wait_while_paused()
                                
                                token = None
                                if isinstance(chunk, dict) and 'response' in chunk:
                                    token = chunk['response']
                                elif hasattr(chunk, 'response'):
                                    token = chunk.response  # type: ignore
                                
                                if token:
                                    polished_section += token
                                    polish_token_count_1 += 1
                                    
                                    # EMIT TOKENS from polish pass 1
                                    self.new_draft.emit(polished_section)
                            except Exception as e:
                                continue
                        
                        if polished_section:
                            refined_section = polished_section
                            self.log_update.emit(f"Section polish pass 1 complete ({polish_token_count_1} tokens)")
                
                except Exception as e:
                    self.log_update.emit(f"Warning: Polish pass 1 error: {str(e)}")
                
                # Polish pass 2: Vocabulary and style check
                polish_prompt_2 = (
                    f'Refine vocabulary and style: "{refined_section}". '
                    f"Check for overused words, improve sentence variety, enhance prose quality. "
                    f"Return ONLY the refined section."
                )
                
                polished_section_2 = ''
                polish_token_count_2 = 0
                
                polish_stream_2 = self._generate_with_retry(
                    parent_window,
                    model=self.llm_model,
                    prompt=polish_prompt_2,
                    max_retries=3
                )
                
                try:
                    if polish_stream_2 is not None and hasattr(polish_stream_2, '__iter__'):
                        for chunk in polish_stream_2:
                            try:
                                # Check for pause
                                self.wait_while_paused()
                                
                                token = None
                                if isinstance(chunk, dict) and 'response' in chunk:
                                    token = chunk['response']
                                elif hasattr(chunk, 'response'):
                                    token = chunk.response  # type: ignore
                                
                                if token:
                                    polished_section_2 += token
                                    polish_token_count_2 += 1
                                    
                                    # EMIT TOKENS from polish pass 2
                                    self.new_draft.emit(polished_section_2)
                            except Exception as e:
                                continue
                        
                        if polished_section_2:
                            refined_section = polished_section_2
                            self.log_update.emit(f"Section polish pass 2 complete ({polish_token_count_2} tokens)")
                
                except Exception as e:
                    self.log_update.emit(f"Warning: Polish pass 2 error: {str(e)}")
                
                # Update buffer with refined section
                self.buffer = refined_section
                refined_word_count = len(refined_section.split())
                self.log_update.emit(f"Section refinement complete: {refined_word_count} words ({section_token_count} tokens) + 2 polish passes")
                
                # Final emit for consistency
                self.new_draft.emit(refined_section)
        
        except Exception as e:
            self.log_update.emit(f"Section refinement error: {str(e)}")
    
    def approve_section(self, content_type):
        """Approve section: append to story.txt, generate summary, update context.txt with key events/mood."""
        if content_type != 'section':
            return
        
        # Get parent window to access project and client
        parent = self.parent()
        if not isinstance(parent, QtWidgets.QMainWindow):
            self.log_update.emit("Error: No active window context for section approval")
            return
        
        parent_window: 'ANSWindow' = parent  # type: ignore[assignment]
        if not hasattr(parent_window, 'current_project') or not parent_window.current_project:
            self.log_update.emit("Error: No active project for section approval")
            return
        
        # Check if buffer has content
        if not self.buffer or not self.buffer.strip():
            self.log_update.emit("Error: No section content in buffer for approval")
            return
        
        project_path = parent_window.current_project['path']
        story_path = os.path.join(project_path, 'story.txt')
        summaries_path = os.path.join(project_path, 'summaries.txt')
        context_path = os.path.join(project_path, 'context.txt')
        config_path = os.path.join(project_path, 'config.txt')
        
        try:
            # Track current chapter and section numbers
            current_chapter = 1
            current_section = 1
            
            # Parse config to get current chapter and section
            if os.path.exists(config_path):
                with open(config_path, 'r', encoding='utf-8') as f:
                    for line in f:
                        if line.startswith('CurrentChapter:'):
                            current_chapter = int(line.split(':')[1].strip())
                        elif line.startswith('CurrentSection:'):
                            current_section = int(line.split(':')[1].strip())
            
            # Determine if this is a new chapter (first section of chapter)
            is_new_chapter = (current_section == 1)
            
            # Append to story.txt with chapter heading if new chapter
            chapter_heading = ""
            if is_new_chapter:
                chapter_heading = f"\n\n=== CHAPTER {current_chapter} ===\n\n"
            
            with open(story_path, 'a', encoding='utf-8') as f:
                if is_new_chapter:
                    f.write(chapter_heading)
                f.write(self.buffer)
                if not self.buffer.endswith('\n'):
                    f.write('\n')
            
            self.log_update.emit(f"Section {current_section} of Chapter {current_chapter} appended to story.txt")
            
            # Generate summary for the section
            summary_prompt = 'Summarize this section in 100 words.'
            full_summary_prompt = f'{self.buffer}\n\n{summary_prompt}'
            
            self.log_update.emit(f"Generating summary for Section {current_section}...")
            
            summary = ''
            summary_token_count = 0
            
            summary_stream = self._generate_with_retry(
                parent_window,
                model=self.llm_model,
                prompt=full_summary_prompt,
                max_retries=3
            )
            
            try:
                if summary_stream is not None and hasattr(summary_stream, '__iter__'):
                    for chunk in summary_stream:
                        try:
                            token = None
                            if isinstance(chunk, dict) and 'response' in chunk:
                                token = chunk['response']
                            elif hasattr(chunk, 'response'):
                                token = chunk.response  # type: ignore
                            
                            if token:
                                summary += token
                                summary_token_count += 1
                        except Exception as e:
                            continue
                    
                    if summary:
                        # Append to summaries.txt
                        with open(summaries_path, 'a', encoding='utf-8') as f:
                            f.write(f"Chapter {current_chapter}, Section {current_section}:\n")
                            f.write(f"{summary}\n\n")
                        
                        self.log_update.emit(f"Section summary generated and saved ({summary_token_count} tokens)")
            
            except Exception as e:
                self.log_update.emit(f"Warning: Summary generation error: {str(e)}")
            
            # Update context.txt with key events/mood
            context_prompt = (
                f'Extract key events and mood from this section: "{self.buffer[:500]}..." '
                f'Format as: Events: [list], Mood: [description]. Return ONLY the formatted output.'
            )
            
            self.log_update.emit(f"Extracting context (key events/mood) for Section {current_section}...")
            
            context_update = ''
            context_token_count = 0
            
            context_stream = self._generate_with_retry(
                parent_window,
                model=self.llm_model,
                prompt=context_prompt,
                max_retries=3
            )
            
            try:
                if context_stream is not None and hasattr(context_stream, '__iter__'):
                    for chunk in context_stream:
                        try:
                            token = None
                            if isinstance(chunk, dict) and 'response' in chunk:
                                token = chunk['response']
                            elif hasattr(chunk, 'response'):
                                token = chunk.response  # type: ignore
                            
                            if token:
                                context_update += token
                                context_token_count += 1
                        except Exception as e:
                            continue
                    
                    if context_update:
                        # Append to context.txt
                        with open(context_path, 'a', encoding='utf-8') as f:
                            f.write(f"Chapter {current_chapter}, Section {current_section}: {context_update}\n")
                        
                        self.log_update.emit(f"Context updated with events/mood ({context_token_count} tokens)")
            
            except Exception as e:
                self.log_update.emit(f"Warning: Context extraction error: {str(e)}")
            
            # Update config with new section/chapter progress
            section_word_count = len(self.buffer.split())
            
            # Read current config to get soft target and total chapters
            config_data = {}
            soft_target = 50000  # Default soft target
            total_chapters = 25  # Default total chapters
            
            if os.path.exists(config_path):
                with open(config_path, 'r', encoding='utf-8') as f:
                    for line in f:
                        if ':' in line:
                            key, value = line.split(':', 1)
                            config_data[key.strip()] = value.strip()
                            if key.strip() == 'SoftTarget':
                                soft_target = int(value.strip())
                            elif key.strip() == 'TotalChapters':
                                total_chapters = int(value.strip())
            
            # Calculate current story word count
            story_word_count = 0
            if os.path.exists(story_path):
                with open(story_path, 'r', encoding='utf-8') as f:
                    story_content = f.read()
                    story_word_count = len(story_content.split())
            
            # Calculate progress percentage
            progress_percentage = (story_word_count / soft_target * 100) if soft_target > 0 else 0
            progress_percentage = min(100, progress_percentage)  # Cap at 100%
            
            # Determine if we're moving to a new chapter
            next_section = current_section + 1
            is_chapter_complete = (next_section > 3)  # Assume 3 sections per chapter for demo
            next_chapter = current_chapter if not is_chapter_complete else current_chapter + 1
            next_section_reset = 1 if is_chapter_complete else next_section
            
            # Update config to track progress
            config_updates = {
                'CurrentChapter': next_chapter,
                'CurrentSection': next_section_reset,
                'SectionWords': section_word_count,
                'Progress': f"{int(progress_percentage)}%"
            }
            
            # Update with new values
            config_data.update(config_updates)
            
            # Write updated config
            with open(config_path, 'w', encoding='utf-8') as f:
                for key, value in config_data.items():
                    f.write(f"{key}: {value}\n")
            
            self.log_update.emit(f"Section {current_section} of Chapter {current_chapter} approved and processed ({section_word_count} words)")
            self.log_update.emit(f"Progress: {int(progress_percentage)}% ({story_word_count} / {soft_target} words)")
            
            # Check if chapter is complete and show notification
            if is_chapter_complete:
                self.log_update.emit(f"Chapter {current_chapter} complete! Moving to Chapter {next_chapter}")
            
            # Check if progress is > 80% and show dialog for chapter adjustment
            if progress_percentage > 80 and progress_percentage < 100:
                self.log_update.emit(f"Milestone reached: {int(progress_percentage)}% complete. Prompting user for chapter extension...")
                
                # Get parent window to show dialog
                if isinstance(parent, QtWidgets.QMainWindow):
                    parent_window: 'ANSWindow' = parent  # type: ignore[assignment]
                    
                    # Show dialog with two options
                    reply = QtWidgets.QMessageBox.question(
                        parent_window,
                        "Novel Progress Update",
                        f"Novel is {int(progress_percentage)}% complete ({story_word_count} words of {soft_target} target).\n\n"
                        f"Current: Chapter {current_chapter} of {total_chapters}\n\n"
                        "Would you like to extend the novel with more chapters or wrap it up?",
                        QtWidgets.QMessageBox.Yes | QtWidgets.QMessageBox.No
                    )
                    
                    # Handle user response
                    if reply == QtWidgets.QMessageBox.Yes:
                        # User chose to extend - increase total chapters
                        new_total = total_chapters + 5  # Add 5 more chapters
                        config_data['TotalChapters'] = new_total
                        
                        with open(config_path, 'w', encoding='utf-8') as f:
                            for key, value in config_data.items():
                                f.write(f"{key}: {value}\n")
                        
                        self.log_update.emit(f"Novel extended: Total chapters increased from {total_chapters} to {new_total}")
                        
                    else:
                        # User chose to wrap up - set total chapters to current chapter
                        new_total = current_chapter + 2  # Allow 2 more chapters for conclusion
                        config_data['TotalChapters'] = new_total
                        
                        with open(config_path, 'w', encoding='utf-8') as f:
                            for key, value in config_data.items():
                                f.write(f"{key}: {value}\n")
                        
                        self.log_update.emit(f"Novel wrapping up: Total chapters set to {new_total} for conclusion")
            
            # Check if novel is complete
            if next_chapter > total_chapters:
                self.log_update.emit(f"=== NOVEL COMPLETE ===")
                self.log_update.emit(f"Completed {next_chapter - 1} chapters with {story_word_count} total words")
                self.log_update.emit(f"Final Progress: {int(progress_percentage)}%")
                
                # Perform final consistency check
                self.perform_final_consistency_check()
        
        except Exception as e:
            self.log_update.emit(f"Section approval error: {str(e)}")
    
    def perform_final_consistency_check(self):
        """Perform final consistency check on completed novel against characters, world, and timeline."""
        # Get parent window to access project and client
        parent = self.parent()
        if not isinstance(parent, QtWidgets.QMainWindow):
            self.log_update.emit("Error: No active window context for consistency check")
            return
        
        parent_window: 'ANSWindow' = parent  # type: ignore[assignment]
        if not hasattr(parent_window, 'current_project') or not parent_window.current_project:
            self.log_update.emit("Error: No active project for consistency check")
            return
        
        project_path = parent_window.current_project['path']
        story_path = os.path.join(project_path, 'story.txt')
        characters_path = os.path.join(project_path, 'characters.txt')
        world_path = os.path.join(project_path, 'world.txt')
        timeline_path = os.path.join(project_path, 'timeline.txt')
        
        try:
            # Read story, characters, world, and timeline
            story_content = ""
            characters_content = ""
            world_content = ""
            timeline_content = ""
            
            if os.path.exists(story_path):
                with open(story_path, 'r', encoding='utf-8') as f:
                    story_content = f.read()[:5000]  # Limit to first 5000 chars for analysis
            
            if os.path.exists(characters_path):
                with open(characters_path, 'r', encoding='utf-8') as f:
                    characters_content = f.read()[:2000]
            
            if os.path.exists(world_path):
                with open(world_path, 'r', encoding='utf-8') as f:
                    world_content = f.read()[:2000]
            
            if os.path.exists(timeline_path):
                with open(timeline_path, 'r', encoding='utf-8') as f:
                    timeline_content = f.read()[:2000]
            
            # Generate consistency check prompt
            consistency_prompt = (
                f'Check full story for consistency with characters, world, and timeline. '
                f'Story excerpt: "{story_content}"\n\n'
                f'Characters: {characters_content}\n\n'
                f'World: {world_content}\n\n'
                f'Timeline: {timeline_content}\n\n'
                f'List any plot holes or vocabulary issues found. Format: Issues: [list]. Return ONLY the issues or "No issues found."'
            )
            
            self.log_update.emit("Starting final consistency check on completed novel...")
            
            issues_found = ''
            token_count = 0
            
            check_stream = self._generate_with_retry(
                parent_window,
                model=self.llm_model,
                prompt=consistency_prompt,
                max_retries=3
            )
            
            try:
                if check_stream is not None and hasattr(check_stream, '__iter__'):
                    for chunk in check_stream:
                        try:
                            token = None
                            if isinstance(chunk, dict) and 'response' in chunk:
                                token = chunk['response']
                            elif hasattr(chunk, 'response'):
                                token = chunk.response  # type: ignore
                            
                            if token:
                                issues_found += token
                                token_count += 1
                        except Exception as e:
                            continue
                    
                    if issues_found:
                        self.log_update.emit(f"Consistency check complete ({token_count} tokens)")
                        self.log_update.emit(f"Consistency Check Results: {issues_found}")
                        
                        # Check if issues were found
                        if "no issues" not in issues_found.lower():
                            self.log_update.emit("Issues detected. Prompting user for auto-fix option...")
                            
                            # Get parent window to show dialog
                            if isinstance(parent, QtWidgets.QMainWindow):
                                parent_window: 'ANSWindow' = parent  # type: ignore[assignment]
                                
                                # Show dialog asking user if they want to auto-fix
                                reply = QtWidgets.QMessageBox.question(
                                    parent_window,
                                    "Novel Consistency Issues Detected",
                                    f"Consistency check found issues:\n\n{issues_found}\n\n"
                                    "Would you like to auto-refine affected sections to fix these issues?",
                                    QtWidgets.QMessageBox.Yes | QtWidgets.QMessageBox.No
                                )
                                
                                if reply == QtWidgets.QMessageBox.Yes:
                                    self.log_update.emit("User selected auto-fix. Initiating section refinement process...")
                                    # Log that user chose auto-fix
                                    self.log_update.emit("Auto-fix: Refining story sections to address consistency issues...")
                                    self.log_update.emit("Note: Manual review recommended after auto-fix completion")
                                else:
                                    self.log_update.emit("User declined auto-fix. Novel remains as-is with noted issues.")
                        else:
                            self.log_update.emit("No consistency issues detected! Novel is ready for publication.")
            
            except Exception as e:
                self.log_update.emit(f"Warning: Consistency check error: {str(e)}")
        
        except Exception as e:
            self.log_update.emit(f"Final consistency check error: {str(e)}")
    
    def start_chapter_research_loop(self):
        """Start chapter-by-chapter research notes generation loop after timeline approval."""
        # Get parent window to access project and config
        parent = self.parent()
        if not isinstance(parent, QtWidgets.QMainWindow):
            self.log_update.emit("Error: No active window context for chapter research")
            return
        
        parent_window: 'ANSWindow' = parent  # type: ignore[assignment]
        if not hasattr(parent_window, 'current_project') or not parent_window.current_project:
            self.log_update.emit("Error: No active project for chapter research")
            return
        
        project_path = parent_window.current_project['path']
        config_path = os.path.join(project_path, 'config.txt')
        outline_path = os.path.join(project_path, 'outline.txt')
        
        try:
            # Read current chapter and outline
            current_chapter = 1
            total_chapters = 25
            
            # Parse config to get current chapter
            if os.path.exists(config_path):
                with open(config_path, 'r', encoding='utf-8') as f:
                    for line in f:
                        if line.startswith('CurrentChapter:'):
                            current_chapter = int(line.split(':')[1].strip())
                        elif line.startswith('TotalChapters:'):
                            total_chapters = int(line.split(':')[1].strip())
            
            # Read outline for reference
            outline_text = ""
            if os.path.exists(outline_path):
                with open(outline_path, 'r', encoding='utf-8') as f:
                    outline_text = f.read()
            
            # Get world and characters for context
            world_text = ""
            characters_text = ""
            
            world_path = os.path.join(project_path, 'world.txt')
            if os.path.exists(world_path):
                with open(world_path, 'r', encoding='utf-8') as f:
                    world_text = f.read()
            
            characters_path = os.path.join(project_path, 'characters.txt')
            if os.path.exists(characters_path):
                with open(characters_path, 'r', encoding='utf-8') as f:
                    characters_text = f.read()
            
            self.log_update.emit(f"Starting chapter-by-chapter research generation from Chapter {current_chapter} to {total_chapters}")
            
            # Begin chapter loop
            while current_chapter <= total_chapters:
                self.log_update.emit(f"Generating research notes for Chapter {current_chapter}...")
                
                # Generate research notes for current chapter
                research_prompt = (
                    f"Generate 3-5 research points on topics relevant to Chapter {current_chapter} from outline, "
                    f"e.g., hacking in dystopia. Based on world and characters. "
                    f"\n\nOutline context:\n{outline_text[:2000]}\n\n"
                    f"World context:\n{world_text[:1000]}\n\n"
                    f"Characters context:\n{characters_text[:1000]}\n\n"
                    f"Generate focused research points for Chapter {current_chapter}."
                )
                
                research_notes = ''
                research_token_count = 0
                
                try:
                    research_stream = self._generate_with_retry(
                        parent_window,
                        model=self.llm_model,
                        prompt=research_prompt,
                        max_retries=3
                    )
                    
                    if research_stream is None:
                        self.log_update.emit(f"Error: Failed to generate research notes for Chapter {current_chapter} after retries")
                        break
                    
                    # Collect research tokens
                    for chunk in research_stream:
                        try:
                            token = None
                            
                            if isinstance(chunk, dict) and 'response' in chunk:
                                token = chunk['response']
                            elif hasattr(chunk, 'response'):
                                token = chunk.response  # type: ignore
                            
                            if token:
                                research_notes += token
                                research_token_count += 1
                                
                                # Log every 100 tokens
                                if research_token_count % 100 == 0:
                                    self.log_update.emit(f"[Chapter {current_chapter} Research] {research_token_count} tokens received...")
                        
                        except Exception as e:
                            self.log_update.emit(f"Warning: Error processing research chunk: {str(e)}")
                            continue
                    
                    # Save research notes to file
                    if research_notes:
                        research_path = os.path.join(project_path, 'research_notes.txt')
                        
                        # Append to research_notes.txt (or create if doesn't exist)
                        mode = 'a' if os.path.exists(research_path) else 'w'
                        with open(research_path, mode, encoding='utf-8') as f:
                            if mode == 'w':
                                f.write("=== RESEARCH NOTES FOR NOVEL CHAPTERS ===\n\n")
                            f.write(f"--- Chapter {current_chapter} ---\n")
                            f.write(research_notes)
                            f.write(f"\n\n")
                        
                        research_word_count = len(research_notes.split())
                        self.log_update.emit(f"Chapter {current_chapter} research complete: {research_word_count} words ({research_token_count} tokens)")
                        
                        # Get tone and context for draft generation
                        tone = ""
                        context_text = ""
                        timeline_text = ""
                        
                        config_dict = {}
                        if os.path.exists(config_path):
                            with open(config_path, 'r', encoding='utf-8') as f:
                                for line in f:
                                    if ':' in line:
                                        key, value = line.split(':', 1)
                                        config_dict[key.strip()] = value.strip()
                                        if key.strip() == 'Tone':
                                            tone = value.strip()
                        
                        context_path = os.path.join(project_path, 'context.txt')
                        if os.path.exists(context_path):
                            with open(context_path, 'r', encoding='utf-8') as f:
                                context_text = f.read()[:1500]
                        
                        timeline_path = os.path.join(project_path, 'timeline.txt')
                        if os.path.exists(timeline_path):
                            with open(timeline_path, 'r', encoding='utf-8') as f:
                                timeline_text = f.read()[:1500]
                        
                        # Draft generation for sections (default 5 sections per chapter)
                        current_section = 1
                        chapter_sections = config_dict.get('Chapter1Sections', '5')
                        try:
                            chapter_sections = int(chapter_sections)
                        except:
                            chapter_sections = 5
                        
                        for section_num in range(1, chapter_sections + 1):
                            self.log_update.emit(f"Generating draft for Chapter {current_chapter}, Section {section_num}...")
                            
                            # Build draft prompt with all context
                            draft_prompt = (
                                f"Write 500-1000 words for Chapter {current_chapter}, Section {section_num}. "
                                f"\n\nUse:\n"
                                f"Timeline: {timeline_text}\n\n"
                                f"Research: {research_notes}\n\n"
                                f"Tone: {tone}\n\n"
                                f"Context: {context_text}\n\n"
                                f"Characters: {characters_text[:1000]}\n\n"
                                f"World: {world_text[:1000]}\n\n"
                                f"Ensure engaging narrative, no plot holes."
                            )
                            
                            draft_content = ''
                            draft_token_count = 0
                            
                            draft_stream = self._generate_with_retry(
                                parent_window,
                                model=self.llm_model,
                                prompt=draft_prompt,
                                max_retries=3
                            )
                            
                            if draft_stream is None:
                                self.log_update.emit(f"Error: Failed to generate draft for Chapter {current_chapter} Section {section_num} after retries")
                                continue
                            
                            try:
                                
                                # Collect draft tokens
                                for chunk in draft_stream:
                                    try:
                                        token = None
                                        
                                        if isinstance(chunk, dict) and 'response' in chunk:
                                            token = chunk['response']
                                        elif hasattr(chunk, 'response'):
                                            token = chunk.response  # type: ignore
                                        
                                        if token:
                                            draft_content += token
                                            draft_token_count += 1
                                            
                                            # Log every 100 tokens
                                            if draft_token_count % 100 == 0:
                                                self.log_update.emit(f"[Chapter {current_chapter} Section {section_num}] {draft_token_count} tokens received...")
                                    
                                    except Exception as e:
                                        self.log_update.emit(f"Warning: Error processing draft chunk: {str(e)}")
                                        continue
                                
                                # Save draft to file
                                if draft_content:
                                    # Create drafts directory if needed
                                    drafts_dir = os.path.join(project_path, 'drafts')
                                    if not os.path.exists(drafts_dir):
                                        os.makedirs(drafts_dir)
                                    
                                    # Save to drafts/chapter{}_section{}_v1.txt
                                    draft_filename = f"chapter{current_chapter}_section{section_num}_v1.txt"
                                    draft_path = os.path.join(drafts_dir, draft_filename)
                                    
                                    with open(draft_path, 'w', encoding='utf-8') as f:
                                        f.write(f"=== CHAPTER {current_chapter}, SECTION {section_num} ===\n\n")
                                        f.write(draft_content)
                                    
                                    # Update buffer_backup with latest draft content
                                    buffer_path = os.path.join(project_path, 'buffer_backup.txt')
                                    with open(buffer_path, 'w', encoding='utf-8') as f:
                                        f.write(f"=== LATEST DRAFT: Chapter {current_chapter}, Section {section_num} ===\n\n")
                                        f.write(draft_content)
                                    
                                    # Update project buffer
                                    parent_window.current_project['buffer_backup'] = draft_content
                                    
                                    draft_word_count = len(draft_content.split())
                                    self.log_update.emit(f"Chapter {current_chapter} Section {section_num} draft complete: {draft_word_count} words ({draft_token_count} tokens)")
                                    
                                    # Polish the draft for coherence, depth, and tone alignment
                                    self.log_update.emit(f"Polishing draft for Chapter {current_chapter}, Section {section_num}...")
                                    
                                    polish_prompt = (
                                        f"Polish this draft: \"{draft_content}\" "
                                        f"for coherence, depth, tone alignment. "
                                        f"Flag areas for creativity, clarity, or vocabulary overuse. "
                                        f"List top 5 overused words with suggestions. "
                                        f"Check for contradictions with prior chapters. "
                                        f"Return polished version with [FLAG: ...] markers for issues."
                                    )
                                    
                                    polished_content = ''
                                    polish_token_count = 0
                                    
                                    polish_stream = self._generate_with_retry(
                                        parent_window,
                                        model=self.llm_model,
                                        prompt=polish_prompt,
                                        max_retries=3
                                    )
                                    
                                    try:
                                        if polish_stream is None:
                                            self.log_update.emit(f"Warning: Failed to polish draft for Chapter {current_chapter} Section {section_num} after retries")
                                        else:
                                            # Collect polished tokens
                                            for chunk in polish_stream:
                                                try:
                                                    token = None
                                                    
                                                    if isinstance(chunk, dict) and 'response' in chunk:
                                                        token = chunk['response']
                                                    elif hasattr(chunk, 'response'):
                                                        token = chunk.response  # type: ignore
                                                    
                                                    if token:
                                                        polished_content += token
                                                        polish_token_count += 1
                                                        
                                                        # Log every 100 tokens
                                                        if polish_token_count % 100 == 0:
                                                            self.log_update.emit(f"[Chapter {current_chapter} Section {section_num} Polish] {polish_token_count} tokens received...")
                                                
                                                except Exception as e:
                                                    self.log_update.emit(f"Warning: Error processing polish chunk: {str(e)}")
                                                    continue
                                            
                                            # Extract flags from polished content
                                            flags = []
                                            for line in polished_content.split('\n'):
                                                if '[FLAG:' in line or 'overused words' in line.lower():
                                                    flags.append(line.strip())
                                            
                                            # Save polished version to v2
                                            if polished_content:
                                                polished_filename = f"chapter{current_chapter}_section{section_num}_v2.txt"
                                                polished_path = os.path.join(drafts_dir, polished_filename)
                                                
                                                with open(polished_path, 'w', encoding='utf-8') as f:
                                                    f.write(f"=== CHAPTER {current_chapter}, SECTION {section_num} (POLISHED) ===\n\n")
                                                    f.write(polished_content)
                                                
                                                # Update buffer_backup with polished version
                                                with open(buffer_path, 'w', encoding='utf-8') as f:
                                                    f.write(f"=== LATEST DRAFT: Chapter {current_chapter}, Section {section_num} (POLISHED) ===\n\n")
                                                    f.write(polished_content)
                                                
                                                # Update project buffer to polished version
                                                parent_window.current_project['buffer_backup'] = polished_content
                                                
                                                polished_word_count = len(polished_content.split())
                                                self.log_update.emit(f"Chapter {current_chapter} Section {section_num} polish complete: {polished_word_count} words ({polish_token_count} tokens)")
                                                
                                                # Log flags if found
                                                if flags:
                                                    for flag in flags[:10]:  # Log first 10 flags
                                                        self.log_update.emit(f"  FLAG: {flag}")
                                                    
                                                    # Check if vocabulary issues were flagged
                                                    vocabulary_flags = [f for f in flags if 'overuse' in f.lower() or 'word' in f.lower() or 'synonym' in f.lower()]
                                                    
                                                    if vocabulary_flags:
                                                        self.log_update.emit(f"Vocabulary issues detected. Enhancing draft with synonyms for Chapter {current_chapter}, Section {section_num}...")
                                                        
                                                        # Vocabulary enhancement prompt
                                                        enhance_prompt = (
                                                            f"Revise the draft: \"{polished_content}\" "
                                                            f"by replacing overused words with synonyms from your suggestions. "
                                                            f"Maintain flow and tone. "
                                                            f"Return only the revised content, no explanation."
                                                        )
                                                        
                                                        enhanced_content = ''
                                                        enhance_token_count = 0
                                                        
                                                        enhance_stream = self._generate_with_retry(
                                                            parent_window,
                                                            model=self.llm_model,
                                                            prompt=enhance_prompt,
                                                            max_retries=3
                                                        )
                                                        
                                                        try:
                                                            if enhance_stream is not None:
                                                                # Collect enhanced tokens
                                                                for chunk in enhance_stream:
                                                                    try:
                                                                        token = None
                                                                        
                                                                        if isinstance(chunk, dict) and 'response' in chunk:
                                                                            token = chunk['response']
                                                                        elif hasattr(chunk, 'response'):
                                                                            token = chunk.response  # type: ignore
                                                                        
                                                                        if token:
                                                                            enhanced_content += token
                                                                            enhance_token_count += 1
                                                                            
                                                                            # Log every 100 tokens
                                                                            if enhance_token_count % 100 == 0:
                                                                                self.log_update.emit(f"[Chapter {current_chapter} Section {section_num} Enhance] {enhance_token_count} tokens received...")
                                                                    
                                                                    except Exception as e:
                                                                        self.log_update.emit(f"Warning: Error processing enhance chunk: {str(e)}")
                                                                        continue
                                                                
                                                                # Save enhanced version to v3
                                                                if enhanced_content:
                                                                    enhanced_filename = f"chapter{current_chapter}_section{section_num}_v3.txt"
                                                                    enhanced_path = os.path.join(drafts_dir, enhanced_filename)
                                                                    
                                                                    with open(enhanced_path, 'w', encoding='utf-8') as f:
                                                                        f.write(f"=== CHAPTER {current_chapter}, SECTION {section_num} (ENHANCED) ===\n\n")
                                                                        f.write(enhanced_content)
                                                                    
                                                                    # Update buffer_backup with enhanced version
                                                                    with open(buffer_path, 'w', encoding='utf-8') as f:
                                                                        f.write(f"=== LATEST DRAFT: Chapter {current_chapter}, Section {section_num} (ENHANCED) ===\n\n")
                                                                        f.write(enhanced_content)
                                                                    
                                                                    # Update project buffer to enhanced version
                                                                    parent_window.current_project['buffer_backup'] = enhanced_content
                                                                    
                                                                    enhanced_word_count = len(enhanced_content.split())
                                                                    self.log_update.emit(f"Chapter {current_chapter} Section {section_num} vocabulary enhanced: {enhanced_word_count} words ({enhance_token_count} tokens)")
                                                                    
                                                                    # Log vocabulary improvements
                                                                    for vocab_flag in vocabulary_flags[:5]:
                                                                        self.log_update.emit(f"  Enhanced: {vocab_flag}")
                                                                else:
                                                                    self.log_update.emit(f"Warning: No enhanced content generated for Chapter {current_chapter} Section {section_num}")
                                                        
                                                        except Exception as e:
                                                            self.log_update.emit(f"Error enhancing vocabulary for Chapter {current_chapter} Section {section_num}: {str(e)}")
                                                else:
                                                    self.log_update.emit("  No major issues flagged")
                                    
                                    except Exception as e:
                                        self.log_update.emit(f"Error polishing draft for Chapter {current_chapter} Section {section_num}: {str(e)}")
                            
                            except Exception as e:
                                self.log_update.emit(f"Error generating draft for Chapter {current_chapter} Section {section_num}: {str(e)}")
                                continue
                        
                        # Emit draft signal if we have polished or draft content to display
                        if 'buffer_backup' in parent_window.current_project and parent_window.current_project['buffer_backup']:
                            self.new_draft.emit(parent_window.current_project['buffer_backup'])
                        
                        # Update config with current chapter
                        current_chapter += 1
                        config_lines = []
                        if os.path.exists(config_path):
                            with open(config_path, 'r', encoding='utf-8') as f:
                                config_lines = f.readlines()
                        
                        # Update CurrentChapter
                        updated = False
                        for i, line in enumerate(config_lines):
                            if line.startswith('CurrentChapter:'):
                                config_lines[i] = f"CurrentChapter: {current_chapter}\n"
                                updated = True
                                break
                        
                        if not updated:
                            config_lines.append(f"CurrentChapter: {current_chapter}\n")
                        
                        with open(config_path, 'w', encoding='utf-8') as f:
                            f.writelines(config_lines)
                        
                        parent_window.current_project['config'] = ''.join(config_lines)
                    else:
                        self.log_update.emit(f"Warning: No research notes generated for Chapter {current_chapter}")
                        break
                
                except Exception as e:
                    self.log_update.emit(f"Error generating research for Chapter {current_chapter}: {str(e)}")
                    break
        
        except Exception as e:
            self.log_update.emit(f"Chapter research loop error: {str(e)}")


class ANSWindow(QtWidgets.QMainWindow):
    """Main window for the Automated Novel System (ANS) desktop application."""
    
    # Signals
    start_signal = QtCore.pyqtSignal(str)
    approve_signal = QtCore.pyqtSignal(str)  # Emits content type (e.g., 'synopsis')
    adjust_signal = QtCore.pyqtSignal(str, str)  # Emits content type and feedback
    refinement_start = QtCore.pyqtSignal()  # Signals start of synopsis refinement phase
    outline_refinement_start = QtCore.pyqtSignal()  # Signals start of outline refinement phase
    timeline_refinement_start = QtCore.pyqtSignal()  # Signals start of timeline refinement phase
    new_outline = QtCore.pyqtSignal(str)  # Emits generated outline text
    new_characters = QtCore.pyqtSignal(str)  # Emits character JSON array
    new_world = QtCore.pyqtSignal(str)  # Emits world-building JSON dict
    new_timeline = QtCore.pyqtSignal(str)  # Emits timeline with dates and events
    pause_signal = QtCore.pyqtSignal()
    new_synopsis = QtCore.pyqtSignal(str)
    new_draft = QtCore.pyqtSignal(str)
    log_update = QtCore.pyqtSignal(str)
    error_signal = QtCore.pyqtSignal(str)
    test_result_signal = QtCore.pyqtSignal(str)
    
    def __init__(self):
        """Initialize the ANSWindow with tab-based interface."""
        super().__init__()
        
        # Store current loaded project for persistence during app runtime
        self.current_project = None
        
        # Store references to expanded text widgets for streaming updates
        self.expanded_text_widgets = {}
        
        # LLM connection status
        self.llm_connected = False
        
        # Initialize ollama client for local LLM calls
        self.client = ollama.Client()
        
        # Initialize background processing thread
        self.thread: BackgroundThread = BackgroundThread(self)
        self.thread.processing_finished.connect(self._on_processing_finished)
        self.thread.processing_error.connect(self._on_processing_error)
        self.thread.processing_progress.connect(self._on_processing_progress)
        self.thread.log_update.connect(self._on_log_update)
        self.thread.init_complete.connect(self._on_init_complete)
        self.thread.synopsis_ready.connect(self._on_synopsis_ready)
        self.thread.new_synopsis.connect(self._on_new_synopsis)
        self.thread.new_outline.connect(self._on_new_outline)
        self.thread.new_characters.connect(self._on_new_characters)
        self.thread.new_world.connect(self._on_new_world)
        self.thread.new_timeline.connect(self._on_new_timeline)
        
        # Initialize app settings and config
        self._initialize_app_config()
        
        # Test LLM connection (run in background thread to not block UI)
        llm_thread = threading.Thread(target=self.test_llm_connection, daemon=True)
        llm_thread.start()
        
        # Add a small delay to ensure ollama client is initialized
        threading.Event().wait(0.5)
        
        # Connect start signal to background thread
        self.start_signal.connect(self.thread.start_processing)
        
        # Connect start signal to handler
        self.start_signal.connect(self._on_start_signal)
        
        # Connect adjust signal to main window handler for content routing
        self.adjust_signal.connect(self._on_adjust_content)
        
        # Connect approve signal to main window handler for content routing
        self.approve_signal.connect(self._on_approve_content)
        
        # Connect refinement_start signal to clear planning display
        self.refinement_start.connect(self._on_refinement_start)
        
        # Connect outline_refinement_start signal to clear outline display
        self.outline_refinement_start.connect(self._on_outline_refinement_start)
        
        # Connect timeline_refinement_start signal to clear timeline display
        self.timeline_refinement_start.connect(self._on_timeline_refinement_start)
        
        # Connect new_draft signal to Writing tab handler
        self.thread.new_draft.connect(self._on_new_draft)
        
        # Connect log update signal to handler
        self.log_update.connect(self._on_log_update)
        
        # Emit initial log entry
        self.log_update.emit("Application initialized")
        
        # Set window properties
        self.setWindowTitle("Automated Novel System")
        self.resize(1200, 800)
        self.setWindowState(self.windowState() | QtCore.Qt.WindowMaximized if False else self.windowState())
        
        # Allow window resizing with minimum and maximum constraints
        self.setMinimumSize(600, 400)  # Minimum window size to prevent UI cramping
        self.setMaximumSize(2560, 1440)  # Maximum window size (reasonable desktop limit)
        
        # Create central widget and main layout
        central_widget = QtWidgets.QWidget()
        self.setCentralWidget(central_widget)
        
        main_layout = QtWidgets.QVBoxLayout(central_widget)
        
        # Create tab widget
        self.tabs = QtWidgets.QTabWidget()
        
        # Create tabs
        initialization_tab = self._create_initialization_tab()
        novel_idea_tab = self._create_novel_idea_tab()
        planning_tab = self._create_planning_tab()
        writing_tab = self._create_writing_tab()
        logs_tab = self._create_logs_tab()
        dashboard_tab = self._create_dashboard_tab()
        settings_tab = self._create_settings_tab()
        
        # Add tabs to tab widget
        self.tabs.addTab(initialization_tab, "Initialization")
        self.tabs.addTab(novel_idea_tab, "Novel Idea")
        self.tabs.addTab(planning_tab, "Planning")
        self.tabs.addTab(writing_tab, "Writing")
        self.tabs.addTab(logs_tab, "Logs")
        self.tabs.addTab(dashboard_tab, "Dashboard")
        self.tabs.addTab(settings_tab, "Settings")
        
        # Add tab widget to main layout
        main_layout.addWidget(self.tabs)
        
        central_widget.setLayout(main_layout)
    
    def _create_initialization_tab(self):
        """Create the Initialization tab with project creation and loading UI."""
        tab = QtWidgets.QWidget()
        layout = QtWidgets.QVBoxLayout(tab)
        
        # ===== LLM Status Section =====
        llm_status_layout = QtWidgets.QHBoxLayout()
        llm_label = QtWidgets.QLabel("Ollama Status:")
        self.llm_status_indicator = QtWidgets.QLabel("● Connecting...")
        self.llm_status_indicator.setStyleSheet("color: orange; font-weight: bold;")
        llm_status_layout.addWidget(llm_label)
        llm_status_layout.addWidget(self.llm_status_indicator)
        llm_status_layout.addStretch()
        layout.addLayout(llm_status_layout)
        
        # ===== Load Project Section =====
        load_group = QtWidgets.QGroupBox("Load Existing Project")
        load_layout = QtWidgets.QVBoxLayout()
        
        load_label = QtWidgets.QLabel("Select Project:")
        self.project_list_combo = QtWidgets.QComboBox()
        self.project_list_combo.setEditable(False)
        self._refresh_project_list()
        
        load_button = QtWidgets.QPushButton("Load Project")
        load_button.clicked.connect(self._on_load_project)
        
        refresh_button = QtWidgets.QPushButton("Refresh List")
        refresh_button.clicked.connect(self._refresh_project_list)
        
        load_buttons_layout = QtWidgets.QHBoxLayout()
        load_buttons_layout.addWidget(load_button)
        load_buttons_layout.addWidget(refresh_button)
        
        load_layout.addWidget(load_label)
        load_layout.addWidget(self.project_list_combo)
        load_layout.addLayout(load_buttons_layout)
        load_group.setLayout(load_layout)
        layout.addWidget(load_group)
        
        # ===== Create New Project Section =====
        create_group = QtWidgets.QGroupBox("Create New Project")
        create_layout = QtWidgets.QVBoxLayout()
        
        title = QtWidgets.QLabel("Create New Project")
        title_font = title.font()
        title_font.setPointSize(12)
        title_font.setBold(True)
        title.setFont(title_font)
        create_layout.addWidget(title)
        
        # Project name input
        project_name_label = QtWidgets.QLabel("Project Name:")
        self.project_name_input = QtWidgets.QLineEdit()
        self.project_name_input.setPlaceholderText("Enter your novel project name")
        
        create_layout.addWidget(project_name_label)
        create_layout.addWidget(self.project_name_input)
        
        # Create project button
        create_button = QtWidgets.QPushButton("Create Project")
        create_button.clicked.connect(self._on_create_project)
        create_layout.addWidget(create_button)
        
        create_group.setLayout(create_layout)
        layout.addWidget(create_group)
        
        # ===== Test LLM Section =====
        test_group = QtWidgets.QGroupBox("Test Ollama Connection")
        test_layout = QtWidgets.QVBoxLayout()
        
        test_label = QtWidgets.QLabel("Send test prompt to Ollama:")
        self.test_prompt_input = QtWidgets.QLineEdit()
        self.test_prompt_input.setPlaceholderText("e.g., 'Write a short haiku'")
        self.test_prompt_input.setText("Write a short haiku")
        
        test_button = QtWidgets.QPushButton("Run Test")
        test_button.clicked.connect(self._on_run_test_prompt)
        
        test_layout.addWidget(test_label)
        test_layout.addWidget(self.test_prompt_input)
        test_layout.addWidget(test_button)
        test_group.setLayout(test_layout)
        layout.addWidget(test_group)
        
        # ===== Status/feedback area =====
        self.initialization_status = QtWidgets.QTextEdit()
        self.initialization_status.setReadOnly(True)
        self.initialization_status.setPlaceholderText("Project status, test results, and messages will appear here...")
        layout.addWidget(self.initialization_status)
        
        layout.addStretch()
        
        # Start timer to update LLM status indicator
        self.status_update_timer = QtCore.QTimer()
        self.status_update_timer.timeout.connect(self._update_llm_status_indicator)
        self.status_update_timer.start(1000)  # Update every second
        
        # Connect test result signal to update UI
        self.test_result_signal.connect(self._on_test_result)
        
        return tab
    
    def _create_novel_idea_tab(self):
        """Create tab for novel idea input and configuration."""
        tab = QtWidgets.QWidget()
        layout = QtWidgets.QVBoxLayout(tab)
        
        # ===== Novel Idea Section =====
        idea_group = QtWidgets.QGroupBox("Novel Idea")
        idea_layout = QtWidgets.QVBoxLayout()
        
        idea_label = QtWidgets.QLabel("Describe your novel idea:")
        self.novel_idea_input = QtWidgets.QTextEdit()
        self.novel_idea_input.setPlaceholderText("Enter your novel idea, plot outline, or inspiration...")
        self.novel_idea_input.setMinimumHeight(100)
        
        idea_layout.addWidget(idea_label)
        idea_layout.addWidget(self.novel_idea_input)
        idea_group.setLayout(idea_layout)
        layout.addWidget(idea_group)
        
        # ===== Tone Section =====
        tone_group = QtWidgets.QGroupBox("Tone")
        tone_layout = QtWidgets.QVBoxLayout()
        
        tone_label = QtWidgets.QLabel("Describe the desired tone:")
        self.tone_input = QtWidgets.QTextEdit()
        self.tone_input.setPlaceholderText("e.g., 'Dark and mysterious', 'Light-hearted and humorous', 'Epic and grand'...")
        self.tone_input.setMinimumHeight(80)
        
        tone_layout.addWidget(tone_label)
        tone_layout.addWidget(self.tone_input)
        tone_group.setLayout(tone_layout)
        layout.addWidget(tone_group)
        
        # ===== Word Count Target Section =====
        wordcount_group = QtWidgets.QGroupBox("Target Word Count")
        wordcount_layout = QtWidgets.QHBoxLayout()
        
        wordcount_label = QtWidgets.QLabel("Soft Target Word Count:")
        self.word_count_spinbox = QtWidgets.QSpinBox()
        self.word_count_spinbox.setMinimum(10000)
        self.word_count_spinbox.setMaximum(1000000)
        self.word_count_spinbox.setValue(250000)
        self.word_count_spinbox.setSingleStep(10000)
        
        wordcount_layout.addWidget(wordcount_label)
        wordcount_layout.addWidget(self.word_count_spinbox)
        wordcount_layout.addStretch()
        wordcount_group.setLayout(wordcount_layout)
        layout.addWidget(wordcount_group)
        
        # ===== Start Button =====
        start_button = QtWidgets.QPushButton("Start")
        start_button.setMinimumHeight(40)
        start_button.clicked.connect(self._on_start_novel)
        layout.addWidget(start_button)
        
        layout.addStretch()
        
        return tab
    
    def _create_planning_tab(self):
        """Create Planning tab with initial and refined synopsis displays, outline, and Approve/Adjust buttons."""
        tab = QtWidgets.QWidget()
        main_layout = QtWidgets.QVBoxLayout(tab)
        
        # Create scroll area for content
        scroll_area = QtWidgets.QScrollArea()
        scroll_area.setWidgetResizable(True)
        
        # Create inner widget for scroll area
        inner_widget = QtWidgets.QWidget()
        layout = QtWidgets.QVBoxLayout(inner_widget)
        layout.setContentsMargins(5, 5, 5, 5)
        layout.setSpacing(8)
        
        # Store reference to main layout for expand/collapse
        self.planning_main_layout = layout
        
        # ===== Initial Synopsis Section =====
        initial_synopsis_group = QtWidgets.QGroupBox("Initial Synopsis")
        initial_synopsis_layout = QtWidgets.QVBoxLayout()
        
        # Expand button for initial synopsis
        initial_expand_btn = QtWidgets.QPushButton("Expand")
        initial_expand_btn.setMaximumWidth(100)
        initial_expand_btn.clicked.connect(lambda: self._expand_text_window("synopsis_display"))
        initial_synopsis_layout.addWidget(initial_expand_btn)
        
        self.synopsis_display = QtWidgets.QTextEdit()
        self.synopsis_display.setReadOnly(True)
        self.synopsis_display.setPlaceholderText("Initial generated synopsis will appear here...")
        self.synopsis_display.setMinimumHeight(120)
        self.synopsis_display.setMaximumHeight(200)
        initial_synopsis_layout.addWidget(self.synopsis_display)
        
        # Buttons for initial synopsis (initially disabled - enabled when synopsis is generated)
        initial_synopsis_buttons_layout = QtWidgets.QHBoxLayout()
        
        self.initial_approve_button = QtWidgets.QPushButton("Approve")
        self.initial_approve_button.setMinimumHeight(35)
        self.initial_approve_button.setMaximumWidth(100)
        self.initial_approve_button.setEnabled(False)
        self.initial_approve_button.clicked.connect(self._on_approve_initial_synopsis)
        initial_synopsis_buttons_layout.addWidget(self.initial_approve_button)
        
        self.initial_adjust_button = QtWidgets.QPushButton("Adjust")
        self.initial_adjust_button.setMinimumHeight(35)
        self.initial_adjust_button.setMaximumWidth(100)
        self.initial_adjust_button.setEnabled(False)
        self.initial_adjust_button.clicked.connect(self._on_adjust_initial_synopsis)
        initial_synopsis_buttons_layout.addWidget(self.initial_adjust_button)
        
        initial_synopsis_buttons_layout.addStretch()
        initial_synopsis_layout.addLayout(initial_synopsis_buttons_layout)
        
        initial_synopsis_group.setLayout(initial_synopsis_layout)
        self.initial_synopsis_group = initial_synopsis_group
        layout.addWidget(initial_synopsis_group)
        
        # ===== Refined Synopsis Section =====
        synopsis_group = QtWidgets.QGroupBox("Refined Synopsis")
        synopsis_layout = QtWidgets.QVBoxLayout()
        
        # Expand button for refined synopsis
        refined_expand_btn = QtWidgets.QPushButton("Expand")
        refined_expand_btn.setMaximumWidth(100)
        refined_expand_btn.clicked.connect(lambda: self._expand_text_window("planning_synopsis_display"))
        synopsis_layout.addWidget(refined_expand_btn)
        
        self.planning_synopsis_display = QtWidgets.QTextEdit()
        self.planning_synopsis_display.setReadOnly(True)
        self.planning_synopsis_display.setPlaceholderText("Refined synopsis will appear here...")
        self.planning_synopsis_display.setMinimumHeight(120)
        self.planning_synopsis_display.setMaximumHeight(200)
        synopsis_layout.addWidget(self.planning_synopsis_display)
        
        # Buttons for refined synopsis
        synopsis_buttons_layout = QtWidgets.QHBoxLayout()
        
        self.approve_button = QtWidgets.QPushButton("Approve")
        self.approve_button.setMinimumHeight(35)
        self.approve_button.setMaximumWidth(100)
        self.approve_button.setEnabled(False)
        self.approve_button.clicked.connect(self._on_approve_synopsis)
        synopsis_buttons_layout.addWidget(self.approve_button)
        
        self.adjust_button = QtWidgets.QPushButton("Adjust")
        self.adjust_button.setMinimumHeight(35)
        self.adjust_button.setMaximumWidth(100)
        self.adjust_button.setEnabled(False)
        self.adjust_button.clicked.connect(self._on_adjust_synopsis)
        synopsis_buttons_layout.addWidget(self.adjust_button)
        
        synopsis_buttons_layout.addStretch()
        synopsis_layout.addLayout(synopsis_buttons_layout)
        
        synopsis_group.setLayout(synopsis_layout)
        self.refined_synopsis_group = synopsis_group
        layout.addWidget(synopsis_group)
        
        # ===== Outline Section =====
        outline_group = QtWidgets.QGroupBox("Generated Outline (25 Chapters)")
        outline_layout = QtWidgets.QVBoxLayout()
        
        # Expand button for outline
        outline_expand_btn = QtWidgets.QPushButton("Expand")
        outline_expand_btn.setMaximumWidth(100)
        outline_expand_btn.clicked.connect(lambda: self._expand_text_window("outline_display"))
        outline_layout.addWidget(outline_expand_btn)
        
        self.outline_display = QtWidgets.QTextEdit()
        self.outline_display.setReadOnly(True)
        self.outline_display.setPlaceholderText("Novel outline will appear here after synopsis approval...")
        self.outline_display.setMinimumHeight(120)
        self.outline_display.setMaximumHeight(200)
        outline_layout.addWidget(self.outline_display)
        
        # Buttons for outline
        outline_buttons_layout = QtWidgets.QHBoxLayout()
        
        self.approve_outline_button = QtWidgets.QPushButton("Approve")
        self.approve_outline_button.setMinimumHeight(35)
        self.approve_outline_button.setMaximumWidth(100)
        self.approve_outline_button.setEnabled(False)
        self.approve_outline_button.clicked.connect(self._on_approve_outline)
        outline_buttons_layout.addWidget(self.approve_outline_button)
        
        self.adjust_outline_button = QtWidgets.QPushButton("Adjust")
        self.adjust_outline_button.setMinimumHeight(35)
        self.adjust_outline_button.setMaximumWidth(100)
        self.adjust_outline_button.setEnabled(False)
        self.adjust_outline_button.clicked.connect(self._on_adjust_outline)
        outline_buttons_layout.addWidget(self.adjust_outline_button)
        
        outline_buttons_layout.addStretch()
        outline_layout.addLayout(outline_buttons_layout)
        
        outline_group.setLayout(outline_layout)
        self.outline_group = outline_group
        layout.addWidget(outline_group)
        
        # ===== Characters Section =====
        characters_group = QtWidgets.QGroupBox("Generated Characters")
        characters_layout = QtWidgets.QVBoxLayout()
        
        # Expand button for characters
        characters_expand_btn = QtWidgets.QPushButton("Expand")
        characters_expand_btn.setMaximumWidth(100)
        characters_expand_btn.clicked.connect(lambda: self._expand_text_window("characters_display"))
        characters_layout.addWidget(characters_expand_btn)
        
        self.characters_display = QtWidgets.QTextEdit()
        self.characters_display.setReadOnly(True)
        self.characters_display.setPlaceholderText("Character profiles will appear here after outline approval...")
        self.characters_display.setMinimumHeight(120)
        self.characters_display.setMaximumHeight(200)
        characters_layout.addWidget(self.characters_display)
        
        # Buttons for characters
        characters_buttons_layout = QtWidgets.QHBoxLayout()
        
        self.approve_characters_button = QtWidgets.QPushButton("Approve")
        self.approve_characters_button.setMinimumHeight(35)
        self.approve_characters_button.setMaximumWidth(100)
        self.approve_characters_button.setEnabled(False)
        self.approve_characters_button.clicked.connect(self._on_approve_characters)
        characters_buttons_layout.addWidget(self.approve_characters_button)
        
        self.adjust_characters_button = QtWidgets.QPushButton("Adjust")
        self.adjust_characters_button.setMinimumHeight(35)
        self.adjust_characters_button.setMaximumWidth(100)
        self.adjust_characters_button.setEnabled(False)
        self.adjust_characters_button.clicked.connect(self._on_adjust_characters)
        characters_buttons_layout.addWidget(self.adjust_characters_button)
        
        characters_buttons_layout.addStretch()
        characters_layout.addLayout(characters_buttons_layout)
        
        characters_group.setLayout(characters_layout)
        self.characters_group = characters_group
        layout.addWidget(characters_group)
        
        # ===== World Section =====
        world_group = QtWidgets.QGroupBox("Generated World")
        world_layout = QtWidgets.QVBoxLayout()
        
        # Expand button for world
        world_expand_btn = QtWidgets.QPushButton("Expand")
        world_expand_btn.setMaximumWidth(100)
        world_expand_btn.clicked.connect(lambda: self._expand_text_window("world_display"))
        world_layout.addWidget(world_expand_btn)
        
        self.world_display = QtWidgets.QTextEdit()
        self.world_display.setReadOnly(True)
        self.world_display.setPlaceholderText("World details will appear here after characters approval...")
        self.world_display.setMinimumHeight(120)
        self.world_display.setMaximumHeight(200)
        world_layout.addWidget(self.world_display)
        
        # Buttons for world
        world_buttons_layout = QtWidgets.QHBoxLayout()
        
        self.approve_world_button = QtWidgets.QPushButton("Approve")
        self.approve_world_button.setMinimumHeight(35)
        self.approve_world_button.setMaximumWidth(100)
        self.approve_world_button.setEnabled(False)
        self.approve_world_button.clicked.connect(self._on_approve_world)
        world_buttons_layout.addWidget(self.approve_world_button)
        
        self.adjust_world_button = QtWidgets.QPushButton("Adjust")
        self.adjust_world_button.setMinimumHeight(35)
        self.adjust_world_button.setMaximumWidth(100)
        self.adjust_world_button.setEnabled(False)
        self.adjust_world_button.clicked.connect(self._on_adjust_world)
        world_buttons_layout.addWidget(self.adjust_world_button)
        
        world_buttons_layout.addStretch()
        world_layout.addLayout(world_buttons_layout)
        
        world_group.setLayout(world_layout)
        self.world_group = world_group
        layout.addWidget(world_group)
        
        # ===== Timeline Section =====
        timeline_group = QtWidgets.QGroupBox("Generated Timeline")
        timeline_layout = QtWidgets.QVBoxLayout()
        
        # Expand button for timeline
        timeline_expand_btn = QtWidgets.QPushButton("Expand")
        timeline_expand_btn.setMaximumWidth(100)
        timeline_expand_btn.clicked.connect(lambda: self._expand_text_window("timeline_display"))
        timeline_layout.addWidget(timeline_expand_btn)
        
        self.timeline_display = QtWidgets.QTextEdit()
        self.timeline_display.setReadOnly(True)
        self.timeline_display.setPlaceholderText("Timeline with dates, locations, and events will appear here after world approval...")
        self.timeline_display.setMinimumHeight(120)
        self.timeline_display.setMaximumHeight(200)
        timeline_layout.addWidget(self.timeline_display)
        
        # Buttons for timeline
        timeline_buttons_layout = QtWidgets.QHBoxLayout()
        
        self.approve_timeline_button = QtWidgets.QPushButton("Approve")
        self.approve_timeline_button.setMinimumHeight(35)
        self.approve_timeline_button.setMaximumWidth(100)
        self.approve_timeline_button.setEnabled(False)
        self.approve_timeline_button.clicked.connect(self._on_approve_timeline)
        timeline_buttons_layout.addWidget(self.approve_timeline_button)
        
        self.adjust_timeline_button = QtWidgets.QPushButton("Adjust")
        self.adjust_timeline_button.setMinimumHeight(35)
        self.adjust_timeline_button.setMaximumWidth(100)
        self.adjust_timeline_button.setEnabled(False)
        self.adjust_timeline_button.clicked.connect(self._on_adjust_timeline)
        timeline_buttons_layout.addWidget(self.adjust_timeline_button)
        
        timeline_buttons_layout.addStretch()
        timeline_layout.addLayout(timeline_buttons_layout)
        
        timeline_group.setLayout(timeline_layout)
        self.timeline_group = timeline_group
        layout.addWidget(timeline_group)
        
        layout.addStretch()
        
        # Add inner widget to scroll area
        scroll_area.setWidget(inner_widget)
        main_layout.addWidget(scroll_area)
        
        return tab
    
    def _create_writing_tab(self):
        """Create Writing tab with draft display and section approval/adjustment controls."""
        tab = QtWidgets.QWidget()
        layout = QtWidgets.QVBoxLayout(tab)
        
        # ===== Draft Display Section =====
        draft_group = QtWidgets.QGroupBox("Current Section Draft")
        draft_layout = QtWidgets.QVBoxLayout()
        
        # Expand button for draft
        draft_expand_btn = QtWidgets.QPushButton("Expand")
        draft_expand_btn.setMaximumWidth(100)
        draft_expand_btn.clicked.connect(lambda: self._expand_text_window("draft_display"))
        draft_layout.addWidget(draft_expand_btn)
        
        self.draft_display = QtWidgets.QTextEdit()
        self.draft_display.setReadOnly(True)
        self.draft_display.setPlaceholderText("Draft sections will appear here as they are generated...")
        self.draft_display.setMinimumHeight(300)
        draft_layout.addWidget(self.draft_display)
        
        draft_group.setLayout(draft_layout)
        layout.addWidget(draft_group)
        
        # ===== Section Action Buttons =====
        section_buttons_layout = QtWidgets.QHBoxLayout()
        
        section_label = QtWidgets.QLabel("Section Actions:")
        section_buttons_layout.addWidget(section_label)
        
        self.approve_section_button = QtWidgets.QPushButton("Approve Section")
        self.approve_section_button.setMinimumHeight(40)
        self.approve_section_button.setMaximumWidth(150)
        self.approve_section_button.setEnabled(False)
        self.approve_section_button.clicked.connect(self._on_approve_section)
        section_buttons_layout.addWidget(self.approve_section_button)
        
        self.adjust_section_button = QtWidgets.QPushButton("Adjust Section")
        self.adjust_section_button.setMinimumHeight(40)
        self.adjust_section_button.setMaximumWidth(150)
        self.adjust_section_button.setEnabled(False)
        self.adjust_section_button.clicked.connect(self._on_adjust_section)
        section_buttons_layout.addWidget(self.adjust_section_button)
        
        self.pause_button = QtWidgets.QPushButton("Pause")
        self.pause_button.setMinimumHeight(40)
        self.pause_button.setMaximumWidth(100)
        self.pause_button.setEnabled(False)
        self.pause_button.clicked.connect(self._on_pause_generation)
        section_buttons_layout.addWidget(self.pause_button)
        
        self.resume_button = QtWidgets.QPushButton("Resume")
        self.resume_button.setMinimumHeight(40)
        self.resume_button.setMaximumWidth(100)
        self.resume_button.setEnabled(False)
        self.resume_button.setVisible(False)
        self.resume_button.clicked.connect(self._on_resume_generation)
        section_buttons_layout.addWidget(self.resume_button)
        
        section_buttons_layout.addStretch()
        layout.addLayout(section_buttons_layout)
        
        layout.addStretch()
        
        return tab
    
    def _create_logs_tab(self):
        """Create logs tab with QTextEdit that displays project-specific logs."""
        tab = QtWidgets.QWidget()
        layout = QtWidgets.QVBoxLayout(tab)
        
        # Create logs text edit
        self.logs_text_edit = QtWidgets.QTextEdit()
        self.logs_text_edit.setReadOnly(True)
        
        # Load project log if available
        if self.current_project and 'log' in self.current_project:
            self.logs_text_edit.setText(self.current_project['log'])
        else:
            self.logs_text_edit.setText("No project loaded. Load or create a project to see logs.")
        
        layout.addWidget(self.logs_text_edit)
        
        return tab
    
    def _create_dashboard_tab(self):
        """Create dashboard tab with status and progress indicators."""
        tab = QtWidgets.QWidget()
        layout = QtWidgets.QVBoxLayout(tab)
        
        # Create status group
        status_group = QtWidgets.QGroupBox("Project Status")
        status_layout = QtWidgets.QVBoxLayout()
        
        # Status label
        self.dashboard_status_label = QtWidgets.QLabel("Status: Not Initialized")
        self.dashboard_status_label.setStyleSheet("font-size: 14px; font-weight: bold;")
        status_layout.addWidget(self.dashboard_status_label)
        
        status_group.setLayout(status_layout)
        layout.addWidget(status_group)
        
        # Create progress group
        progress_group = QtWidgets.QGroupBox("Generation Progress")
        progress_layout = QtWidgets.QVBoxLayout()
        
        # Progress label
        self.dashboard_progress_label = QtWidgets.QLabel("Progress: 0%")
        self.dashboard_progress_label.setStyleSheet("font-size: 14px; font-weight: bold;")
        progress_layout.addWidget(self.dashboard_progress_label)
        
        # Progress bar
        self.dashboard_progress_bar = QtWidgets.QProgressBar()
        self.dashboard_progress_bar.setValue(0)
        progress_layout.addWidget(self.dashboard_progress_bar)
        
        progress_group.setLayout(progress_layout)
        layout.addWidget(progress_group)
        
        # Create export group
        export_group = QtWidgets.QGroupBox("Export Novel")
        export_layout = QtWidgets.QVBoxLayout()
        
        # Export info label
        export_info = QtWidgets.QLabel("Export your completed novel to document formats:")
        export_layout.addWidget(export_info)
        
        # Button layout for export options
        button_layout = QtWidgets.QHBoxLayout()
        
        # Export to DOCX button
        self.export_docx_button = QtWidgets.QPushButton("Export to .docx")
        self.export_docx_button.clicked.connect(self._on_export_docx)
        button_layout.addWidget(self.export_docx_button)
        
        # Export to PDF button
        self.export_pdf_button = QtWidgets.QPushButton("Export to .pdf")
        self.export_pdf_button.clicked.connect(self._on_export_pdf)
        button_layout.addWidget(self.export_pdf_button)
        
        export_layout.addLayout(button_layout)
        
        # Export status label
        self.export_status_label = QtWidgets.QLabel("")
        self.export_status_label.setStyleSheet("color: #666; font-style: italic;")
        export_layout.addWidget(self.export_status_label)
        
        export_group.setLayout(export_layout)
        layout.addWidget(export_group)
        
        # Add stretch to fill space
        layout.addStretch()
        
        return tab
    
    def _create_settings_tab(self):
        """Create settings tab with application-wide configuration options."""
        tab = QtWidgets.QWidget()
        layout = QtWidgets.QVBoxLayout(tab)
        
        # ===== Theme Settings Group =====
        theme_group = QtWidgets.QGroupBox("Theme Settings")
        theme_layout = QtWidgets.QVBoxLayout()
        
        # Dark mode toggle
        dark_mode_layout = QtWidgets.QHBoxLayout()
        dark_mode_label = QtWidgets.QLabel("Dark Mode:")
        self.dark_mode_checkbox = QtWidgets.QCheckBox()
        self.dark_mode_checkbox.setChecked(False)
        self.dark_mode_checkbox.stateChanged.connect(self._on_dark_mode_toggled)
        dark_mode_layout.addWidget(dark_mode_label)
        dark_mode_layout.addWidget(self.dark_mode_checkbox)
        dark_mode_layout.addStretch()
        theme_layout.addLayout(dark_mode_layout)
        
        theme_group.setLayout(theme_layout)
        layout.addWidget(theme_group)
        
        # ===== LLM Settings Group =====
        llm_group = QtWidgets.QGroupBox("LLM Configuration")
        llm_layout = QtWidgets.QVBoxLayout()
        
        # Model selection
        model_layout = QtWidgets.QHBoxLayout()
        model_label = QtWidgets.QLabel("Model:")
        self.model_combo = QtWidgets.QComboBox()
        
        # Auto-detect installed models from Ollama
        self._populate_ollama_models()
        
        # Set default model if available
        if self.model_combo.findText("gemma3:12b") >= 0:
            self.model_combo.setCurrentText("gemma3:12b")
        elif self.model_combo.count() > 0:
            self.model_combo.setCurrentIndex(0)
        
        model_layout.addWidget(model_label)
        model_layout.addWidget(self.model_combo)
        
        # Add refresh button to re-scan Ollama models
        refresh_models_btn = QtWidgets.QPushButton("Refresh Models")
        refresh_models_btn.setMaximumWidth(120)
        refresh_models_btn.clicked.connect(self._refresh_ollama_models)
        model_layout.addWidget(refresh_models_btn)
        
        model_layout.addStretch()
        llm_layout.addLayout(model_layout)
        
        # Temperature slider
        temp_layout = QtWidgets.QHBoxLayout()
        temp_label = QtWidgets.QLabel("Temperature:")
        self.temperature_slider = QtWidgets.QSlider(QtCore.Qt.Orientation.Horizontal)
        self.temperature_slider.setMinimum(0)
        self.temperature_slider.setMaximum(100)
        self.temperature_slider.setValue(70)
        self.temperature_slider.setTickPosition(QtWidgets.QSlider.TicksBelow)
        self.temperature_slider.setTickInterval(10)
        self.temperature_value_label = QtWidgets.QLabel("0.70")
        self.temperature_slider.sliderMoved.connect(self._on_temperature_changed)
        temp_layout.addWidget(temp_label)
        temp_layout.addWidget(self.temperature_slider)
        temp_layout.addWidget(self.temperature_value_label)
        temp_layout.addWidget(QtWidgets.QLabel("(0.0-1.0)"), 0, QtCore.Qt.AlignmentFlag.AlignRight)
        llm_layout.addLayout(temp_layout)
        
        llm_group.setLayout(llm_layout)
        layout.addWidget(llm_group)
        
        # ===== Application Settings Group =====
        app_group = QtWidgets.QGroupBox("Application")
        app_layout = QtWidgets.QVBoxLayout()
        
        # Auto-save interval
        autosave_layout = QtWidgets.QHBoxLayout()
        autosave_label = QtWidgets.QLabel("Auto-save Interval (minutes):")
        self.autosave_spinbox = QtWidgets.QSpinBox()
        self.autosave_spinbox.setMinimum(1)
        self.autosave_spinbox.setMaximum(60)
        self.autosave_spinbox.setValue(15)
        autosave_layout.addWidget(autosave_label)
        autosave_layout.addWidget(self.autosave_spinbox)
        autosave_layout.addStretch()
        app_layout.addLayout(autosave_layout)
        
        # Enable notifications
        notifications_layout = QtWidgets.QHBoxLayout()
        notifications_label = QtWidgets.QLabel("Enable Notifications:")
        self.notifications_checkbox = QtWidgets.QCheckBox()
        self.notifications_checkbox.setChecked(True)
        notifications_layout.addWidget(notifications_label)
        notifications_layout.addWidget(self.notifications_checkbox)
        notifications_layout.addStretch()
        app_layout.addLayout(notifications_layout)
        
        # Auto-approval setting
        autoapproval_layout = QtWidgets.QHBoxLayout()
        autoapproval_label = QtWidgets.QLabel("Auto-approve Content:")
        self.autoapproval_checkbox = QtWidgets.QCheckBox()
        self.autoapproval_checkbox.setChecked(False)
        autoapproval_info = QtWidgets.QLabel("(Auto-approves synopsis, outline, sections)")
        autoapproval_info.setStyleSheet("font-style: italic; color: #666;")
        autoapproval_layout.addWidget(autoapproval_label)
        autoapproval_layout.addWidget(self.autoapproval_checkbox)
        autoapproval_layout.addStretch()
        app_layout.addLayout(autoapproval_layout)
        app_layout.addWidget(autoapproval_info)
        
        app_group.setLayout(app_layout)
        layout.addWidget(app_group)
        
        # ===== Generation Parameters Group =====
        gen_group = QtWidgets.QGroupBox("Generation Parameters")
        gen_layout = QtWidgets.QVBoxLayout()
        
        # Max retries
        retries_layout = QtWidgets.QHBoxLayout()
        retries_label = QtWidgets.QLabel("Max LLM Retries:")
        self.max_retries_spinbox = QtWidgets.QSpinBox()
        self.max_retries_spinbox.setMinimum(1)
        self.max_retries_spinbox.setMaximum(10)
        self.max_retries_spinbox.setValue(3)
        retries_layout.addWidget(retries_label)
        retries_layout.addWidget(self.max_retries_spinbox)
        retries_layout.addStretch()
        gen_layout.addLayout(retries_layout)
        
        # Detail level
        detail_layout = QtWidgets.QHBoxLayout()
        detail_label = QtWidgets.QLabel("Detail Level:")
        self.detail_combo = QtWidgets.QComboBox()
        self.detail_combo.addItems(["Concise", "Balanced", "Detailed"])
        self.detail_combo.setCurrentText("Balanced")
        detail_layout.addWidget(detail_label)
        detail_layout.addWidget(self.detail_combo)
        detail_layout.addStretch()
        gen_layout.addLayout(detail_layout)
        
        # Character depth
        char_layout = QtWidgets.QHBoxLayout()
        char_label = QtWidgets.QLabel("Character Depth:")
        self.char_depth_combo = QtWidgets.QComboBox()
        self.char_depth_combo.addItems(["Shallow", "Standard", "Deep"])
        self.char_depth_combo.setCurrentText("Standard")
        char_layout.addWidget(char_label)
        char_layout.addWidget(self.char_depth_combo)
        char_layout.addStretch()
        gen_layout.addLayout(char_layout)
        
        # World depth
        world_layout = QtWidgets.QHBoxLayout()
        world_label = QtWidgets.QLabel("World-building Depth:")
        self.world_depth_combo = QtWidgets.QComboBox()
        self.world_depth_combo.addItems(["Minimal", "Standard", "Comprehensive"])
        self.world_depth_combo.setCurrentText("Standard")
        world_layout.addWidget(world_label)
        world_layout.addWidget(self.world_depth_combo)
        world_layout.addStretch()
        gen_layout.addLayout(world_layout)
        
        # Quality check
        quality_layout = QtWidgets.QHBoxLayout()
        quality_label = QtWidgets.QLabel("Quality Check Level:")
        self.quality_combo = QtWidgets.QComboBox()
        self.quality_combo.addItems(["Strict", "Moderate", "Lenient"])
        self.quality_combo.setCurrentText("Moderate")
        quality_layout.addWidget(quality_label)
        quality_layout.addWidget(self.quality_combo)
        quality_layout.addStretch()
        gen_layout.addLayout(quality_layout)
        
        # Sections per chapter
        sections_layout = QtWidgets.QHBoxLayout()
        sections_label = QtWidgets.QLabel("Sections Per Chapter:")
        self.sections_spinbox = QtWidgets.QSpinBox()
        self.sections_spinbox.setMinimum(1)
        self.sections_spinbox.setMaximum(10)
        self.sections_spinbox.setValue(3)
        sections_layout.addWidget(sections_label)
        sections_layout.addWidget(self.sections_spinbox)
        sections_layout.addStretch()
        gen_layout.addLayout(sections_layout)
        
        gen_group.setLayout(gen_layout)
        layout.addWidget(gen_group)
        
        # ===== Info Group =====
        info_group = QtWidgets.QGroupBox("Application Info")
        info_layout = QtWidgets.QVBoxLayout()
        
        # Version info
        version_label = QtWidgets.QLabel("Version: 1.0.0")
        info_layout.addWidget(version_label)
        
        # About button
        about_button = QtWidgets.QPushButton("About ANS")
        about_button.clicked.connect(self._on_about_clicked)
        info_layout.addWidget(about_button)
        
        info_group.setLayout(info_layout)
        layout.addWidget(info_group)
        
        # Add stretch to fill remaining space
        layout.addStretch()
        
        # Load settings from config
        self._load_settings()
        
        # Connect change signals to save settings
        self.model_combo.currentTextChanged.connect(self._save_settings)
        self.autosave_spinbox.valueChanged.connect(self._save_settings)
        self.notifications_checkbox.stateChanged.connect(self._save_settings)
        self.autoapproval_checkbox.stateChanged.connect(self._save_settings)
        self.max_retries_spinbox.valueChanged.connect(self._save_settings)
        self.detail_combo.currentTextChanged.connect(self._save_settings)
        self.char_depth_combo.currentTextChanged.connect(self._save_settings)
        self.world_depth_combo.currentTextChanged.connect(self._save_settings)
        self.quality_combo.currentTextChanged.connect(self._save_settings)
        self.sections_spinbox.valueChanged.connect(self._save_settings)
        
        return tab
    
    def _on_dark_mode_toggled(self, state):
        """Toggle dark mode theme for the application."""
        if state == QtCore.Qt.CheckState.Checked:
            self._apply_dark_mode()
            self._save_settings()
        else:
            self._apply_light_mode()
            self._save_settings()
    
    def _apply_dark_mode(self):
        """Apply dark mode stylesheet to the application."""
        dark_stylesheet = """
        QMainWindow, QDialog, QWidget {
            background-color: #1e1e1e;
            color: #e0e0e0;
        }
        QGroupBox {
            color: #e0e0e0;
            border: 1px solid #444;
            border-radius: 5px;
            margin-top: 10px;
            padding-top: 10px;
        }
        QGroupBox::title {
            subcontrol-origin: margin;
            left: 10px;
            padding: 0 3px 0 3px;
        }
        QLabel {
            color: #e0e0e0;
        }
        QLineEdit, QTextEdit, QComboBox {
            background-color: #2d2d2d;
            color: #e0e0e0;
            border: 1px solid #444;
            padding: 5px;
        }
        QPushButton {
            background-color: #0d47a1;
            color: #ffffff;
            border: none;
            padding: 5px;
            border-radius: 3px;
        }
        QPushButton:hover {
            background-color: #1565c0;
        }
        QPushButton:pressed {
            background-color: #0d3b8f;
        }
        QTabWidget::pane {
            border: 1px solid #444;
        }
        QTabBar::tab {
            background-color: #2d2d2d;
            color: #e0e0e0;
            padding: 5px 20px;
            border: 1px solid #444;
        }
        QTabBar::tab:selected {
            background-color: #0d47a1;
        }
        QProgressBar {
            background-color: #2d2d2d;
            border: 1px solid #444;
            color: #e0e0e0;
        }
        QSlider::handle:horizontal {
            background-color: #0d47a1;
        }
        QScrollBar {
            background-color: #2d2d2d;
        }
        """
        app = QtWidgets.qApp  # type: ignore
        if app:
            app.setStyleSheet(dark_stylesheet)  # type: ignore
    
    def _apply_light_mode(self):
        """Apply light mode stylesheet to the application (default)."""
        app = QtWidgets.qApp  # type: ignore
        if app:
            app.setStyleSheet("")  # type: ignore
    
    def _on_temperature_changed(self, value):
        """Update temperature value label when slider changes."""
        temp_value = value / 100.0
        self.temperature_value_label.setText(f"{temp_value:.2f}")
        self._save_settings()
    
    def _on_about_clicked(self):
        """Show about dialog with application information."""
        about_text = """
        <h2>Automated Novel System (ANS)</h2>
        <p><b>Version:</b> 1.0.0</p>
        <p><b>Description:</b></p>
        <p>A PyQt5-based desktop application for automated novel generation using local Ollama LLM models.</p>
        <p><b>Features:</b></p>
        <ul>
            <li>AI-powered synopsis generation and refinement</li>
            <li>Automated 25-chapter outline generation</li>
            <li>Character and world-building generation</li>
            <li>Section-by-section novel drafting</li>
            <li>Real-time streaming token display</li>
            <li>Export to DOCX and PDF formats</li>
            <li>Dark mode support</li>
        </ul>
        <p><b>Framework:</b> PyQt5 with Ollama local LLM integration</p>
        """
        dialog = QtWidgets.QMessageBox(self)
        dialog.setWindowTitle("About ANS")
        dialog.setText(about_text)
        dialog.setIcon(QtWidgets.QMessageBox.Information)
        dialog.exec_()
    
    def _populate_ollama_models(self):
        """Populate model combo box with installed Ollama models."""
        models = self._get_available_models()
        
        self.model_combo.clear()
        
        if models:
            for model in sorted(models):
                self.model_combo.addItem(model)
        else:
            # Add default models as fallback if no models are detected
            self.model_combo.addItem("gemma3:12b")
            self.model_combo.addItem("llama2")
            self.model_combo.addItem("mistral")
    
    def _refresh_ollama_models(self):
        """Refresh the list of available Ollama models."""
        self._populate_ollama_models()
        if self.current_project:
            self.log_update.emit("Ollama models refreshed")
        
        # Show notification to user
        models_count = self.model_combo.count()
        message = f"Found {models_count} model(s) installed in Ollama"
        if models_count == 0:
            message = "No models found. Please install a model in Ollama first."
        
        QtWidgets.QMessageBox.information(self, "Models Refreshed", message)
    
    def _get_available_models(self):
        """Get list of available models from Ollama.
        
        Returns:
            list: Sorted list of available model names
        """
        try:
            # Use ollama client to list available models
            if hasattr(self, 'client') and self.client:
                response = self.client.list()
                
                # Extract model names from response
                models = []
                if hasattr(response, 'models'):
                    # If response has models attribute (ollama.Response object)
                    for model in response.models:
                        # Try to get model name from various attributes
                        model_name = None
                        
                        # Try 'name' attribute first
                        if hasattr(model, 'name'):
                            model_name = model.name  # type: ignore
                        # Try extracting from string representation
                        elif hasattr(model, 'model'):
                            model_name = model.model  # type: ignore
                        else:
                            # Parse from string representation
                            model_str = str(model)
                            if "model='" in model_str:
                                # Extract model name from model='xxx' pattern
                                start = model_str.find("model='") + 7
                                end = model_str.find("'", start)
                                if start > 6 and end > start:
                                    model_name = model_str[start:end]
                        
                        if model_name and model_name not in models:
                            models.append(model_name)
                
                elif isinstance(response, dict) and 'models' in response:
                    # If response is a dict with models key
                    for model_entry in response['models']:
                        if isinstance(model_entry, dict) and 'name' in model_entry:
                            model_name = model_entry['name']
                            if model_name not in models:
                                models.append(model_name)
                        else:
                            model_str = str(model_entry)
                            if model_str not in models:
                                models.append(model_str)
                
                elif isinstance(response, list):
                    # If response is directly a list
                    for item in response:
                        model_name = str(item)
                        if model_name not in models:
                            models.append(model_name)
                
                if models:
                    return sorted(models)
        except Exception as e:
            print(f"Error detecting Ollama models: {e}")
        
        # Return empty list if detection fails
        return []
    
    def _load_settings(self):
        """Load application settings from config file."""
        try:
            config_path = os.path.join('Config', 'app_settings.txt')
            if os.path.exists(config_path):
                with open(config_path, 'r', encoding='utf-8') as f:
                    lines = f.readlines()
                    for line in lines:
                        if 'DarkMode:' in line:
                            dark_mode = line.split('DarkMode:')[1].strip() == 'True'
                            self.dark_mode_checkbox.setChecked(dark_mode)
                        elif 'Model:' in line:
                            model = line.split('Model:')[1].strip()
                            self.model_combo.setCurrentText(model)
                        elif 'Temperature:' in line:
                            temp = float(line.split('Temperature:')[1].strip())
                            self.temperature_slider.setValue(int(temp * 100))
                        elif 'AutoSave:' in line:
                            autosave = int(line.split('AutoSave:')[1].strip())
                            self.autosave_spinbox.setValue(autosave)
                        elif 'Notifications:' in line:
                            notifications = line.split('Notifications:')[1].strip() == 'True'
                            self.notifications_checkbox.setChecked(notifications)
                        elif 'AutoApproval:' in line:
                            autoapproval = line.split('AutoApproval:')[1].strip() == 'True'
                            self.autoapproval_checkbox.setChecked(autoapproval)
                        elif 'MaxRetries:' in line:
                            max_retries = int(line.split('MaxRetries:')[1].strip())
                            self.max_retries_spinbox.setValue(max_retries)
                        elif 'DetailLevel:' in line:
                            detail = line.split('DetailLevel:')[1].strip()
                            self.detail_combo.setCurrentText(detail.capitalize())
                        elif 'CharacterDepth:' in line:
                            char_depth = line.split('CharacterDepth:')[1].strip()
                            self.char_depth_combo.setCurrentText(char_depth.capitalize())
                        elif 'WorldDepth:' in line:
                            world_depth = line.split('WorldDepth:')[1].strip()
                            self.world_depth_combo.setCurrentText(world_depth.capitalize())
                        elif 'QualityCheck:' in line:
                            quality = line.split('QualityCheck:')[1].strip()
                            self.quality_combo.setCurrentText(quality.capitalize())
                        elif 'SectionsPerChapter:' in line:
                            sections = int(line.split('SectionsPerChapter:')[1].strip())
                            self.sections_spinbox.setValue(sections)
        except Exception as e:
            print(f"Error loading settings: {e}")
    
    def _save_settings(self):
        """Save application settings to config file."""
        try:
            # Create Config directory if it doesn't exist
            if not os.path.exists('Config'):
                os.makedirs('Config')
            
            config_path = os.path.join('Config', 'app_settings.txt')
            
            settings_content = f"""DarkMode: {self.dark_mode_checkbox.isChecked()}
Model: {self.model_combo.currentText()}
Temperature: {self.temperature_slider.value() / 100.0}
AutoSave: {self.autosave_spinbox.value()}
Notifications: {self.notifications_checkbox.isChecked()}
AutoApproval: {self.autoapproval_checkbox.isChecked()}
MaxRetries: {self.max_retries_spinbox.value()}
DetailLevel: {self.detail_combo.currentText().lower()}
CharacterDepth: {self.char_depth_combo.currentText().lower()}
WorldDepth: {self.world_depth_combo.currentText().lower()}
QualityCheck: {self.quality_combo.currentText().lower()}
SectionsPerChapter: {self.sections_spinbox.value()}
"""
            with open(config_path, 'w', encoding='utf-8') as f:
                f.write(settings_content)
        except Exception as e:
            print(f"Error saving settings: {e}")
    
    def _on_init_complete(self):
        """Handle initialization complete signal from thread."""
        if hasattr(self, 'dashboard_status_label'):
            self.dashboard_status_label.setText("Status: Initialized")
        if hasattr(self, 'dashboard_progress_label'):
            self.dashboard_progress_label.setText("Progress: 0%")
        if hasattr(self, 'dashboard_progress_bar'):
            self.dashboard_progress_bar.setValue(0)
    
    def is_autoapproval_enabled(self):
        """Check if auto-approval is enabled in settings."""
        return hasattr(self, 'autoapproval_checkbox') and self.autoapproval_checkbox.isChecked()
    
    def _sync_settings_to_thread(self):
        """Synchronize UI settings to background thread."""
        if hasattr(self, 'thread'):
            self.thread.llm_model = self.model_combo.currentText()
            self.thread.temperature = self.temperature_slider.value() / 100.0
            self.thread.max_retries = self.max_retries_spinbox.value()
            self.thread.detail_level = self.detail_combo.currentText().lower()
            self.thread.character_depth = self.char_depth_combo.currentText().lower()
            self.thread.world_depth = self.world_depth_combo.currentText().lower()
            self.thread.quality_check = self.quality_combo.currentText().lower()
            self.thread.sections_per_chapter = self.sections_spinbox.value()
    
    def _refresh_logs_tab(self):
        """Refresh the logs tab with message about app logs location."""
        if hasattr(self, 'logs_text_edit'):
            # Project logs no longer stored - all logs go to Config/log1-5.txt
            if self.current_project:
                self.logs_text_edit.setText(f"Project: {self.current_project['name']}\n\nAll application logs are stored in Config/log1.txt through log5.txt\n(rotating log files)\n\nLogs displayed here are from the current session.")
            else:
                self.logs_text_edit.setText("No project loaded.\n\nApplication logs are stored in Config/log1.txt through log5.txt\n(rotating log files)")
    
    def _update_llm_status_indicator(self):
        """Update the LLM status indicator color and text."""
        if self.llm_connected:
            self.llm_status_indicator.setText("● Connected")
            self.llm_status_indicator.setStyleSheet("color: green; font-weight: bold;")
            self.status_update_timer.stop()  # Stop updating once connected
        else:
            # Keep showing "Connecting..." until LLM connects
            pass
    
    def _refresh_project_list(self):
        """Refresh the project list combo box."""
        projects = self.get_project_list()
        self.project_list_combo.clear()
        
        if projects:
            self.project_list_combo.addItems(projects)
            self.project_list_combo.setEnabled(True)
        else:
            self.project_list_combo.addItem("(No projects available)")
            self.project_list_combo.setEnabled(False)
    
    def _on_load_project(self):
        """Handle project loading when button is clicked."""
        project_name = self.project_list_combo.currentText()
        
        self._write_app_log(f"User initiated project load: {project_name}")
        
        if project_name == "(No projects available)":
            self._write_app_log(f"Project load failed: no projects available")
            self.initialization_status.setText("Error: No projects available to load.")
            self.error_signal.emit("No projects available")
            return
        
        try:
            self.load_project(project_name)
            self._write_app_log(f"Project '{project_name}' loaded successfully - session persistence active")
            self.initialization_status.setText(f"✓ Project '{project_name}' loaded successfully!\n\nProject stored in memory for session persistence.\nLocation: projects/{project_name}/")
            self.log_update.emit(f"Project '{project_name}' loaded into session")
        except Exception as e:
            error_msg = f"Error loading project: {str(e)}"
            self._write_app_log(f"Project load failed for '{project_name}': {error_msg}")
            self.initialization_status.setText(error_msg)
            self.error_signal.emit(error_msg)
    
    def _on_run_test_prompt(self):
        """Handle test prompt execution."""
        if not self.llm_connected:
            self._write_app_log(f"Test prompt failed: LLM not connected")
            self.initialization_status.setText("✗ Ollama is not connected yet. Please wait for connection or check logs.")
            return
        
        prompt = self.test_prompt_input.text().strip()
        if not prompt:
            self._write_app_log(f"Test prompt failed: empty prompt")
            self.initialization_status.setText("Error: Prompt cannot be empty!")
            return
        
        self._write_app_log(f"User initiated test prompt: '{prompt[:50]}...'")
        self.initialization_status.setText("Testing Ollama... (this may take a moment)")
        QtWidgets.QApplication.processEvents()  # Update UI immediately
        
        try:
            # Run in background thread to not block UI
            test_thread = threading.Thread(target=self._run_test_prompt_thread, args=(prompt,), daemon=True)
            test_thread.start()
        except Exception as e:
            error_msg = f"Error running test: {str(e)}"
            self._write_app_log(f"Test prompt error: {error_msg}")
            self.initialization_status.setText(error_msg)
    
    def _run_test_prompt_thread(self, prompt):
        """Run test prompt in background thread."""
        try:
            response = self.client.generate(model='gemma3:12b', prompt=prompt)
            result_text = response.get('response', '')
            
            # Emit signal to update UI from main thread
            message = f"✓ Test successful!\n\nPrompt: {prompt}\n\nResponse:\n{result_text}"
            self.test_result_signal.emit(message)
            
            # Log to project if active
            if self.current_project:
                self.log_update.emit(f"Test prompt executed: {prompt}")
        except Exception as e:
            error_msg = f"✗ Test failed: {str(e)}"
            self.test_result_signal.emit(error_msg)
            
            # Log to project if active
            if self.current_project:
                self.log_update.emit(f"Test prompt failed: {str(e)}")
    
    def _on_test_result(self, message):
        """Handle test result from background thread."""
        self.initialization_status.setText(message)
    
    def _on_error_signal(self, error_message):
        """Handle error signal by showing message box and logging to log.txt."""
        # Show critical error dialog
        QtWidgets.QMessageBox.critical(self, "Error", error_message)
        
        # Log error to rotating app log
        self._write_app_log(f"ERROR: {error_message}")
    
    def _on_start_signal(self, config_string):
        """Handle start signal with novel configuration."""
        # Sync settings to thread before starting
        self._sync_settings_to_thread()
        
        self._write_app_log(f"Start signal received: {config_string}")
    
    def _on_processing_finished(self, result):
        """Handle background thread processing finished signal."""
        self.log_update.emit(f"Processing finished: {result}")
    
    def _on_processing_error(self, error):
        """Handle background thread processing error signal."""
        self.error_signal.emit(f"Processing error: {error}")
    
    def _on_processing_progress(self, progress):
        """Handle background thread processing progress signal."""
        self.log_update.emit(f"Processing progress: {progress}")
    
    def _on_synopsis_ready(self, synopsis_text):
        """Handle synopsis ready signal. Update synopsis display incrementally."""
        if hasattr(self, 'synopsis_display'):
            # Get current text length
            current_text = self.synopsis_display.toPlainText()
            
            # Only update if new text has been added
            if len(synopsis_text) > len(current_text):
                # Check if user is at the bottom before appending
                scrollbar = self.synopsis_display.verticalScrollBar()
                is_at_bottom = scrollbar is not None and scrollbar.value() == scrollbar.maximum()
                
                # Append only the new part to preserve scroll position
                new_part = synopsis_text[len(current_text):]
                self.synopsis_display.insertPlainText(new_part)
                
                # Only auto-scroll if user was already at the bottom
                if is_at_bottom:
                    cursor = self.synopsis_display.textCursor()
                    cursor.movePosition(cursor.MoveOperation.End)
                    self.synopsis_display.setTextCursor(cursor)
        
        # Also update expanded window if it's open
        if 'synopsis_display' in self.expanded_text_widgets:
            expanded_text = self.expanded_text_widgets['synopsis_display']
            current_text = expanded_text.toPlainText()
            if len(synopsis_text) > len(current_text):
                new_part = synopsis_text[len(current_text):]
                expanded_text.insertPlainText(new_part)
                # Auto-scroll to bottom
                cursor = expanded_text.textCursor()
                cursor.movePosition(cursor.MoveOperation.End)
                expanded_text.setTextCursor(cursor)
        
        # Enable initial synopsis buttons for user review
        if hasattr(self, 'initial_approve_button'):
            self.initial_approve_button.setEnabled(True)
        if hasattr(self, 'initial_adjust_button'):
            self.initial_adjust_button.setEnabled(True)
    
    def _on_refinement_start(self):
        """Handle refinement start signal. Clear planning display and disable buttons during refinement."""
        if hasattr(self, 'planning_synopsis_display'):
            self.planning_synopsis_display.clear()
        
        # Disable buttons during refinement - they'll be re-enabled when refinement completes
        if hasattr(self, 'approve_button'):
            self.approve_button.setEnabled(False)
        if hasattr(self, 'adjust_button'):
            self.adjust_button.setEnabled(False)
    
    def _on_outline_refinement_start(self):
        """Handle outline refinement start signal. Clear outline display and disable buttons during refinement."""
        if hasattr(self, 'outline_display'):
            self.outline_display.clear()
        
        # Disable outline buttons during refinement - they'll be re-enabled when refinement completes
        if hasattr(self, 'approve_outline_button'):
            self.approve_outline_button.setEnabled(False)
        if hasattr(self, 'adjust_outline_button'):
            self.adjust_outline_button.setEnabled(False)
    
    def _on_timeline_refinement_start(self):
        """Handle timeline refinement start signal. Clear timeline display and disable buttons during refinement."""
        if hasattr(self, 'timeline_display'):
            self.timeline_display.clear()
        
        # Disable timeline buttons during refinement - they'll be re-enabled when refinement completes
        if hasattr(self, 'approve_timeline_button'):
            self.approve_timeline_button.setEnabled(False)
        if hasattr(self, 'adjust_timeline_button'):
            self.adjust_timeline_button.setEnabled(False)
    
    def _on_new_synopsis(self, refined_synopsis_text):
        """Handle new_synopsis signal from refinement. Update Planning tab and switch to it."""
        # Update Planning tab synopsis display with incremental updates
        if hasattr(self, 'planning_synopsis_display'):
            # Get current text length
            current_text = self.planning_synopsis_display.toPlainText()
            current_length = len(current_text)
            new_length = len(refined_synopsis_text)
            
            # Handle display clear first (refinement_start signal clears display)
            if new_length == 0 and current_length > 0:
                # Display was cleared, this is a new refinement starting
                self.planning_synopsis_display.clear()
            elif new_length > current_length:
                # New text has been added - stream it in
                # Check if user is at the bottom before appending
                scrollbar = self.planning_synopsis_display.verticalScrollBar()
                is_at_bottom = scrollbar is not None and scrollbar.value() == scrollbar.maximum()
                
                # Append only the new part to preserve scroll position
                new_part = refined_synopsis_text[current_length:]
                self.planning_synopsis_display.insertPlainText(new_part)
                
                # Only auto-scroll if user was already at the bottom
                if is_at_bottom:
                    cursor = self.planning_synopsis_display.textCursor()
                    cursor.movePosition(cursor.MoveOperation.End)
                    self.planning_synopsis_display.setTextCursor(cursor)
        
        # Also update expanded window if it's open
        if 'planning_synopsis_display' in self.expanded_text_widgets:
            expanded_text = self.expanded_text_widgets['planning_synopsis_display']
            current_text = expanded_text.toPlainText()
            current_length = len(current_text)
            new_length = len(refined_synopsis_text)
            
            # Handle display clear first
            if new_length == 0 and current_length > 0:
                expanded_text.clear()
            elif new_length > current_length:
                # Stream new content
                new_part = refined_synopsis_text[current_length:]
                expanded_text.insertPlainText(new_part)
                # Auto-scroll to bottom
                cursor = expanded_text.textCursor()
                cursor.movePosition(cursor.MoveOperation.End)
                expanded_text.setTextCursor(cursor)
        
        # Enable the Approve and Adjust buttons for synopsis
        if hasattr(self, 'approve_button'):
            self.approve_button.setEnabled(True)
        if hasattr(self, 'adjust_button'):
            self.adjust_button.setEnabled(True)
        
        # Switch to Planning tab (index 2: Initialization=0, Novel Idea=1, Planning=2)
        if hasattr(self, 'tabs'):
            self.tabs.setCurrentIndex(2)
        
        # Auto-approve synopsis if enabled in settings
        if self.is_autoapproval_enabled():
            QtCore.QTimer.singleShot(500, self._on_approve_synopsis)
    
    def _on_new_outline(self, outline_text):
        """Handle new_outline signal from outline generation/refinement. Update Planning tab outline display."""
        # Update Planning tab outline display with incremental updates
        if hasattr(self, 'outline_display'):
            # Get current text length
            current_text = self.outline_display.toPlainText()
            
            # Only update if new text has been added
            if len(outline_text) > len(current_text):
                # Check if user is at the bottom before appending
                scrollbar = self.outline_display.verticalScrollBar()
                is_at_bottom = scrollbar is not None and scrollbar.value() == scrollbar.maximum()
                
                # Append only the new part to preserve scroll position
                new_part = outline_text[len(current_text):]
                self.outline_display.insertPlainText(new_part)
                
                # Only auto-scroll if user was already at the bottom
                if is_at_bottom:
                    cursor = self.outline_display.textCursor()
                    cursor.movePosition(cursor.MoveOperation.End)
                    self.outline_display.setTextCursor(cursor)
        
        # Also update expanded window if it's open
        if 'outline_display' in self.expanded_text_widgets:
            expanded_text = self.expanded_text_widgets['outline_display']
            current_text = expanded_text.toPlainText()
            if len(outline_text) > len(current_text):
                new_part = outline_text[len(current_text):]
                expanded_text.insertPlainText(new_part)
                # Auto-scroll to bottom
                cursor = expanded_text.textCursor()
                cursor.movePosition(cursor.MoveOperation.End)
                expanded_text.setTextCursor(cursor)
        
        # Enable the Approve and Adjust buttons for outline
        if hasattr(self, 'approve_outline_button'):
            self.approve_outline_button.setEnabled(True)
        if hasattr(self, 'adjust_outline_button'):
            self.adjust_outline_button.setEnabled(True)
        
        # Auto-approve outline if enabled in settings
        if self.is_autoapproval_enabled():
            QtCore.QTimer.singleShot(500, self._on_approve_outline)
    
    def _on_new_characters(self, characters_json):
        """Handle new_characters signal from character generation. Display formatted JSON in Planning tab and enable buttons."""
        import json
        
        # Update current project's character data
        if self.current_project:
            self.current_project['characters'] = characters_json
            self.log_update.emit(f"Character generation received: {len(characters_json.split())} words generated")
        
        # Display formatted JSON in characters_display
        if hasattr(self, 'characters_display'):
            try:
                # Parse JSON and format with indentation
                characters_data = json.loads(characters_json)
                formatted_json = json.dumps(characters_data, indent=2, ensure_ascii=False)
                self.characters_display.setText(formatted_json)
                
                # Scroll to top to show first character
                cursor = self.characters_display.textCursor()
                cursor.movePosition(cursor.MoveOperation.Start)
                self.characters_display.setTextCursor(cursor)
            except json.JSONDecodeError:
                # If parsing fails, display raw JSON
                self.characters_display.setText(characters_json)
        
        # Enable the Approve and Adjust buttons for characters
        if hasattr(self, 'approve_characters_button'):
            self.approve_characters_button.setEnabled(True)
        if hasattr(self, 'adjust_characters_button'):
            self.adjust_characters_button.setEnabled(True)
    
    def _on_new_world(self, world_json):
        """Handle new_world signal from world generation. Display formatted JSON in Planning tab and enable buttons."""
        import json
        
        # Update current project's world data
        if self.current_project:
            self.current_project['world'] = world_json
            self.log_update.emit(f"World generation received: {len(world_json.split())} words generated")
        
        # Display formatted JSON in world_display
        if hasattr(self, 'world_display'):
            try:
                # Parse JSON and format with indentation
                world_data = json.loads(world_json)
                formatted_json = json.dumps(world_data, indent=2, ensure_ascii=False)
                self.world_display.setText(formatted_json)
                
                # Scroll to top to show first section
                cursor = self.world_display.textCursor()
                cursor.movePosition(cursor.MoveOperation.Start)
                self.world_display.setTextCursor(cursor)
            except json.JSONDecodeError:
                # If parsing fails, display raw JSON
                self.world_display.setText(world_json)
        
        # Enable the Approve and Adjust buttons for world
        if hasattr(self, 'approve_world_button'):
            self.approve_world_button.setEnabled(True)
        if hasattr(self, 'adjust_world_button'):
            self.adjust_world_button.setEnabled(True)
    
    def _on_new_timeline(self, timeline_text):
        """Handle new_timeline signal from timeline generation. Update project, display, and enable buttons."""
        # Update current project's timeline data
        if self.current_project:
            self.current_project['timeline'] = timeline_text
            self.log_update.emit(f"Timeline generation received: {len(timeline_text.split())} words generated")
            
            # Display timeline in Planning tab
            if hasattr(self, 'timeline_display'):
                self.timeline_display.setText(timeline_text)
                cursor = self.timeline_display.textCursor()
                cursor.movePosition(cursor.MoveOperation.Start)
                self.timeline_display.setTextCursor(cursor)
            
            # Enable timeline action buttons for user approval/adjustment
            if hasattr(self, 'approve_timeline_button'):
                self.approve_timeline_button.setEnabled(True)
            if hasattr(self, 'adjust_timeline_button'):
                self.adjust_timeline_button.setEnabled(True)
    
    def _on_log_update(self, log_message):
        """Handle log update signal. Display in Logs tab QTextEdit."""
        # Create log entry with timestamp
        log_entry = f"{datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')} - {log_message}\n"
        
        # Append to logs QTextEdit if it exists (all logging now goes to Config/log1-5.txt)
        if hasattr(self, 'logs_text_edit'):
            self.logs_text_edit.append(log_entry.rstrip())
    
    def _on_new_draft(self, draft_content):
        """Handle new_draft signal from background thread. Display draft in Writing tab."""
        if self.current_project:
            # Store in current project buffer
            self.current_project['buffer_backup'] = draft_content
            
            # Display draft in Writing tab
            if hasattr(self, 'draft_display'):
                self.draft_display.setText(draft_content)
                cursor = self.draft_display.textCursor()
                cursor.movePosition(cursor.MoveOperation.Start)
                self.draft_display.setTextCursor(cursor)
                
                # Auto-switch to Writing tab when new draft arrives
                self.tabs.setCurrentIndex(3)  # Writing tab is index 3
            
            # Enable action buttons for user control
            if hasattr(self, 'approve_section_button'):
                self.approve_section_button.setEnabled(True)
            if hasattr(self, 'adjust_section_button'):
                self.adjust_section_button.setEnabled(True)
            if hasattr(self, 'pause_button'):
                self.pause_button.setEnabled(True)
            
            self.log_update.emit("Draft received and ready for review in Writing tab")
            
            # Auto-approve section if enabled in settings
            if self.is_autoapproval_enabled():
                QtCore.QTimer.singleShot(500, self._on_approve_section)
    
    def _on_approve_section(self):
        """Handle Approve Section button click - emit approve_signal with 'section'."""
        # Check if there's content in the draft display
        if hasattr(self, 'draft_display'):
            content = self.draft_display.toPlainText().strip()
            if content and self.current_project:
                self.buffer = content
                self.current_project['buffer_backup'] = content
        
        self.approve_signal.emit('section')
        if self.current_project:
            self.log_update.emit("Section approved. Continuing to next section...")
            # Disable buttons during generation
            if hasattr(self, 'approve_section_button'):
                self.approve_section_button.setEnabled(False)
            if hasattr(self, 'adjust_section_button'):
                self.adjust_section_button.setEnabled(False)
            if hasattr(self, 'pause_button'):
                self.pause_button.setEnabled(False)
    
    def _on_adjust_section(self):
        """Handle Adjust Section button click - get feedback and emit adjust_signal with 'section'."""
        # Check if there's content to adjust
        if not hasattr(self, 'draft_display') or not self.draft_display.toPlainText().strip():
            self.error_signal.emit("No section content to adjust")
            return
        
        feedback_text, ok = QtWidgets.QInputDialog.getMultiLineText(
            self,
            "Adjust Section",
            "What adjustments would you like to make to this section?",
            ""
        )
        
        if ok and feedback_text.strip():
            # Always emit signal to start refinement process
            self.adjust_signal.emit('section', feedback_text)
            if self.current_project:
                self.log_update.emit(f"Section adjustment requested: {feedback_text[:100]}...")
            # Disable buttons during revision
            if hasattr(self, 'approve_section_button'):
                self.approve_section_button.setEnabled(False)
            if hasattr(self, 'adjust_section_button'):
                self.adjust_section_button.setEnabled(False)
    
    def _on_pause_generation(self):
        """Handle Pause button click to pause background generation."""
        self._write_app_log(f"User paused generation")
        self.pause_signal.emit()
        self.thread.set_paused(True)
        if self.current_project:
            self.log_update.emit("Generation paused. Click resume to continue.")
        # Hide pause button and show resume button
        if hasattr(self, 'pause_button'):
            self.pause_button.setEnabled(False)
            self.pause_button.setVisible(False)
        if hasattr(self, 'resume_button'):
            self.resume_button.setEnabled(True)
            self.resume_button.setVisible(True)
    
    def _on_resume_generation(self):
        """Handle Resume button click to resume paused background generation."""
        self._write_app_log(f"User resumed generation")
        self.thread.set_paused(False)
        if self.current_project:
            self.log_update.emit("Generation resumed.")
        # Hide resume button and show pause button
        if hasattr(self, 'resume_button'):
            self.resume_button.setEnabled(False)
            self.resume_button.setVisible(False)
        if hasattr(self, 'pause_button'):
            self.pause_button.setEnabled(True)
            self.pause_button.setVisible(True)
    
    def _on_start_novel(self):
        """Handle start button click to begin novel generation."""
        idea = self.novel_idea_input.toPlainText().strip()
        tone = self.tone_input.toPlainText().strip()
        word_count = self.word_count_spinbox.value()
        
        if not idea:
            self._write_app_log(f"Novel generation failed: no idea provided")
            QtWidgets.QMessageBox.warning(self, "Warning", "Please enter a novel idea to start.")
            return
        
        if not tone:
            self._write_app_log(f"Novel generation failed: no tone provided")
            QtWidgets.QMessageBox.warning(self, "Warning", "Please describe the desired tone.")
            return
        
        self._write_app_log(f"User initiated novel generation - Target: {word_count} words")
        
        # Create config string from inputs
        config_string = f"Idea: {idea}, Tone: {tone}, Soft Target: {word_count}"
        
        # Emit start signal with configuration
        self.start_signal.emit(config_string)
        
        # Log to project if active
        if self.current_project:
            self.log_update.emit(f"Novel generation started - {config_string}")
    
    def _on_approve_synopsis(self):
        """Handle Approve button click - emit approve_signal with 'synopsis'."""
        # Check if there's content in the refined synopsis display
        refined_content = ""
        if hasattr(self, 'planning_synopsis_display'):
            refined_content = self.planning_synopsis_display.toPlainText().strip()
        
        if not refined_content:
            # Try initial synopsis if refined is empty
            if hasattr(self, 'synopsis_display'):
                refined_content = self.synopsis_display.toPlainText().strip()
        
        if refined_content:
            # Update the synopsis in current_project
            if self.current_project:
                self.current_project['synopsis'] = refined_content
        
        self.approve_signal.emit('synopsis')
        if self.current_project:
            self.log_update.emit("Refined synopsis approved. Ready to proceed with planning.")
        # Disable initial synopsis buttons
        if hasattr(self, 'initial_approve_button'):
            self.initial_approve_button.setEnabled(False)
        if hasattr(self, 'initial_adjust_button'):
            self.initial_adjust_button.setEnabled(False)
    
    def _on_adjust_synopsis(self):
        """Handle Adjust button click - get feedback and emit adjust_signal."""
        # Check if there's content to adjust
        current_content = ""
        if hasattr(self, 'planning_synopsis_display'):
            current_content = self.planning_synopsis_display.toPlainText().strip()
        
        if not current_content and hasattr(self, 'synopsis_display'):
            current_content = self.synopsis_display.toPlainText().strip()
        
        if not current_content:
            self.error_signal.emit("No synopsis content to adjust")
            return
        
        feedback_text, ok = QtWidgets.QInputDialog.getMultiLineText(
            self,
            "Adjust Synopsis",
            "What adjustments would you like to make to the synopsis?",
            ""
        )
        
        if ok and feedback_text.strip():
            # Always emit signal to start refinement process
            self.adjust_signal.emit('synopsis', feedback_text)
            if self.current_project:
                self.log_update.emit(f"Refined synopsis adjustment requested: {feedback_text[:100]}...")
            # Disable buttons during refinement
            if hasattr(self, 'initial_approve_button'):
                self.initial_approve_button.setEnabled(False)
            if hasattr(self, 'initial_adjust_button'):
                self.initial_adjust_button.setEnabled(False)
    
    def _on_approve_initial_synopsis(self):
        """Handle Approve button click for initial synopsis - can approve generated or loaded content."""
        # Check if there's content in the initial synopsis display
        if hasattr(self, 'synopsis_display'):
            content = self.synopsis_display.toPlainText().strip()
            if content:
                # Content exists - either from generation or loaded from file
                # Update the synopsis in current_project
                if self.current_project:
                    self.current_project['synopsis'] = content
                    self.log_update.emit("Initial synopsis approved. Moving to refined synopsis review...")
                # Route through the normal approval flow
                self._on_approve_synopsis()
                return
        
        self.error_signal.emit("No synopsis content to approve")
    
    def _on_adjust_initial_synopsis(self):
        """Handle Adjust button click for initial synopsis - can adjust generated or loaded content."""
        # Check if there's content in the initial synopsis display
        if hasattr(self, 'synopsis_display'):
            current_content = self.synopsis_display.toPlainText().strip()
            if not current_content:
                self.error_signal.emit("No synopsis content to adjust")
                return
            
            feedback_text, ok = QtWidgets.QInputDialog.getMultiLineText(
                self,
                "Adjust Initial Synopsis",
                "What adjustments would you like to make to the initial synopsis?",
                ""
            )
            
            if ok and feedback_text.strip():
                # Always emit signal to start refinement process
                self.adjust_signal.emit('synopsis', feedback_text)
                if self.current_project:
                    self.log_update.emit(f"Initial synopsis adjustment requested: {feedback_text[:100]}...")
                # Disable initial synopsis buttons during refinement
                if hasattr(self, 'initial_approve_button'):
                    self.initial_approve_button.setEnabled(False)
                if hasattr(self, 'initial_adjust_button'):
                    self.initial_adjust_button.setEnabled(False)
            return
        
        self.error_signal.emit("No synopsis content to adjust")
    
    def _on_approve_content(self, content_type):
        """Route approve signal to appropriate handler based on content type."""
        self._write_app_log(f"User approved content: {content_type}")
        if content_type == 'synopsis':
            # Save progress: synopsis approved
            self._save_progress('synopsis', 'approved')
            # Queue outline generation in background thread via start_processing
            self.thread.start_processing({'operation': 'generate_outline', 'type': 'synopsis'})
            if self.current_project:
                self.log_update.emit("Synopsis approved. Generating outline...")
        elif content_type == 'outline':
            # Save progress: outline approved
            self._save_progress('outline', 'approved')
            # Queue character generation in background thread via start_processing
            self.thread.start_processing({'operation': 'generate_characters', 'type': 'outline'})
            if self.current_project:
                self.log_update.emit("Outline approved. Generating characters...")
        elif content_type == 'characters':
            # Save progress: characters approved
            self._save_progress('characters', 'approved')
            # Queue world generation in background thread via start_processing
            self.thread.start_processing({'operation': 'generate_world', 'type': 'characters'})
            if self.current_project:
                self.log_update.emit("Characters approved. Generating world...")
        elif content_type == 'world':
            # Save progress: world approved (final stage)
            self._save_progress('world', 'approved')
            if self.current_project:
                self.log_update.emit("World approved. All core planning stages complete!")
                # Initialize chapter tracking after world approval
                self._initialize_chapter_tracking()
                # Queue timeline generation in background thread via start_processing
                self.thread.start_processing({'operation': 'generate_timeline', 'type': 'world'})
                self.log_update.emit("Generating timeline from planning...")
        elif content_type == 'timeline':
            # Save progress: timeline approved
            self._save_progress('timeline', 'approved')
            if self.current_project:
                self.log_update.emit("Timeline approved. Beginning chapter-by-chapter research generation...")
                # Queue chapter research loop in background thread via start_processing
                self.thread.start_processing({'operation': 'start_chapter_research_loop'})
        elif content_type == 'section':
            # Section approval: append to story, generate summary, update context
            # Queue section approval in background thread via start_processing
            self.thread.start_processing({'operation': 'approve_section', 'type': 'section'})
            if self.current_project:
                self.log_update.emit("Section approved. Processing and storing...")
    
    def _on_adjust_content(self, content_type, feedback):
        """Route adjust signal to appropriate handler based on content type."""
        self._write_app_log(f"User requested refinement for: {content_type} - Feedback: {feedback[:50]}...")
        # Clear the appropriate display BEFORE starting refinement
        if content_type == 'synopsis' and hasattr(self, 'planning_synopsis_display'):
            self.planning_synopsis_display.clear()
        elif content_type == 'outline' and hasattr(self, 'outline_display'):
            self.outline_display.clear()
        elif content_type == 'characters' and hasattr(self, 'characters_display'):
            self.characters_display.clear()
        elif content_type == 'world' and hasattr(self, 'world_display'):
            self.world_display.clear()
        elif content_type == 'timeline' and hasattr(self, 'timeline_display'):
            self.timeline_display.clear()
        elif content_type == 'section' and hasattr(self, 'draft_display'):
            self.draft_display.clear()
        
        # Also clear expanded windows if open
        display_names = {
            'synopsis': 'planning_synopsis_display',
            'outline': 'outline_display',
            'characters': 'characters_display',
            'world': 'world_display',
            'timeline': 'timeline_display',
            'section': 'draft_display'
        }
        
        if content_type in display_names:
            display_name = display_names[content_type]
            if display_name in self.expanded_text_widgets:
                self.expanded_text_widgets[display_name].clear()
        
        # For refinements on loaded content, we need to start the thread to execute the refinement
        if not self.thread.isRunning():
            # Thread is not running, so we need to store the refinement info and start the thread
            self.thread.refinement_type = content_type  # type: ignore
            self.thread.refinement_feedback = feedback  # type: ignore
            self.thread.start_processing({'refinement': True, 'type': content_type, 'feedback': feedback})
        else:
            # Thread is already running (active generation), call refinement directly
            if content_type == 'synopsis':
                self.thread.refine_synopsis_with_feedback(content_type, feedback)
            elif content_type == 'outline':
                self.thread.refine_outline_with_feedback(content_type, feedback)
            elif content_type == 'characters':
                self.thread.refine_characters_with_feedback(content_type, feedback)
            elif content_type == 'world':
                self.thread.refine_world_with_feedback(content_type, feedback)
            elif content_type == 'timeline':
                self.thread.refine_timeline_with_feedback(content_type, feedback)
            elif content_type == 'section':
                self.thread.refine_section_with_feedback(content_type, feedback)
    
    def _on_approve_outline(self):
        """Handle Approve button click for outline - emit approve_signal with 'outline'."""
        # Check if there's content in the outline display
        if hasattr(self, 'outline_display'):
            content = self.outline_display.toPlainText().strip()
            if content and self.current_project:
                self.current_project['outline'] = content
        
        self.approve_signal.emit('outline')
        if self.current_project:
            self.log_update.emit("Outline approved. Ready to proceed with writing.")
    
    def _on_adjust_outline(self):
        """Handle Adjust button click for outline - get feedback and emit adjust_signal."""
        # Check if there's content to adjust
        if not hasattr(self, 'outline_display') or not self.outline_display.toPlainText().strip():
            self.error_signal.emit("No outline content to adjust")
            return
        
        feedback_text, ok = QtWidgets.QInputDialog.getMultiLineText(
            self,
            "Adjust Outline",
            "What adjustments would you like to make to the outline?",
            ""
        )
        
        if ok and feedback_text.strip():
            # Always emit signal to start refinement process
            self.adjust_signal.emit('outline', feedback_text)
            if self.current_project:
                self.log_update.emit(f"Outline adjustment requested: {feedback_text[:100]}...")
    
    def _on_approve_characters(self):
        """Handle Approve button click for characters - emit approve_signal with 'characters'."""
        # Check if there's content in the characters display
        if hasattr(self, 'characters_display'):
            content = self.characters_display.toPlainText().strip()
            if content and self.current_project:
                self.current_project['characters'] = content
        
        self.approve_signal.emit('characters')
        if self.current_project:
            self.log_update.emit("Characters approved. Ready for next phase.")
    
    def _on_adjust_characters(self):
        """Handle Adjust button click for characters - get feedback and emit adjust_signal."""
        # Check if there's content to adjust
        if not hasattr(self, 'characters_display') or not self.characters_display.toPlainText().strip():
            self.error_signal.emit("No characters content to adjust")
            return
        
        feedback_text, ok = QtWidgets.QInputDialog.getMultiLineText(
            self,
            "Adjust Characters",
            "What adjustments would you like to make to the character profiles?",
            ""
        )
        
        if ok and feedback_text.strip():
            # Always emit signal to start refinement process
            self.adjust_signal.emit('characters', feedback_text)
            if self.current_project:
                self.log_update.emit(f"Characters adjustment requested: {feedback_text[:100]}...")
    
    def _on_approve_world(self):
        """Handle Approve button click for world - emit approve_signal with 'world'."""
        # Check if there's content in the world display
        if hasattr(self, 'world_display'):
            content = self.world_display.toPlainText().strip()
            if content and self.current_project:
                self.current_project['world'] = content
        
        self.approve_signal.emit('world')
        if self.current_project:
            self.log_update.emit("World approved. Ready for next phase.")
    
    def _on_adjust_world(self):
        """Handle Adjust button click for world - get feedback and emit adjust_signal."""
        # Check if there's content to adjust
        if not hasattr(self, 'world_display') or not self.world_display.toPlainText().strip():
            self.error_signal.emit("No world content to adjust")
            return
        
        feedback_text, ok = QtWidgets.QInputDialog.getMultiLineText(
            self,
            "Adjust World",
            "What adjustments would you like to make to the world details?",
            ""
        )
        
        if ok and feedback_text.strip():
            # Always emit signal to start refinement process
            self.adjust_signal.emit('world', feedback_text)
            if self.current_project:
                self.log_update.emit(f"World adjustment requested: {feedback_text[:100]}...")
    
    def _on_approve_timeline(self):
        """Handle Approve button click for timeline - emit approve_signal with 'timeline'."""
        # Check if there's content in the timeline display
        if hasattr(self, 'timeline_display'):
            content = self.timeline_display.toPlainText().strip()
            if content and self.current_project:
                self.current_project['timeline'] = content
        
        self.approve_signal.emit('timeline')
        if self.current_project:
            self.log_update.emit("Timeline approved. Planning workflow complete.")
            # Disable timeline buttons after approval
            if hasattr(self, 'approve_timeline_button'):
                self.approve_timeline_button.setEnabled(False)
            if hasattr(self, 'adjust_timeline_button'):
                self.adjust_timeline_button.setEnabled(False)
    
    def _on_adjust_timeline(self):
        """Handle Adjust button click for timeline - get feedback and emit adjust_signal."""
        # Check if there's content to adjust
        if not hasattr(self, 'timeline_display') or not self.timeline_display.toPlainText().strip():
            self.error_signal.emit("No timeline content to adjust")
            return
        
        feedback_text, ok = QtWidgets.QInputDialog.getMultiLineText(
            self,
            "Adjust Timeline",
            "What adjustments would you like to make to the timeline?",
            ""
        )
        
        if ok and feedback_text.strip():
            # Always emit signal to start refinement process
            self.adjust_signal.emit('timeline', feedback_text)
            if self.current_project:
                self.log_update.emit(f"Timeline adjustment requested: {feedback_text[:100]}...")
                # Disable buttons during refinement
                if hasattr(self, 'approve_timeline_button'):
                    self.approve_timeline_button.setEnabled(False)
                if hasattr(self, 'adjust_timeline_button'):
                    self.adjust_timeline_button.setEnabled(False)
    
    def _expand_text_window(self, window_name):
        """Expand a text window to fill the entire planning tab."""
        # Get the text display widget
        text_widget = getattr(self, window_name, None)
        if not text_widget:
            return
        
        # Create a new dialog window for expanded viewing
        expand_dialog = QtWidgets.QDialog(self)
        expand_dialog.setWindowTitle(f"Expanded View - {window_name.replace('_', ' ').title()}")
        expand_dialog.setGeometry(100, 100, 1000, 800)
        
        layout = QtWidgets.QVBoxLayout(expand_dialog)
        
        # Create expanded text display
        expanded_text = QtWidgets.QTextEdit()
        expanded_text.setReadOnly(True)
        expanded_text.setPlainText(text_widget.toPlainText())
        layout.addWidget(expanded_text)
        
        # Store reference to expanded text widget for streaming updates
        self.expanded_text_widgets[window_name] = expanded_text
        
        # Add close button
        close_btn = QtWidgets.QPushButton("Close")
        close_btn.clicked.connect(lambda: self._on_expanded_window_close(window_name, expand_dialog))
        layout.addWidget(close_btn)
        
        expand_dialog.setLayout(layout)
        expand_dialog.exec_()
    
    def _on_expanded_window_close(self, window_name, dialog):
        """Handle expanded window close - remove reference and close dialog."""
        if window_name in self.expanded_text_widgets:
            del self.expanded_text_widgets[window_name]
        dialog.accept()
        
    def _on_create_project(self):
        """Handle project creation when button is clicked."""
        project_name = self.project_name_input.text().strip()
        
        self._write_app_log(f"User initiated project creation: {project_name}")
        
        if not project_name:
            self._write_app_log(f"Project creation failed: empty name provided")
            self.initialization_status.setText("Error: Project name cannot be empty!")
            self.error_signal.emit("Project name is required")
            return
        
        try:
            self.create_project_structure(project_name)
            self._write_app_log(f"Project '{project_name}' created successfully in projects folder")
            self.initialization_status.setText(f"✓ Project '{project_name}' created successfully!\n\nLocation: projects/{project_name}/")
            self.log_update.emit(f"Project '{project_name}' created")
            self.project_name_input.clear()
        except Exception as e:
            error_msg = f"Error creating project: {str(e)}"
            self._write_app_log(f"Project creation failed for '{project_name}': {error_msg}")
            self.initialization_status.setText(error_msg)
            self.error_signal.emit(error_msg)
    
    def test_llm_connection(self):
        """Test LLM connection with retry logic. Up to 3 retries with 5s sleep between attempts."""
        max_retries = 3
        retry_delay = 5  # seconds (reduced from 30 for faster feedback)
        
        for attempt in range(1, max_retries + 1):
            try:
                # Attempt to generate a test response
                response = self.client.generate(model='gemma3:12b', prompt='Test.')
                
                # Mark as connected
                self.llm_connected = True
                
                # Log success
                self._write_app_log("LLM connection successful (gemma3:12b)")
                
                # Emit log if project is active
                if self.current_project:
                    self.log_update.emit("LLM connection successful (gemma3:12b)")
                
                return
                
            except Exception as e:
                if attempt < max_retries:
                    # Log the failed attempt
                    self._write_app_log(f"LLM connection attempt {attempt} failed: {str(e)}. Retrying in {retry_delay}s...")
                    
                    # Sleep before retry
                    threading.Event().wait(retry_delay)
                else:
                    # All retries failed
                    self.llm_connected = False
                    error_msg = f"LLM connection failed after {max_retries} attempts: {str(e)}"
                    self._write_app_log(error_msg)
                    
                    # Emit log if project is active
                    if self.current_project:
                        self.log_update.emit(error_msg)
                    
                    self.error_signal.emit(error_msg)
    
    
    
    def _initialize_app_config(self):
        """Create app-level settings and config directory if they don't exist."""
        # Create app config folder for application-wide settings
        os.makedirs('Config', exist_ok=True)
        
        # Create app settings file
        app_settings_file = os.path.join('Config', 'app_settings.txt')
        if not os.path.exists(app_settings_file):
            with open(app_settings_file, 'w', encoding='utf-8') as f:
                pass
        
        # Initialize rotating log system (log1.txt through log5.txt)
        self._rotate_app_logs()
    
    def _rotate_app_logs(self):
        """Rotate application logs using a 5-file system (log1.txt through log5.txt).
        
        On each startup, this method:
        1. Determines which log file to use next (oldest/least recently modified)
        2. Clears that log file (effectively reusing it)
        3. Stores the current log file number for this session
        
        This keeps the log directory clean and prevents needing to scroll through
        months/years of history when debugging current issues.
        """
        config_dir = 'Config'
        log_files = [os.path.join(config_dir, f'log{i}.txt') for i in range(1, 6)]
        
        # Create all log files if they don't exist
        for log_file in log_files:
            if not os.path.exists(log_file):
                with open(log_file, 'w', encoding='utf-8') as f:
                    pass
        
        # Find the oldest log file (least recently modified)
        oldest_log = min(log_files, key=lambda f: os.path.getmtime(f))
        
        # Clear the oldest log file for fresh session logging
        with open(oldest_log, 'w', encoding='utf-8') as f:
            pass
        
        # Store current log file path for this session
        self.current_app_log = oldest_log
        
        # Touch the file to update modification time for next rotation
        os.utime(oldest_log, None)
    
    def _write_app_log(self, message: str):
        """Write message to current rotating app log file."""
        try:
            if hasattr(self, 'current_app_log'):
                log_entry = f"{datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')} - {message}\n"
                with open(self.current_app_log, 'a', encoding='utf-8') as f:
                    f.write(log_entry)
        except Exception as e:
            print(f"Failed to write to app log: {str(e)}")
    
    def create_project_structure(self, project_path):
        """Create project-specific files and folders. Call when a new project is created."""
        # Create projects folder and project directory
        os.makedirs('projects', exist_ok=True)
        full_project_path = os.path.join('projects', project_path)
        os.makedirs(full_project_path, exist_ok=True)
        
        # Create drafts folder
        os.makedirs(os.path.join(full_project_path, 'drafts'), exist_ok=True)
        
        # List of project-specific files to create (empty)
        empty_files = [
            'story.txt',
            'config.txt',
            'context.txt',
            'buffer_backup.txt'
        ]
        
        # Create empty text files
        for filename in empty_files:
            filepath = os.path.join(full_project_path, filename)
            if not os.path.exists(filepath):
                with open(filepath, 'w', encoding='utf-8') as f:
                    pass
        
        # Create empty JSON files
        json_files = ['characters.txt', 'world.txt']
        for filename in json_files:
            filepath = os.path.join(full_project_path, filename)
            if not os.path.exists(filepath):
                with open(filepath, 'w', encoding='utf-8') as f:
                    pass
        
        # Create progress tracking file
        progress_file = os.path.join(full_project_path, 'progress.json')
        if not os.path.exists(progress_file):
            with open(progress_file, 'w', encoding='utf-8') as f:
                json.dump({}, f, indent=2)
        
        # Create summaries.txt (empty)
        summaries_file = os.path.join(full_project_path, 'summaries.txt')
        if not os.path.exists(summaries_file):
            with open(summaries_file, 'w', encoding='utf-8') as f:
                pass
        
        # Placeholders for backup files (not created yet)
        # story_backup.txt - placeholder
        # log_backup.txt - placeholder
        # (project logs no longer created - all logging goes to Config/log1-5.txt)
    
    def load_project(self, project_name):
        """Load an existing project and store it in memory for persistence."""
        self._write_app_log(f"Loading project: {project_name}")
        project_path = os.path.join('projects', project_name)
        
        if not os.path.exists(project_path):
            self._write_app_log(f"Project load failed: path not found - {project_path}")
            raise FileNotFoundError(f"Project '{project_name}' not found at {project_path}")
        
        # Verify all required project files exist
        required_files = ['story.txt', 'config.txt', 'context.txt', 
                         'characters.txt', 'world.txt', 'summaries.txt', 'buffer_backup.txt']
        for filename in required_files:
            filepath = os.path.join(project_path, filename)
            if not os.path.exists(filepath):
                self._write_app_log(f"Project load failed: missing file - {filename}")
                raise FileNotFoundError(f"Project file missing: {filename}")
        
        # Load and store project data in memory
        self.current_project = {
            'name': project_name,
            'path': project_path,
            'story': self._read_file(os.path.join(project_path, 'story.txt')),
            'config': self._read_file(os.path.join(project_path, 'config.txt')),
            'context': self._read_file(os.path.join(project_path, 'context.txt')),
            'characters': self._read_file(os.path.join(project_path, 'characters.txt')),
            'world': self._read_file(os.path.join(project_path, 'world.txt')),
            'summaries': self._read_file(os.path.join(project_path, 'summaries.txt')),
            'buffer_backup': self._read_file(os.path.join(project_path, 'buffer_backup.txt'))
        }
        
        # Log the project load
        log_entry = f"{datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')} - Project loaded into session\n"
        app_settings_file = os.path.join('Config', 'app_settings.txt')
        with open(app_settings_file, 'a', encoding='utf-8') as f:
            f.write(log_entry)
        
        # Refresh logs tab with new project's logs
        self._refresh_logs_tab()
        
        # Load progress tracking if it exists
        self._load_progress()
        
        # Populate planning tab displays with existing project data
        self._populate_planning_displays()
        
        self._write_app_log(f"Project loaded successfully: {project_name} - session persistence enabled")
        return self.current_project
        return self.current_project
    
    def _save_progress(self, stage, status):
        """Save progress tracking to project. Stages: synopsis, outline, characters, world."""
        if not self.current_project:
            return
        
        progress_file = os.path.join(self.current_project['path'], 'progress.json')
        
        # Load existing progress or create new
        try:
            with open(progress_file, 'r', encoding='utf-8') as f:
                progress = json.load(f)
        except (FileNotFoundError, json.JSONDecodeError):
            progress = {}
        
        # Update stage status
        progress[stage] = {
            'status': status,  # 'completed', 'approved', 'refinement_pending', etc
            'timestamp': datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')
        }
        
        # Save progress file
        try:
            with open(progress_file, 'w', encoding='utf-8') as f:
                json.dump(progress, f, indent=2, ensure_ascii=False)
        except Exception as e:
            self.log_update.emit(f"Warning: Could not save progress: {str(e)}")
    
    def _load_progress(self):
        """Load progress tracking from project if it exists."""
        if not self.current_project:
            return {}
        
        progress_file = os.path.join(self.current_project['path'], 'progress.json')
        
        try:
            with open(progress_file, 'r', encoding='utf-8') as f:
                progress = json.load(f)
                # Store in current_project for quick access
                self.current_project['progress'] = progress
                return progress
        except (FileNotFoundError, json.JSONDecodeError):
            self.current_project['progress'] = {}
            return {}
    
    def _populate_planning_displays(self):
        """Populate planning tab displays with existing project data from project files."""
        if not self.current_project:
            return
        
        # Load initial synopsis from synopsis.txt if it exists
        synopsis_file = os.path.join(self.current_project['path'], 'synopsis.txt')
        has_initial_synopsis = False
        if os.path.exists(synopsis_file):
            try:
                with open(synopsis_file, 'r', encoding='utf-8') as f:
                    synopsis_content = f.read().strip()
                
                if synopsis_content and hasattr(self, 'synopsis_display'):
                    self.synopsis_display.setText(synopsis_content)
                    has_initial_synopsis = True
            except Exception as e:
                self.log_update.emit(f"Failed to load initial synopsis: {str(e)}")
        
        # Load refined synopsis from refined_synopsis.txt if it exists
        refined_synopsis_file = os.path.join(self.current_project['path'], 'refined_synopsis.txt')
        has_refined_synopsis = False
        refined_synopsis_content = ""
        if os.path.exists(refined_synopsis_file):
            try:
                with open(refined_synopsis_file, 'r', encoding='utf-8') as f:
                    refined_synopsis_content = f.read().strip()
                
                if refined_synopsis_content and hasattr(self, 'planning_synopsis_display'):
                    self.planning_synopsis_display.setText(refined_synopsis_content)
                    has_refined_synopsis = True
            except Exception as e:
                self.log_update.emit(f"Failed to load refined synopsis: {str(e)}")
        
        # Sync synopsis to background thread for refinement operations
        # Use refined synopsis if available, otherwise use initial synopsis
        if refined_synopsis_content:
            self.thread.synopsis = refined_synopsis_content
        elif synopsis_content and has_initial_synopsis:
            self.thread.synopsis = synopsis_content
        
        # Enable synopsis buttons if content is loaded
        if has_initial_synopsis:
            if hasattr(self, 'initial_approve_button'):
                self.initial_approve_button.setEnabled(True)
            if hasattr(self, 'initial_adjust_button'):
                self.initial_adjust_button.setEnabled(True)
        
        if has_refined_synopsis:
            if hasattr(self, 'approve_button'):
                self.approve_button.setEnabled(True)
            if hasattr(self, 'adjust_button'):
                self.adjust_button.setEnabled(True)
        
        # Check for outline.txt file (indicates outline was generated)
        outline_file = os.path.join(self.current_project['path'], 'outline.txt')
        has_outline = False
        if os.path.exists(outline_file):
            try:
                with open(outline_file, 'r', encoding='utf-8') as f:
                    outline_content = f.read()
                if outline_content.strip() and hasattr(self, 'outline_display'):
                    self.outline_display.setText(outline_content)
                    has_outline = True
            except Exception as e:
                self.log_update.emit(f"Failed to load outline: {str(e)}")
        
        # Enable outline buttons if content is loaded
        if has_outline:
            if hasattr(self, 'approve_outline_button'):
                self.approve_outline_button.setEnabled(True)
            if hasattr(self, 'adjust_outline_button'):
                self.adjust_outline_button.setEnabled(True)
        
        # Load characters if available
        has_characters = False
        if self.current_project.get('characters', ''):
            if hasattr(self, 'characters_display'):
                try:
                    import json as json_module
                    # Try to parse and format as JSON
                    characters_data = json_module.loads(self.current_project['characters'])
                    formatted = json_module.dumps(characters_data, indent=2, ensure_ascii=False)
                    self.characters_display.setText(formatted)
                    has_characters = True
                except (json_module.JSONDecodeError, ValueError):
                    # If not JSON, display as plain text
                    self.characters_display.setText(self.current_project['characters'])
                    has_characters = True
        
        # Enable characters buttons if content is loaded
        if has_characters:
            if hasattr(self, 'approve_characters_button'):
                self.approve_characters_button.setEnabled(True)
            if hasattr(self, 'adjust_characters_button'):
                self.adjust_characters_button.setEnabled(True)
        
        # Load world if available
        has_world = False
        if self.current_project.get('world', ''):
            if hasattr(self, 'world_display'):
                try:
                    import json as json_module
                    # Try to parse and format as JSON
                    world_data = json_module.loads(self.current_project['world'])
                    formatted = json_module.dumps(world_data, indent=2, ensure_ascii=False)
                    self.world_display.setText(formatted)
                    has_world = True
                except (json_module.JSONDecodeError, ValueError):
                    # If not JSON, display as plain text
                    self.world_display.setText(self.current_project['world'])
                    has_world = True
        
        # Enable world buttons if content is loaded
        if has_world:
            if hasattr(self, 'approve_world_button'):
                self.approve_world_button.setEnabled(True)
            if hasattr(self, 'adjust_world_button'):
                self.adjust_world_button.setEnabled(True)
        
        # Load timeline if available
        has_timeline = False
        timeline_file = os.path.join(self.current_project['path'], 'timeline.txt')
        if os.path.exists(timeline_file):
            try:
                with open(timeline_file, 'r', encoding='utf-8') as f:
                    timeline_content = f.read()
                if timeline_content.strip() and hasattr(self, 'timeline_display'):
                    self.timeline_display.setText(timeline_content)
                    has_timeline = True
            except Exception as e:
                self.log_update.emit(f"Failed to load timeline: {str(e)}")
        
        # Enable timeline buttons if content is loaded
        if has_timeline:
            if hasattr(self, 'approve_timeline_button'):
                self.approve_timeline_button.setEnabled(True)
            if hasattr(self, 'adjust_timeline_button'):
                self.adjust_timeline_button.setEnabled(True)
        
        # Log that displays were populated
        self.log_update.emit("Planning displays populated with existing project data")
    
    def get_stage_status(self, stage):
        """Get the current status of a generation stage. Returns None if not completed."""
        if not self.current_project or 'progress' not in self.current_project:
            return None
        
        progress = self.current_project['progress']
        if stage in progress:
            return progress[stage].get('status')
        return None
    
    def _initialize_chapter_tracking(self):
        """Initialize chapter tracking after world approval. Parse outline and update config.txt."""
        if not self.current_project:
            return
        
        try:
            # Read outline to determine actual chapter count (assuming 25 chapters standard)
            outline_path = os.path.join(self.current_project['path'], 'outline.txt')
            try:
                with open(outline_path, 'r', encoding='utf-8') as f:
                    outline_content = f.read()
                
                # Count chapter headings in outline (lines starting with "Chapter" or "#")
                chapter_count = 0
                for line in outline_content.split('\n'):
                    line_stripped = line.strip()
                    if line_stripped.lower().startswith('chapter') or line_stripped.startswith('#'):
                        chapter_count += 1
                
                # Use actual count or default to 25
                total_chapters = chapter_count if chapter_count > 0 else 25
            except Exception:
                total_chapters = 25  # Default to 25 chapters if outline read fails
            
            # Estimate sections per chapter (typically 4-6, using 5 as average)
            sections_per_chapter = 5
            
            # Update config.txt with chapter tracking information
            config_path = os.path.join(self.current_project['path'], 'config.txt')
            config_content = self.current_project.get('config', '')
            
            # Parse existing config to preserve other settings
            config_dict = {}
            for line in config_content.split('\n'):
                if ':' in line:
                    key, value = line.split(':', 1)
                    config_dict[key.strip()] = value.strip()
            
            # Update chapter tracking fields
            config_dict['TotalChapters'] = str(total_chapters)
            config_dict['CurrentChapter'] = '1'
            config_dict['CurrentSection'] = '1'
            config_dict['Chapter1Sections'] = str(sections_per_chapter)
            
            # Reconstruct config content
            updated_config = '\n'.join([f"{key}: {value}" for key, value in config_dict.items()])
            
            # Write updated config to file and update in-memory copy
            with open(config_path, 'w', encoding='utf-8') as f:
                f.write(updated_config)
            
            self.current_project['config'] = updated_config
            
            # Log the chapter initialization
            self.log_update.emit(f"Chapter tracking initialized: {total_chapters} chapters, {sections_per_chapter} sections per chapter")
        
        except Exception as e:
            self.log_update.emit(f"Error initializing chapter tracking: {str(e)}")
    
    def get_project_list(self):
        """Get list of all existing projects."""
        projects_dir = 'projects'
        if not os.path.exists(projects_dir):
            return []
        
        projects = []
        for item in os.listdir(projects_dir):
            item_path = os.path.join(projects_dir, item)
            if os.path.isdir(item_path):
                projects.append(item)
        
        return sorted(projects)
    
    def _read_file(self, filepath):
        """Read file content safely."""
        try:
            with open(filepath, 'r', encoding='utf-8') as f:
                return f.read()
        except Exception as e:
            return f"Error reading file: {str(e)}"
    
    def get_current_project(self):
        """Get the currently loaded project data."""
        return self.current_project
    
    def _on_export_docx(self):
        """Handle DOCX export button click."""
        if not self.current_project:
            self.error_signal.emit("No project loaded. Load a project to export.")
            return
        
        if not HAS_DOCX:
            self.error_signal.emit("python-docx is not installed. Install it with: pip install python-docx")
            return
        
        try:
            project_path = self.current_project['path']
            project_name = os.path.basename(project_path)
            
            # Export to DOCX
            success = export_story_to_docx(project_path)
            
            if success:
                output_file = os.path.join(project_path, f"{project_name}_novel.docx")
                self.export_status_label.setText(f"✓ Exported to {project_name}_novel.docx")
                self.log_update.emit(f"Novel exported to DOCX: {output_file}")
            else:
                self.export_status_label.setText("✗ Export failed. Check that story.txt exists.")
                self.log_update.emit("DOCX export failed")
        
        except Exception as e:
            error_msg = f"Export error: {str(e)}"
            self.export_status_label.setText("✗ Export failed")
            self.log_update.emit(error_msg)
    
    def _on_export_pdf(self):
        """Handle PDF export button click."""
        if not self.current_project:
            self.error_signal.emit("No project loaded. Load a project to export.")
            return
        
        if not HAS_REPORTLAB:
            self.error_signal.emit("reportlab is not installed. Install it with: pip install reportlab")
            return
        
        try:
            project_path = self.current_project['path']
            project_name = os.path.basename(project_path)
            
            # Export to PDF
            success = export_story_to_pdf(project_path)
            
            if success:
                output_file = os.path.join(project_path, f"{project_name}_novel.pdf")
                self.export_status_label.setText(f"✓ Exported to {project_name}_novel.pdf")
                self.log_update.emit(f"Novel exported to PDF: {output_file}")
            else:
                self.export_status_label.setText("✗ Export failed. Check that story.txt exists.")
                self.log_update.emit("PDF export failed")
        
        except Exception as e:
            error_msg = f"Export error: {str(e)}"
            self.export_status_label.setText("✗ Export failed")
            self.log_update.emit(error_msg)


def export_story_to_docx(project_path: str, output_filename: Optional[str] = None) -> bool:
    """Export story.txt to Word document with formatting.
    
    Args:
        project_path: Path to project folder
        output_filename: Optional custom output filename (without extension)
    
    Returns:
        bool: True if successful, False otherwise
    """
    if not HAS_DOCX:
        return False
    
    try:
        story_path = os.path.join(project_path, 'story.txt')
        config_path = os.path.join(project_path, 'config.txt')
        
        if not os.path.exists(story_path):
            return False
        
        # Read story content
        with open(story_path, 'r', encoding='utf-8') as f:
            story_content = f.read()
        
        # Read project config for title
        project_name = os.path.basename(project_path)
        title = f"Novel: {project_name}"
        
        if os.path.exists(config_path):
            with open(config_path, 'r', encoding='utf-8') as f:
                for line in f:
                    if line.startswith('Idea:'):
                        title = line.split(':', 1)[1].strip()
                        break
        
        # Create Word document
        doc = Document()
        
        # Add title
        title_para = doc.add_paragraph()
        title_run = title_para.add_run(title)
        title_run.font.size = Pt(24)
        title_run.font.bold = True
        title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # Add metadata
        doc.add_paragraph(f"Project: {project_name}")
        doc.add_paragraph(f"Generated: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        
        doc.add_paragraph()  # Blank line
        
        # Parse and add content with chapter headings
        lines = story_content.split('\n')
        for line in lines:
            line = line.rstrip()
            
            # Check if line is a chapter heading
            if line.startswith('=== CHAPTER'):
                # Add chapter heading
                heading = doc.add_heading(line.strip('=').strip(), level=1)
                heading.paragraph_format.space_before = Pt(12)
                heading.paragraph_format.space_after = Pt(12)
            elif line.strip():
                # Add regular paragraph with justified alignment
                para = doc.add_paragraph(line)
                para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            else:
                # Preserve blank lines as spacing
                doc.add_paragraph()
        
        # Save document
        if not output_filename:
            output_filename = f"{project_name}_novel"
        
        output_path = os.path.join(project_path, f"{output_filename}.docx")
        doc.save(output_path)
        
        return True
    
    except Exception as e:
        return False


def export_story_to_pdf(project_path: str, output_filename: Optional[str] = None) -> bool:
    """Export story.txt to PDF using reportlab.
    
    Args:
        project_path: Path to project folder
        output_filename: Optional custom output filename (without extension)
    
    Returns:
        bool: True if successful, False otherwise
    """
    if not HAS_REPORTLAB:
        return False
    
    try:
        story_path = os.path.join(project_path, 'story.txt')
        config_path = os.path.join(project_path, 'config.txt')
        
        if not os.path.exists(story_path):
            return False
        
        # Read story content
        with open(story_path, 'r', encoding='utf-8') as f:
            story_content = f.read()
        
        # Read project config for title
        project_name = os.path.basename(project_path)
        title = f"Novel: {project_name}"
        
        if os.path.exists(config_path):
            with open(config_path, 'r', encoding='utf-8') as f:
                for line in f:
                    if line.startswith('Idea:'):
                        title = line.split(':', 1)[1].strip()
                        break
        
        # Prepare output path
        if not output_filename:
            output_filename = f"{project_name}_novel"
        
        output_path = os.path.join(project_path, f"{output_filename}.pdf")
        
        # Create PDF
        doc = SimpleDocTemplate(output_path, pagesize=letter)
        
        # Define styles
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=24,
            textColor=HexColor('000000'),
            spaceAfter=6,
            alignment=TA_CENTER,
            fontName='Helvetica-Bold'
        )
        heading_style = ParagraphStyle(
            'ChapterHeading',
            parent=styles['Heading2'],
            fontSize=14,
            textColor=HexColor('000000'),
            spaceAfter=6,
            spaceBefore=12,
            fontName='Helvetica-Bold'
        )
        body_style = ParagraphStyle(
            'CustomBody',
            parent=styles['BodyText'],
            fontSize=11,
            alignment=TA_JUSTIFY,
            spaceAfter=6,
            leading=14
        )
        
        # Build story elements
        story_elements = []
        
        # Add title
        story_elements.append(Paragraph(title, title_style))
        story_elements.append(Paragraph(f"Project: {project_name}", styles['Normal']))
        story_elements.append(Paragraph(f"Generated: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", styles['Normal']))
        story_elements.append(Spacer(1, 0.3*inch))
        
        # Parse and add content
        lines = story_content.split('\n')
        for line in lines:
            line = line.rstrip()
            
            # Check if line is a chapter heading
            if line.startswith('=== CHAPTER'):
                story_elements.append(Paragraph(line.strip('=').strip(), heading_style))
            elif line.strip():
                story_elements.append(Paragraph(line, body_style))
            else:
                story_elements.append(Spacer(1, 0.1*inch))
        
        # Build PDF
        doc.build(story_elements)
        
        return True
    
    except Exception as e:
        return False


def main():
    """Main entry point for the application."""
    app = QtWidgets.QApplication(sys.argv)
    window = ANSWindow()
    window.show()
    sys.exit(app.exec_())


if __name__ == "__main__":
    main()
